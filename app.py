# -*- coding: utf-8 -*-
"""
Individual Behavior Support Plan Generator
============================================
Streamlit app for nursery/child development professionals.
Generates premium PDF and Word (.docx) reports.
"""

import streamlit as st
import io
import base64
import json
import requests
from datetime import datetime, date
from PIL import Image


# -----------------------------------------------------------------------------
# AI GENERATION (Groq)
# -----------------------------------------------------------------------------
def _groq_call(prompt: str, system: str, max_tokens: int = 1200) -> str:
    """Call Groq API and return the text response, or empty string on failure."""
    try:
        api_key = st.secrets.get("GROQ_API_KEY", "")
        if not api_key:
            return ""
        resp = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": "llama-3.3-70b-versatile",
                "max_tokens": max_tokens,
                "temperature": 0.55,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user",   "content": prompt},
                ],
            },
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"].strip()
    except Exception:
        return ""


@st.cache_data(show_spinner=False, ttl=3600)
def generate_dynamic_activities(skills_tuple: tuple, behaviors_tuple: tuple, child_name: str) -> list:
    """
    Return a list of (title, description) tuples for Session Activities,
    tailored to the selected skills and behaviors using Groq AI.
    Falls back to a smart rule-based list if AI is unavailable.
    """
    skills = list(skills_tuple)
    behaviors = list(behaviors_tuple)

    system = (
        "You are a senior behavioral therapist writing clinical session activity descriptions "
        "for a child behavior support plan used in a nursery/early childhood setting. "
        "Your language is professional, evidence-informed, and warm. "
        "You write in British/international English. "
        "Respond ONLY with a JSON array — no preamble, no markdown fences. "
        "Each element: {\"title\": \"...\", \"description\": \"...\"}. "
        "Titles are 3-6 words. Descriptions are 1-2 sentences, clinical and specific."
    )

    skills_str = ", ".join(skills) if skills else "general developmental support"
    behaviors_str = ", ".join(behaviors) if behaviors else "general behavioral concerns"

    prompt = (
        f"Child: {child_name or 'the child'}\n"
        f"Target skills: {skills_str}\n"
        f"Identified behaviors: {behaviors_str}\n\n"
        f"Generate exactly 6 session activity types that a behavioral therapist would use "
        f"to address THESE SPECIFIC skills and behaviors. "
        f"Each activity must directly reference or logically address at least one of the listed skills or behaviors. "
        f"Do not generate generic activities — make every activity clearly connected to the listed targets."
    )

    raw = _groq_call(prompt, system, max_tokens=900)
    if raw:
        try:
            raw_clean = raw.strip()
            if raw_clean.startswith("```"):
                raw_clean = "\n".join(raw_clean.split("\n")[1:])
                raw_clean = raw_clean.split("```")[0].strip()
            items = json.loads(raw_clean)
            result = [(i["title"], i["description"]) for i in items if "title" in i and "description" in i]
            if len(result) >= 4:
                return result
        except Exception:
            pass

    # ---- Rule-based fallback ----
    return _build_fallback_activities(skills, behaviors)


@st.cache_data(show_spinner=False, ttl=3600)
def generate_dynamic_progress(skills_tuple: tuple, behaviors_tuple: tuple, child_name: str) -> list:
    """
    Return a list of (icon, title, description) tuples for Indicators of Progress,
    tailored to the selected skills and behaviors using Groq AI.
    Falls back to a smart rule-based list if AI is unavailable.
    """
    skills = list(skills_tuple)
    behaviors = list(behaviors_tuple)

    system = (
        "You are a senior behavioral therapist writing clinical progress indicators "
        "for a child behavior support plan in a nursery/early childhood setting. "
        "Your language is professional, evidence-informed, and warm. "
        "You write in British/international English. "
        "Respond ONLY with a JSON array — no preamble, no markdown fences. "
        "Each element: {\"icon\": \"one emoji\", \"title\": \"...\", \"description\": \"...\"}. "
        "Titles are 3-6 words. Descriptions are 1-2 sentences, specific and observable."
    )

    skills_str = ", ".join(skills) if skills else "general developmental support"
    behaviors_str = ", ".join(behaviors) if behaviors else "general behavioral concerns"

    prompt = (
        f"Child: {child_name or 'the child'}\n"
        f"Target skills: {skills_str}\n"
        f"Identified behaviors: {behaviors_str}\n\n"
        f"Generate exactly 6 observable indicators of progress that a therapist would look for "
        f"when these specific skills improve and these specific behaviors reduce. "
        f"Each indicator must be directly linked to the listed skills or behaviors — "
        f"describe what success actually looks like for THIS child's targets. "
        f"Use a relevant single emoji icon per indicator."
    )

    raw = _groq_call(prompt, system, max_tokens=900)
    if raw:
        try:
            raw_clean = raw.strip()
            if raw_clean.startswith("```"):
                raw_clean = "\n".join(raw_clean.split("\n")[1:])
                raw_clean = raw_clean.split("```")[0].strip()
            items = json.loads(raw_clean)
            result = [(i.get("icon", "✦"), i["title"], i["description"])
                      for i in items if "title" in i and "description" in i]
            if len(result) >= 4:
                return result
        except Exception:
            pass

    # ---- Rule-based fallback ----
    return _build_fallback_progress(skills, behaviors)


# ---- Fallback generators (no AI required) ------------------------------------

_SKILL_ACTIVITY_MAP = {
    "Emotional Regulation":     ("Emotion regulation practice",
        "Structured use of co-regulation tools — breathing techniques, calm corners, and emotion cards — "
        "to help the child identify and manage emotional states with gradually decreasing adult support."),
    "Impulse Control":          ("Inhibitory control games",
        "Stop-and-go games, freeze activities, and delayed-response tasks that train the neural pathways "
        "underlying the ability to pause before acting on an impulse."),
    "Frustration Tolerance":    ("Structured frustration exposure",
        "Controlled introduction of manageable difficulty within preferred activities to safely expand "
        "the child's window of tolerance for disappointment and waiting."),
    "Communication Skills":     ("Replacement behavior rehearsal",
        "Explicit teaching and repeated practice of verbal and non-verbal communication alternatives "
        "for each identified challenging behavior, reducing the need for behavioral escalation."),
    "Social Skills":            ("Peer interaction scaffolding",
        "Structured cooperative activities and role-play scenarios designed to build turn-taking, "
        "conflict resolution, and collaborative engagement with familiar peers."),
    "Self-Awareness":           ("Reflective emotional dialogue",
        "Guided post-activity reflection conversations that build the child's vocabulary for internal "
        "states and develop insight into the connection between triggers and responses."),
    "Coping Strategies":        ("Personal coping toolkit building",
        "Collaborative construction and practice of an individualised toolkit of coping strategies "
        "— sensory tools, breathing cards, and calming rituals — accessible independently."),
    "Attention and Concentration": ("Graduated attention training",
        "Progressive focus-and-listen activities with systematically increasing duration to build "
        "sustained on-task engagement and resistance to environmental distractions."),
    "Decision Making":          ("Choice-making and consequence exploration",
        "Structured decision scenarios using visual choice boards and guided reflection on outcomes, "
        "building the child's confidence and capacity for independent problem-solving."),
    "Self-Boundaries":          ("Body safety and boundary role-play",
        "Age-appropriate role-play using puppets and scenarios to develop recognition of personal "
        "space, assertive communication of limits, and respect for others' boundaries."),
    "Hyperactivity Reduction":  ("Energy regulation and dial activities",
        "Structured movement-release followed by guided calm-down practice using the Energy Dial "
        "framework, building the child's self-awareness of arousal level and modulation strategies."),
    "Basic Knowledge":          ("Concept consolidation through play",
        "Structured discovery activities targeting colours, shapes, numbers, and categories within "
        "a reinforcement-rich, play-based framework to build foundational academic readiness."),
    "Emotional Resilience":     ("Bounce-back story and reflection",
        "Shared narratives and guided discussion around characters facing setbacks, building the child's "
        "growth mindset and capacity for emotional recovery after disappointment."),
    "Independence and Self-Care": ("Task-analysed self-care practice",
        "Step-by-step practice of a targeted self-care skill using visual schedules and a systematic "
        "prompt-fading approach toward full independence."),
    "Transitions and Flexibility": ("Visual schedule transition practice",
        "Structured within-session transitions using countdown timers and advance warnings to build "
        "tolerance for activity changes and unexpected routine shifts."),
    "Play Skills":              ("Imaginative and cooperative play scaffolding",
        "Child-led play extended by the therapist through collaborative storytelling and shared goals, "
        "developing play complexity, narrative language, and peer engagement capacity."),
}

_BEHAVIOR_ACTIVITY_MAP = {
    "Throwing objects":             ("Object-use replacement rehearsal",
        "Explicit practice of safe, appropriate alternative responses to the throwing trigger, "
        "using role-play and immediate reinforcement of functionally equivalent replacement behaviors."),
    "Hitting / kicking others":     ("Aggression replacement training",
        "Teaching specific verbal and physical alternatives to hitting and kicking within structured "
        "low-stakes scenarios, with consistent reinforcement of non-aggressive responses."),
    "Tantrums / meltdowns":         ("Dysregulation de-escalation practice",
        "Rehearsal of early self-regulation strategies at the first signs of escalation, with gradual "
        "reduction of adult support needed to return to a regulated baseline."),
    "Biting":                       ("Oral replacement behavior practice",
        "Identification and rehearsal of safe oral alternatives and communication replacements for "
        "biting triggers, paired with consistent reinforcement of appropriate responses."),
    "Screaming / shouting":         ("Voice regulation and communication practice",
        "Explicit teaching of volume awareness and appropriate verbal expression strategies as "
        "replacements for screaming, using visual volume scales and immediate reinforcement."),
    "Refusing transitions":         ("Transition preparation and acceptance",
        "Structured use of advance warnings, visual schedules, and countdown tools to reduce "
        "distress and increase compliance during activity transitions."),
    "Running away / elopement":     ("Safe boundary and stop-signal rehearsal",
        "Teaching and practicing a reliable stop response to adult signals within safe, "
        "controlled environments, with systematic reinforcement of boundary compliance."),
    "Potty training difficulties":  ("Toileting routine and independence support",
        "Consistent implementation of a structured toileting schedule with visual supports, "
        "positive reinforcement of steps completed, and systematic prompt fading."),
    "Rule following challenges":    ("Rule rehearsal and compliance practice",
        "Explicit teaching of specific rules within structured game contexts, using visual rule "
        "cards and reinforcement to build generalised compliance across settings."),
    "Self-injurious behavior":      ("Safe hands and self-protection rehearsal",
        "Teaching safe-hands strategies and alternative sensory responses as replacements for "
        "self-injurious behaviors, with close monitoring and immediate reinforcement."),
    "Verbal aggression":            ("Assertive communication training",
        "Explicit practice of appropriate verbal expression strategies as replacements for "
        "verbally aggressive responses, with reinforcement of calm, assertive language."),
    "Dropping to the floor":        ("Floor-dropping replacement rehearsal",
        "Teaching and practicing communicative alternatives (requesting a break, saying 'no') "
        "as replacements for floor-dropping, with reinforcement of each approximation."),
    "Interrupting / difficulty waiting for turn": ("Turn-taking and waiting practice",
        "Structured turn-taking games with visual timers that build the child's capacity to wait "
        "and take turns, with reinforcement of each successful wait interval."),
    "Separation anxiety behaviors":  ("Separation tolerance and transition support",
        "Graduated separation practice with predictable goodbye routines, visual reassurance tools, "
        "and reinforcement of increasing calm during separation transitions."),
    "Difficulty sharing / possessive behavior": ("Sharing and joint attention practice",
        "Structured cooperative activities that require sharing materials, with reinforcement "
        "of sharing behavior and guided processing of feelings around possession."),
}

_SKILL_PROGRESS_MAP = {
    "Emotional Regulation":     ("🌡", "Improved emotional self-regulation",
        "An observable reduction in the intensity and duration of emotional dysregulation episodes, "
        "with the child beginning to use regulation tools with decreasing adult support."),
    "Impulse Control":          ("⏱", "Increased response latency",
        "A measurable pause between trigger and behavioral response, indicating developing inhibitory "
        "control and the capacity to reflect before acting."),
    "Frustration Tolerance":    ("📉", "Broader tolerance window",
        "The child engaging with frustrating tasks for longer durations without escalating, "
        "and recovering more quickly when disappointment occurs."),
    "Communication Skills":     ("💬", "Increased functional communication",
        "Higher frequency of verbal or non-verbal communication to express needs, particularly "
        "in situations that previously triggered challenging behavior."),
    "Social Skills":            ("🤝", "Improved peer engagement",
        "Observable increase in cooperative play, turn-taking compliance, and use of words "
        "to negotiate or resolve conflict with familiar peers."),
    "Self-Awareness":           ("🪞", "Growing emotional literacy",
        "The child using feeling words spontaneously or in response to gentle prompting, "
        "and beginning to connect physical sensations to emotional states."),
    "Coping Strategies":        ("✋", "Spontaneous coping tool use",
        "The child independently reaching for a coping tool — breathing, sensory item, or "
        "calm corner — without adult prompting in moments of distress."),
    "Attention and Concentration": ("🎯", "Extended on-task engagement",
        "Measurably longer periods of focused engagement on structured tasks, with the child "
        "resisting distractions and returning to task independently after interruption."),
    "Decision Making":          ("🧩", "More considered decision-making",
        "Observable slowing of impulsive choices and the child beginning to consider options "
        "and consequences before acting, even in brief and simple scenarios."),
    "Self-Boundaries":          ("🛡", "Assertive boundary communication",
        "The child using words or gestures to communicate personal limits clearly and appropriately, "
        "and respecting the stated limits of peers and adults."),
    "Hyperactivity Reduction":  ("🎚", "Improved arousal self-regulation",
        "The child demonstrating awareness of their energy level and applying calm-down strategies "
        "with decreasing adult prompting across classroom and session contexts."),
    "Basic Knowledge":          ("📚", "Increasing concept mastery",
        "Reliable and independent identification of targeted foundational concepts — colours, shapes, "
        "numbers — with generalisation to new objects and real-life contexts."),
    "Emotional Resilience":     ("🔄", "Faster emotional recovery",
        "A progressive reduction in the time needed to return to a calm, engaged baseline after a "
        "setback, with the child beginning to verbalise recovery strategies."),
    "Independence and Self-Care": ("🌱", "Growing functional independence",
        "Completion of increasing steps within the target self-care skill independently, with a "
        "measurable reduction in the number of adult prompts required."),
    "Transitions and Flexibility": ("🔁", "Smoother activity transitions",
        "A reduction in distress, resistance, and recovery time during transitions between activities, "
        "with the child responding to advance warnings and visual cues with increasing calm."),
    "Play Skills":              ("🎭", "Richer and more sustained play",
        "Longer episodes of engaged, purposeful play with increasing imaginative complexity and "
        "capacity for shared play narratives with a peer or adult partner."),
}

_BEHAVIOR_PROGRESS_MAP = {
    "Throwing objects":             ("📦", "Reduction in throwing incidents",
        "A measurable decrease in the frequency of object-throwing, with the child beginning to "
        "use alternative responses at the point of trigger."),
    "Hitting / kicking others":     ("🤲", "Decrease in physical aggression",
        "Fewer incidents of hitting or kicking, with observable use of verbal alternatives "
        "or adult-seeking behavior in previously triggering situations."),
    "Tantrums / meltdowns":         ("🌊", "Shorter and less intense meltdowns",
        "A reduction in meltdown duration and intensity, with faster return to regulated baseline "
        "and earlier use of co-regulation support."),
    "Screaming / shouting":         ("🔉", "Improved vocal regulation",
        "Reduction in screaming and shouting episodes, with the child using appropriate voice "
        "volume or alternative communication in high-demand situations."),
    "Refusing transitions":         ("🚶", "Improved transition compliance",
        "Smoother acceptance of activity transitions with advance warning, and reduction in refusal "
        "episodes and floor-dropping at transition points."),
    "Running away / elopement":     ("🛑", "Reliable stop response",
        "Consistent response to adult stop signals in familiar environments, with reduction "
        "in elopement incidents and improved boundary awareness."),
    "Potty training difficulties":  ("✅", "Increasing toileting independence",
        "Greater consistency in following the toileting routine, with fewer accidents and "
        "increasing ability to signal toileting needs independently."),
    "Rule following challenges":    ("📋", "Improved rule compliance",
        "Greater consistency in following stated rules within structured contexts, with the child "
        "referencing rule cards independently and requiring fewer reminders."),
    "Self-injurious behavior":      ("💚", "Reduction in self-injury",
        "Fewer incidents of self-injurious behavior, with the child beginning to use safe "
        "alternative responses when distressed, and faster recovery with adult support."),
    "Verbal aggression":            ("🗣", "Calmer verbal expression",
        "Reduction in verbally aggressive episodes, with the child using assertive but appropriate "
        "language to express frustration or disagreement."),
    "Interrupting / difficulty waiting for turn": ("⏳", "Improved waiting capacity",
        "Measurably longer successful waits in structured turn-taking activities, with the child "
        "using verbal expression ('my turn soon') rather than interrupting."),
    "Separation anxiety behaviors":  ("🌤", "Calmer separation transitions",
        "Reduction in distress at separation points, with the child engaging with the predictable "
        "goodbye routine and settling more quickly after separation."),
    "Difficulty sharing / possessive behavior": ("🤲", "Increased sharing behavior",
        "Higher frequency of spontaneous sharing in low-stakes contexts, and calmer acceptance "
        "of turns and shared access to preferred materials."),
}


def _build_fallback_activities(skills: list, behaviors: list) -> list:
    """Build a tailored activity list from skill/behavior maps without AI."""
    seen = set()
    result = []
    # Skill-based
    for s in skills:
        if s in _SKILL_ACTIVITY_MAP and s not in seen:
            result.append(_SKILL_ACTIVITY_MAP[s])
            seen.add(s)
    # Behavior-based (up to filling 6 total)
    for b in behaviors:
        if len(result) >= 6:
            break
        if b in _BEHAVIOR_ACTIVITY_MAP and b not in seen:
            result.append(_BEHAVIOR_ACTIVITY_MAP[b])
            seen.add(b)
    # Generic pad items if needed
    generic = [
        ("Positive reinforcement systems",
         "Consistent, specific, and immediate reinforcement of target behaviors and approximations, "
         "using individually motivating reinforcers identified in collaboration with the child and family."),
        ("Reflective dialogue",
         "Age-appropriate post-activity discussion to build the child's capacity for emotional "
         "retrospection and self-awareness without shame or blame."),
        ("Co-regulation activities",
         "Guided breathing, sensory regulation tools, and calming rituals to help the child "
         "experience regulated states before working toward independent regulation."),
    ]
    for g in generic:
        if len(result) >= 6:
            break
        result.append(g)
    return result[:6] if result else [
        ("Structured supportive play", "Play-based activities tailored to the child's developmental profile."),
        ("Positive reinforcement systems", "Consistent reinforcement of target behaviors."),
    ]


def _build_fallback_progress(skills: list, behaviors: list) -> list:
    """Build a tailored progress indicator list from skill/behavior maps without AI."""
    seen = set()
    result = []
    for s in skills:
        if s in _SKILL_PROGRESS_MAP and s not in seen:
            result.append(_SKILL_PROGRESS_MAP[s])
            seen.add(s)
    for b in behaviors:
        if len(result) >= 6:
            break
        if b in _BEHAVIOR_PROGRESS_MAP and b not in seen:
            result.append(_BEHAVIOR_PROGRESS_MAP[b])
            seen.add(b)
    generic = [
        ("🌍", "Generalisation across settings",
         "Transfer of skill use beyond the formal session context into the broader nursery "
         "environment and, where reported, the home setting."),
        ("✋", "Spontaneous replacement behaviors",
         "The child independently initiating a replacement behavior without adult prompting "
         "in the natural environment."),
    ]
    for g in generic:
        if len(result) >= 6:
            break
        result.append(g)
    return result[:6] if result else [
        ("📉", "Reduced behavioral intensity",
         "A measurable decrease in the severity of behavioral episodes over a consistent time period."),
        ("🔄", "Faster recovery",
         "A progressive reduction in recovery time following dysregulation."),
    ]

# -----------------------------------------------------------------------------
# PAGE CONFIG
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="Behavior Support Plan Generator",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -----------------------------------------------------------------------------
# GLOBAL STYLES
# -----------------------------------------------------------------------------
def inject_styles():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;0,500;0,600;1,300;1,400&family=Jost:wght@300;400;500;600&display=swap');

    :root {
        --sand:    #F5F0E8;
        --linen:   #F0EBE1;
        --cream:   #FAF8F4;
        --bark:    #8B7355;
        --bark-lt: #B09A7A;
        --stone:   #6B6456;
        --ink:     #2A2520;
        --border:  #DDD8CE;
        --sage:    #7A8A72;
        --radius:  12px;
        --shadow:  0 2px 24px rgba(42,37,32,0.08);
    }

    html, body, [class*="css"] {
        font-family: 'Jost', sans-serif !important;
        background-color: var(--sand) !important;
    }

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background: var(--cream) !important;
        border-right: 1px solid var(--border) !important;
    }
    section[data-testid="stSidebar"] .stMarkdown h3 {
        font-family: 'Jost', sans-serif !important;
        font-size: 0.65rem !important;
        font-weight: 600 !important;
        letter-spacing: 0.14em !important;
        text-transform: uppercase !important;
        color: var(--bark) !important;
        margin-top: 1.2rem !important;
        margin-bottom: 0.3rem !important;
    }

    /* Main area */
    .main .block-container {
        padding: 1.5rem 2rem 3rem !important;
        max-width: 100% !important;
    }

    /* Hide default Streamlit chrome */
    #MainMenu, footer, header { visibility: hidden; }

    /* App header */
    .app-header {
        display: flex; align-items: center; gap: 0.75rem;
        padding: 1rem 0 1.5rem;
        border-bottom: 1px solid var(--border);
        margin-bottom: 1.5rem;
    }
    .app-header-mark {
        width: 40px; height: 40px;
        background: var(--ink); border-radius: 10px;
        display: flex; align-items: center; justify-content: center;
        font-size: 1.1rem; color: #fff;
    }
    .app-header-title {
        font-family: 'Cormorant Garamond', serif !important;
        font-size: 1.5rem; font-weight: 500;
        color: var(--ink); line-height: 1.2;
    }
    .app-header-sub { font-size: 0.78rem; color: var(--stone); }

    /* Buttons */
    .stButton > button {
        font-family: 'Jost', sans-serif !important;
        font-weight: 500 !important;
        border-radius: 9px !important;
        transition: all 0.18s !important;
    }
    .stDownloadButton > button {
        font-family: 'Jost', sans-serif !important;
        font-weight: 500 !important;
        border-radius: 9px !important;
        width: 100% !important;
    }

    /* Inputs */
    .stTextInput input, .stTextArea textarea, .stSelectbox select,
    .stNumberInput input, .stDateInput input {
        border-radius: 8px !important;
        font-family: 'Jost', sans-serif !important;
    }

    /* Section dividers in sidebar */
    .sidebar-divider {
        border: none;
        border-top: 1px solid var(--border);
        margin: 0.5rem 0;
    }

    /* Expander styling */
    .streamlit-expanderHeader {
        font-family: 'Jost', sans-serif !important;
        font-weight: 500 !important;
        font-size: 0.88rem !important;
        background: var(--linen) !important;
        border-radius: 8px !important;
    }

    /* Info box */
    .stInfo { border-radius: 10px !important; }

    /* Document preview container */
    .doc-preview-wrap {
        background: #fff;
        border-radius: 4px;
        box-shadow: 0 8px 40px rgba(42,37,32,0.12);
        overflow: hidden;
        max-width: 820px;
        margin: 0 auto;
        font-family: 'Cormorant Garamond', serif;
    }

    .doc-hdr {
        background: #2A2520;
        padding: 2.5rem 3rem 2rem;
        color: #fff;
        position: relative;
    }
    .doc-hdr-top {
        display: flex; justify-content: space-between;
        align-items: flex-start; margin-bottom: 1.75rem;
    }
    .doc-nursery-nm {
        font-family: 'Jost', sans-serif;
        font-size: 0.85rem; font-weight: 300;
        letter-spacing: 0.1em; text-transform: uppercase;
        color: rgba(255,255,255,0.6);
    }
    .doc-badge {
        background: rgba(255,255,255,0.1);
        border: 1px solid rgba(255,255,255,0.2);
        border-radius: 4px; padding: 0.2rem 0.6rem;
        font-family: 'Jost', sans-serif;
        font-size: 0.65rem; letter-spacing: 0.12em;
        text-transform: uppercase; color: rgba(255,255,255,0.5);
    }
    .doc-main-title {
        font-size: 2.1rem; font-weight: 300; line-height: 1.2;
        margin-bottom: 0.3rem;
    }
    .doc-main-sub {
        font-size: 0.95rem; font-weight: 300;
        color: rgba(255,255,255,0.6); font-style: italic;
    }

    .doc-info-bar {
        display: grid; grid-template-columns: repeat(4, 1fr);
        background: #F0EBE1;
        border-bottom: 1px solid #DDD8CE;
    }
    .doc-info-cell {
        padding: 0.9rem 1.1rem;
        border-right: 1px solid #DDD8CE;
    }
    .doc-info-cell:last-child { border-right: none; }
    .doc-info-lbl {
        font-family: 'Jost', sans-serif;
        font-size: 0.6rem; font-weight: 600;
        letter-spacing: 0.12em; text-transform: uppercase;
        color: #8B7355; margin-bottom: 0.2rem;
    }
    .doc-info-val {
        font-size: 0.95rem; font-weight: 400; color: #2A2520;
    }

    .doc-body { padding: 2.2rem 3rem; }

    .doc-sec { margin-bottom: 2rem; }
    .doc-sec-hdr {
        display: flex; align-items: center; gap: 0.65rem;
        margin-bottom: 0.85rem; padding-bottom: 0.5rem;
        border-bottom: 1px solid #DDD8CE;
    }
    .doc-sec-num {
        width: 24px; height: 24px;
        background: #2A2520; border-radius: 50%;
        display: flex; align-items: center; justify-content: center;
        font-family: 'Jost', sans-serif;
        font-size: 0.68rem; font-weight: 600; color: #fff;
        flex-shrink: 0;
    }
    .doc-sec-title {
        font-size: 1.2rem; font-weight: 400; color: #2A2520;
    }

    .doc-p {
        font-size: 0.94rem; line-height: 1.8; color: #3A3530;
        font-weight: 300; margin-bottom: 0.5rem;
    }

    .doc-skill-blk {
        background: #F0EBE1; border-radius: 9px;
        padding: 1rem 1.2rem; margin-bottom: 0.7rem;
        border-left: 3px solid #8B7355;
    }
    .doc-skill-title {
        font-size: 1rem; font-weight: 500; color: #2A2520;
        margin-bottom: 0.3rem;
    }
    .doc-skill-body {
        font-size: 0.87rem; line-height: 1.7;
        color: #4A4540; font-weight: 300;
    }

    .doc-bullet-item {
        display: flex; align-items: flex-start; gap: 0.55rem;
        padding: 0.45rem 0; border-bottom: 1px solid #F0EBE1;
        font-size: 0.9rem; line-height: 1.65; color: #3A3530; font-weight: 300;
    }
    .doc-bullet-item:last-child { border-bottom: none; }
    .doc-dot {
        width: 5px; height: 5px; background: #8B7355;
        border-radius: 50%; margin-top: 0.6rem; flex-shrink: 0;
    }

    .doc-tag {
        display: inline-block; padding: 0.2rem 0.6rem;
        border-radius: 99px; background: #F0EBE1;
        border: 1px solid #DDD8CE;
        font-family: 'Jost', sans-serif; font-size: 0.75rem;
        color: #6B6456; margin: 0.15rem;
    }
    .doc-level {
        display: inline-block; padding: 0.15rem 0.55rem;
        border-radius: 99px; font-family: 'Jost', sans-serif;
        font-size: 0.72rem; font-weight: 500; text-transform: capitalize;
        margin-left: 0.3rem;
    }
    .doc-level-low    { background: #EAF2E8; color: #4A6A42; }
    .doc-level-moderate { background: #FEF6E4; color: #8A6420; }
    .doc-level-high   { background: #FAECE8; color: #8A3A2A; }

    .doc-monitor-tbl { width: 100%; border-collapse: collapse; }
    .doc-monitor-tbl th {
        background: #2A2520; color: #fff; padding: 0.5rem 0.8rem;
        font-family: 'Jost', sans-serif; font-size: 0.67rem;
        font-weight: 500; letter-spacing: 0.07em; text-transform: uppercase;
        text-align: left;
    }
    .doc-monitor-tbl td {
        padding: 0.6rem 0.8rem; font-size: 0.9rem;
        border-bottom: 1px solid #DDD8CE; color: #2A2520;
    }
    .doc-monitor-tbl tr:nth-child(even) td { background: #F5F0E8; }

    .doc-progress-grid {
        display: grid; grid-template-columns: 1fr 1fr; gap: 0.6rem;
    }
    .doc-progress-card {
        background: #F0EBE1; border-radius: 9px;
        padding: 0.85rem 1rem; border: 1px solid #DDD8CE;
    }
    .doc-progress-icon { font-size: 1.1rem; margin-bottom: 0.25rem; }
    .doc-progress-title {
        font-family: 'Jost', sans-serif;
        font-size: 0.82rem; font-weight: 500; color: #2A2520; margin-bottom: 0.15rem;
    }
    .doc-progress-body {
        font-size: 0.8rem; line-height: 1.55; color: #6B6456; font-weight: 300;
    }

    .doc-collab-grid {
        display: grid; grid-template-columns: repeat(3, 1fr); gap: 0.6rem;
    }
    .doc-collab-card {
        background: #F0EBE1; border-radius: 9px;
        padding: 0.9rem; text-align: center; border: 1px solid #DDD8CE;
    }
    .doc-collab-icon { font-size: 1.4rem; margin-bottom: 0.3rem; }
    .doc-collab-title {
        font-family: 'Jost', sans-serif;
        font-size: 0.85rem; font-weight: 500; color: #2A2520; margin-bottom: 0.25rem;
    }
    .doc-collab-body {
        font-size: 0.78rem; line-height: 1.55; color: #6B6456; font-weight: 300;
    }

    .doc-footer {
        display: flex; justify-content: space-between; align-items: center;
        padding: 1rem 3rem; background: #F0EBE1;
        border-top: 1px solid #DDD8CE;
        font-family: 'Jost', sans-serif; font-size: 0.7rem; color: #6B6456;
    }

    /* Export section */
    .export-box {
        background: var(--cream);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 1.25rem 1.5rem;
        margin-top: 1.5rem;
    }
    .export-title {
        font-family: 'Jost', sans-serif;
        font-size: 0.65rem; font-weight: 600;
        letter-spacing: 0.14em; text-transform: uppercase;
        color: var(--bark); margin-bottom: 0.75rem;
    }
    </style>
    """, unsafe_allow_html=True)


# -----------------------------------------------------------------------------
# SKILL DEFINITIONS
# -----------------------------------------------------------------------------
SKILL_DETAILS = {
    "Emotional Regulation": (
        "Emotional Regulation",
        "The child will develop the ability to identify, understand, and manage emotional responses in a "
        "developmentally appropriate manner. Activities will focus on building an internal emotional vocabulary, "
        "recognizing physiological cues of dysregulation, and applying self-soothing strategies with guided adult "
        "support before transitioning toward greater independence."
    ),
    "Impulse Control": (
        "Impulse Control (Inhibitory Control)",
        "Inhibitory control — the capacity to pause, reflect, and redirect an impulse before acting — is a "
        "foundational executive function skill. Sessions will target this through structured stop-and-think games, "
        "delayed gratification activities, and consistent adult modeling of reflective pausing. Progress will be "
        "measured by an observable increase in the child's latency between trigger and behavioral response."
    ),
    "Frustration Tolerance": (
        "Frustration Tolerance",
        "The child will be gradually and safely exposed to manageable frustration experiences within a supported "
        "environment. Sessions will teach the child to recognize frustration as a temporary, tolerable state, and "
        "to deploy a toolkit of coping strategies — including verbal expression, physical regulation activities, "
        "and requesting adult support — rather than escalating to disruptive behavior."
    ),
    "Communication Skills": (
        "Functional Communication Skills",
        "A primary driver of challenging behavior in early childhood is limited access to functional communication. "
        "This plan will prioritize the development of verbal and non-verbal communication pathways that allow the "
        "child to express needs, preferences, and distress in socially acceptable ways. Replacement behavior "
        "training will be central to this goal."
    ),
    "Social Skills": (
        "Social Skills and Peer Interaction",
        "The child will develop age-appropriate skills in reading social cues, turn-taking, cooperative play, and "
        "conflict resolution. Sessions will use structured peer interaction activities, social stories, and "
        "role-play scenarios to build a repertoire of positive social behaviors that can be generalized to the "
        "classroom environment."
    ),
    "Self-Awareness": (
        "Self-Awareness and Emotional Literacy",
        "The child will build an increasing awareness of their own emotional states, bodily sensations associated "
        "with emotion, and the impact of their behavior on others. This foundational skill supports all other "
        "targeted areas and will be developed through reflective activities, visual tools, and consistent, "
        "non-judgmental adult dialogue."
    ),
    "Coping Strategies": (
        "Coping Strategy Development",
        "The child will be introduced to and will practice a personalized toolkit of coping strategies appropriate "
        "to their developmental level. Strategies may include deep breathing techniques, movement breaks, sensory "
        "regulation tools, and structured calming corners. The goal is for the child to begin accessing these "
        "tools with decreasing levels of adult support over time."
    ),
    "Attention and Concentration": (
        "Attention and Concentration",
        "The child will develop the capacity to sustain focused attention on a task or activity for increasing "
        "durations, appropriate to their developmental stage. Sessions will use structured, play-based attention "
        "exercises with graduated difficulty, visual focus cues, and reinforcement systems. The goal is for the "
        "child to build on-task behaviour and resist distracting stimuli with progressively less adult scaffolding."
    ),
    "Decision Making": (
        "Decision Making and Problem Solving",
        "The child will develop age-appropriate skills in weighing simple choices, anticipating consequences, "
        "and selecting adaptive responses to everyday situations. Sessions will use guided choice-making tasks, "
        "visual choice boards, and structured problem-solving conversations that build the child's confidence "
        "in making independent decisions within safe, supported boundaries."
    ),
    "Self-Boundaries": (
        "Self-Boundaries and Personal Space Awareness",
        "The child will learn to recognise and communicate their own physical and emotional boundaries, and "
        "to respect the boundaries of others. Sessions will address body autonomy, identifying comfortable and "
        "uncomfortable interactions, and using assertive yet appropriate language to express limits. "
        "This skill is foundational to both peer relationships and personal safety."
    ),
    "Hyperactivity Reduction": (
        "Hyperactivity Reduction and Self-Regulation of Activity Level",
        "The child will develop awareness of their own arousal level and build a repertoire of strategies to "
        "modulate physical activity when the environment requires it. Sessions will include movement-regulation "
        "games, sensory input activities, structured transitions, and explicit teaching of body-awareness cues. "
        "Progress is tracked through observable increases in the child's capacity to sustain calm engagement."
    ),
    "Basic Knowledge": (
        "Basic Knowledge and Foundational Concepts",
        "The child will develop and consolidate foundational cognitive concepts including colours, shapes, "
        "numbers, letters, spatial relationships, and sequencing. These building blocks are addressed within "
        "a structured play-based framework that prioritises engagement, repetition, and positive reinforcement, "
        "ensuring skills are solidly established before new complexity is introduced."
    ),
    "Emotional Resilience": (
        "Emotional Resilience and Recovery",
        "The child will develop the capacity to bounce back from setbacks, disappointments, and challenging "
        "emotional experiences without prolonged dysregulation. Sessions will build a growth mindset around "
        "mistakes, practice perspective-taking after difficult moments, and reinforce the child's narrative "
        "of themselves as capable and resilient."
    ),
    "Independence and Self-Care": (
        "Independence and Self-Care Skills",
        "The child will develop age-appropriate independence in daily living tasks such as dressing, tidying, "
        "and personal hygiene routines. Sessions will use task analysis, visual schedules, and reinforcement "
        "systems to break each skill into manageable steps, building confidence and reducing reliance on "
        "adult prompting over time."
    ),
    "Transitions and Flexibility": (
        "Transitions and Cognitive Flexibility",
        "The child will develop improved tolerance for changes in routine and smoother transitions between "
        "activities. Sessions will use visual transition tools, countdown strategies, and preparatory cues "
        "to help the child anticipate and accept change. The goal is to increase the child's overall "
        "cognitive flexibility and reduce distress when expectations shift."
    ),
    "Play Skills": (
        "Play Skills and Imaginative Engagement",
        "The child will develop structured and imaginative play skills, including the ability to engage in "
        "purposeful solo play and cooperative play with peers. Sessions will scaffold progressively complex "
        "play sequences and introduce shared play narratives, building both enjoyment and the social "
        "competencies that emerge naturally through child-led play."
    ),
}

BEHAVIOR_OPTIONS = [
    "Throwing objects", "Hitting / kicking others", "Dropping to the floor",
    "Tantrums / meltdowns", "Biting", "Screaming / shouting",
    "Refusing transitions", "Destroying materials", "Running away / elopement",
    "Self-injurious behavior", "Spitting", "Verbal aggression",
    "Potty training difficulties", "Rule following challenges",
    "Interrupting / difficulty waiting for turn", "Food refusal / mealtime difficulties",
    "Separation anxiety behaviors", "Excessive clinginess",
    "Difficulty sharing / possessive behavior", "Wandering / poor environmental awareness",
]

MONITORING_ROWS = [
    ("Behavior type and topography",
     "Description of the specific behavior observed, its physical characteristics, and the context in which it occurred."),
    ("Trigger identification",
     "The antecedent events or conditions (internal and environmental) that preceded the behavior."),
    ("Response intensity",
     "Severity of the behavioral episode rated on a consistent 1–5 scale to track changes over time."),
    ("Recovery time",
     "Time elapsed from onset of behavioral escalation to return to a regulated, engaged baseline state."),
    ("Use of replacement behaviors",
     "Whether the child initiated or responded to prompts to use a replacement behavior, and the degree of adult support required."),
    ("Adult response fidelity",
     "Whether the planned intervention protocol was implemented as designed during the session."),
]

PROGRESS_CARDS = [
    ("⏱", "Increased Latency to Reaction",
     "An observable pause between the trigger event and the behavioral response, indicating developing inhibitory control."),
    ("📉", "Reduced Behavioral Intensity",
     "A measurable decrease in the severity of behavioral episodes over a consistent time period."),
    ("🔄", "Faster Recovery",
     "A progressive reduction in the time required to return to a calm, engaged state following dysregulation."),
    ("💬", "Improved Communication",
     "Increased frequency of using verbal or non-verbal communication to express needs, particularly in high-demand situations."),
    ("✋", "Spontaneous Replacement Behaviors",
     "The child independently initiating a replacement behavior without adult prompting in the natural environment."),
    ("🌍", "Generalisation Across Settings",
     "Transfer of skill use beyond the formal session context into the broader nursery environment and, where reported, the home setting."),
]

ACTIVITY_ITEMS = [
    ("Co-regulation activities",
     "Guided breathing techniques, sensory regulation tools, and calming rituals to help the child experience and "
     "internalise regulated states before working toward independent regulation."),
    ("Inhibitory control games",
     "Stop-and-go games, impulse delay exercises, and structured waiting activities that build the neurological "
     "pathways underlying self-control."),
    ("Structured play challenges",
     "Activities deliberately designed to include manageable difficulty, turn-taking requirements, and peer "
     "interaction to build frustration tolerance and social competence."),
    ("Replacement behavior rehearsal",
     "Explicit teaching and repeated practice of socially acceptable replacement behaviors for each identified "
     "concern behavior (e.g., using words instead of hitting; requesting a break instead of throwing)."),
    ("Positive reinforcement systems",
     "Consistent, specific, and immediate reinforcement of target behaviors and approximations toward target "
     "behaviors, using individually motivating reinforcers identified in collaboration with the child and family."),
    ("Reflective dialogue",
     "Age-appropriate post-incident discussion to build the child's capacity for emotional retrospection "
     "and self-awareness without shame or blame."),
]


# -----------------------------------------------------------------------------
# DOCUMENT PREVIEW (HTML)
# -----------------------------------------------------------------------------
def render_html_preview(d):
    """Render the full document as HTML for the Streamlit preview."""
    today = datetime.now().strftime("%d %B %Y")
    start = (d["start_date"].strftime("%d %B %Y")
             if d.get("start_date") else "—")
    all_behaviors = d["behaviors"] + ([d["custom_behavior"]] if d.get("custom_behavior") else [])

    # Logo
    if d.get("logo_bytes"):
        b64 = base64.b64encode(d["logo_bytes"]).decode()
        logo_html = f'<img src="data:image/png;base64,{b64}" style="max-height:48px;max-width:130px;object-fit:contain;filter:brightness(0) invert(1);opacity:0.85;" />'
    else:
        logo_html = f'<span class="doc-nursery-nm">{d.get("nursery_name","Nursery Name")}</span>'

    # Behaviors
    tag_html = "".join(f'<span class="doc-tag">{b}</span>' for b in all_behaviors) if all_behaviors \
               else '<span class="doc-tag" style="opacity:0.4">None specified</span>'

    level_cls = {"low":"doc-level-low","moderate":"doc-level-moderate","high":"doc-level-high"}
    freq_badge = f'<span class="doc-level {level_cls.get(d["frequency"],"")}">{d["frequency"]}</span>'
    int_badge  = f'<span class="doc-level {level_cls.get(d["intensity"],"")}">{d["intensity"]}</span>'

    # Skills
    skills_html = ""
    all_skills = d.get("skills", []) + ([d["custom_skill"]] if d.get("custom_skill") else [])
    for s in all_skills:
        if s in SKILL_DETAILS:
            title, body = SKILL_DETAILS[s]
            skills_html += f"""
            <div class="doc-skill-blk">
                <div class="doc-skill-title">{title}</div>
                <div class="doc-skill-body">{body}</div>
            </div>"""
        else:
            # Custom / Other skill
            skills_html += f"""
            <div class="doc-skill-blk">
                <div class="doc-skill-title">{s}</div>
                <div class="doc-skill-body">This is a custom skill area identified for this child's programme. 
                Specific goals and strategies will be developed collaboratively with the team based on the 
                child's individual profile and needs.</div>
            </div>"""
    if not skills_html:
        skills_html = '<p class="doc-p" style="opacity:0.4">No skills selected.</p>'

    # Activities — dynamically tailored to skills + behaviors
    all_skills_for_ai  = tuple(all_skills)
    all_behaviors_for_ai = tuple(all_behaviors)
    child_name_for_ai  = d.get("child_name", "")

    dynamic_activities = generate_dynamic_activities(all_skills_for_ai, all_behaviors_for_ai, child_name_for_ai)
    activities_html = "".join(
        f'<div class="doc-bullet-item"><div class="doc-dot"></div>'
        f'<div><strong>{t}:</strong> {b}</div></div>'
        for t, b in dynamic_activities
    )

    # Monitoring table
    monitor_rows = "".join(
        f'<tr><td><strong>{d_}</strong></td><td>{r}</td></tr>'
        for d_, r in MONITORING_ROWS
    )

    # Progress cards — dynamically tailored to skills + behaviors
    dynamic_progress = generate_dynamic_progress(all_skills_for_ai, all_behaviors_for_ai, child_name_for_ai)
    progress_cards = "".join(
        f'<div class="doc-progress-card"><div class="doc-progress-icon">{ic}</div>'
        f'<div class="doc-progress-title">{t}</div>'
        f'<div class="doc-progress-body">{b}</div></div>'
        for ic, t, b in dynamic_progress
    )

    # Notes
    notes_html = ""
    if d.get("notes") and d["notes"].strip():
        notes_html = f"""
        <div style="margin-top:1.2rem;padding:1rem 1.2rem;background:#F0EBE1;border-radius:9px;border:1px solid #DDD8CE;">
            <strong style="font-family:'Jost',sans-serif;font-size:0.82rem;color:#8B7355;">Additional Notes</strong>
            <p class="doc-p" style="margin-top:0.4rem;">{d["notes"]}</p>
        </div>"""

    footer_text = d.get("footer_text") or "Confidential — For internal and family use only."
    child_name  = d.get("child_name") or "—"
    nursery     = d.get("nursery_name") or "Nursery"
    pattern     = d.get("pattern") or "No specific pattern noted."

    html = f"""
    <div class="doc-preview-wrap">
        <!-- HEADER -->
        <div class="doc-hdr">
            <div class="doc-hdr-top">
                {logo_html}
                <span class="doc-badge">Confidential</span>
            </div>
            <div class="doc-main-title">Individual Behavior<br>Support Plan</div>
            <div class="doc-main-sub">Developmental Support Programme — {nursery}</div>
        </div>

        <!-- INFO BAR -->
        <div class="doc-info-bar">
            <div class="doc-info-cell">
                <div class="doc-info-lbl">Child Name</div>
                <div class="doc-info-val">{child_name}</div>
            </div>
            <div class="doc-info-cell">
                <div class="doc-info-lbl">Date of Birth</div>
                <div class="doc-info-val">{d["date_of_birth"].strftime("%d %B %Y") if d.get("date_of_birth") else "—"}</div>
            </div>
            <div class="doc-info-cell">
                <div class="doc-info-lbl">Age</div>
                <div class="doc-info-val">{d.get("age","—")}</div>
            </div>
            <div class="doc-info-cell">
                <div class="doc-info-lbl">Class / Group</div>
                <div class="doc-info-val">{d.get("class_group","—")}</div>
            </div>
        </div>

        <!-- BODY -->
        <div class="doc-body">

            <!-- 1. PURPOSE -->
            <div class="doc-sec">
                <div class="doc-sec-hdr">
                    <div class="doc-sec-num">1</div>
                    <div class="doc-sec-title">Programme Purpose &amp; Scope</div>
                </div>
                <p class="doc-p">This Individual Behavior Support Plan has been developed to provide structured,
                evidence-informed support for <strong>{child_name}'s</strong> developmental needs. It is important
                to note that this plan is <em>not</em> a diagnostic document, nor does it imply the presence of any
                clinical condition or disorder.</p>
                <p class="doc-p">The purpose of this programme is to build specific developmental skills — including
                emotional regulation, impulse control, frustration tolerance, and functional communication — within
                the safe, consistent, and nurturing environment of the nursery setting.</p>
                <p class="doc-p">All interventions are grounded in a positive, strengths-based framework that
                emphasises the child's capacity for growth. Progress is expected to occur gradually and
                non-linearly; the programme is designed to be responsive and adaptive based on ongoing observation
                and collaboration with the family.</p>
                <p class="doc-p" style="font-family:'Jost',sans-serif;font-size:0.85rem;color:#6B6456;">
                    <strong>Sessions per month:</strong> {d.get("sessions_per_month","—")} &nbsp;&nbsp;
                    <strong>Frequency:</strong> {d.get("session_frequency","—")}
                </p>
            </div>

            <!-- 2. BEHAVIORAL CONCERNS -->
            <div class="doc-sec">
                <div class="doc-sec-hdr">
                    <div class="doc-sec-num">2</div>
                    <div class="doc-sec-title">Identified Behavioral Concerns</div>
                </div>
                <p class="doc-p">The following behaviors have been identified as current areas of focus. These
                behaviors are understood as the child's current best attempt to communicate unmet needs or to manage
                experiences that exceed their present regulatory capacity:</p>
                <div style="margin:0.6rem 0 0.8rem;">{tag_html}</div>
                <p class="doc-p">
                    <strong>Frequency:</strong>{freq_badge} &nbsp;&nbsp;
                    <strong>Intensity:</strong>{int_badge}
                </p>
                <p class="doc-p"><strong>Recent Pattern / Context:</strong> {pattern}</p>
            </div>

            <!-- 3. SKILLS -->
            <div class="doc-sec">
                <div class="doc-sec-hdr">
                    <div class="doc-sec-num">3</div>
                    <div class="doc-sec-title">Core Skills Targeted</div>
                </div>
                <p class="doc-p">The following skill areas have been selected as the primary targets for this
                programme. Each skill area is interconnected; progress in one domain typically supports development
                across others.</p>
                {skills_html}
            </div>

            <!-- 4. ACTIVITIES -->
            <div class="doc-sec">
                <div class="doc-sec-hdr">
                    <div class="doc-sec-num">4</div>
                    <div class="doc-sec-title">Session Activities</div>
                </div>
                <p class="doc-p">Sessions will incorporate a variety of structured, play-based, and reflective
                activities designed to build the targeted skills in an engaging and developmentally appropriate
                manner:</p>
                {activities_html}
            </div>

            <!-- 5. MONITORING -->
            <div class="doc-sec">
                <div class="doc-sec-hdr">
                    <div class="doc-sec-num">5</div>
                    <div class="doc-sec-title">Session Monitoring Framework</div>
                </div>
                <p class="doc-p">Each session will be monitored using the following framework to track behavioral
                patterns, responses to intervention, and progress over time:</p>
                <table class="doc-monitor-tbl">
                    <thead>
                        <tr><th>Monitoring Dimension</th><th>What Is Recorded</th></tr>
                    </thead>
                    <tbody>{monitor_rows}</tbody>
                </table>
            </div>

            <!-- 6. PROGRESS -->
            <div class="doc-sec">
                <div class="doc-sec-hdr">
                    <div class="doc-sec-num">6</div>
                    <div class="doc-sec-title">Indicators of Progress</div>
                </div>
                <p class="doc-p">Progress will be evaluated through direct observation, session notes, and regular
                review meetings. The following indicators will be used as primary markers of meaningful progress:</p>
                <div class="doc-progress-grid">{progress_cards}</div>
            </div>

            <!-- 7. COLLABORATION -->
            <div class="doc-sec">
                <div class="doc-sec-hdr">
                    <div class="doc-sec-num">7</div>
                    <div class="doc-sec-title">Family &amp; Team Collaboration</div>
                </div>
                <p class="doc-p">This programme operates on the understanding that meaningful and lasting
                developmental change occurs when the child's key environments are aligned in their approach and
                consistent in their response.</p>
                <div class="doc-collab-grid">
                    <div class="doc-collab-card">
                        <div class="doc-collab-icon">🏫</div>
                        <div class="doc-collab-title">Nursery Team</div>
                        <div class="doc-collab-body">Teaching staff are briefed on programme goals and monitoring
                        frameworks. Fortnightly check-ins ensure responsiveness and fidelity to the plan.</div>
                    </div>
                    <div class="doc-collab-card">
                        <div class="doc-collab-icon">👩‍🏫</div>
                        <div class="doc-collab-title">Behavioral Therapist</div>
                        <div class="doc-collab-body">The assigned behavioral therapist maintains session monitoring records
                        and serves as primary liaison between the programme and the classroom environment.{f" Therapist: <strong>{d.get('therapist_name')}</strong>." if d.get('therapist_name') else ""}</div>
                    </div>
                    <div class="doc-collab-card">
                        <div class="doc-collab-icon">👨‍👩‍👧</div>
                        <div class="doc-collab-title">Family Partnership</div>
                        <div class="doc-collab-body">Parents receive regular written updates and recommended home
                        reinforcement strategies. Family input is actively valued and integral to all reviews.</div>
                    </div>
                </div>
                {notes_html}
            </div>

        </div><!-- /doc-body -->

        <!-- FOOTER -->
        <div class="doc-footer">
            <span>{footer_text}</span>
            <span style="color:#8B7355;">Generated: {today} | {nursery}</span>
        </div>
    </div>
    """
    return html


# -----------------------------------------------------------------------------
# PDF GENERATION  (reportlab)
# -----------------------------------------------------------------------------
def build_pdf(d) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, HRFlowable, KeepTogether, Image as RLImage
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from io import BytesIO

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=2.2*cm, rightMargin=2.2*cm,
                             topMargin=2.5*cm, bottomMargin=2.5*cm)

    # Palette
    INK        = colors.HexColor("#2A2520")
    BARK       = colors.HexColor("#8B7355")
    BARK_LT    = colors.HexColor("#B09A7A")
    STONE      = colors.HexColor("#6B6456")
    LINEN      = colors.HexColor("#F0EBE1")
    WARM       = colors.HexColor("#FAF8F4")
    WARM2      = colors.HexColor("#F5F0E8")
    BORDER     = colors.HexColor("#DDD8CE")
    WHITE      = colors.white
    LEVEL_COLS = {
        "low":      (colors.HexColor("#EAF2E8"), colors.HexColor("#4A6A42")),
        "moderate": (colors.HexColor("#FEF6E4"), colors.HexColor("#8A6420")),
        "high":     (colors.HexColor("#FAECE8"), colors.HexColor("#8A3A2A")),
    }

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    sTitle   = S("T",  fontName="Helvetica-Bold",  fontSize=20, textColor=INK,
                       leading=26, alignment=TA_CENTER, spaceAfter=4)
    sSub2    = S("Su", fontName="Helvetica",        fontSize=10, textColor=STONE,
                       leading=14, alignment=TA_CENTER, spaceAfter=2)
    sSection = S("Se", fontName="Helvetica-Bold",   fontSize=10, textColor=BARK,
                       leading=14, spaceBefore=14, spaceAfter=4)
    sBody    = S("Bo", fontName="Helvetica",         fontSize=9.5, textColor=INK,
                       leading=15, spaceAfter=4)
    sBullet  = S("Bu", fontName="Helvetica",         fontSize=9.5, textColor=INK,
                       leading=15, leftIndent=14, spaceAfter=3)
    sSmall   = S("Sm", fontName="Helvetica",         fontSize=8.5, textColor=STONE, leading=13)
    sFooter  = S("Fo", fontName="Helvetica-Oblique", fontSize=7.5, textColor=STONE,
                       alignment=TA_CENTER)
    sSkillT  = S("ST", fontName="Helvetica-Bold",    fontSize=10, textColor=INK, leading=14, spaceAfter=2)
    sSkillB  = S("SB", fontName="Helvetica",         fontSize=9,  textColor=colors.HexColor("#4A4540"),
                       leading=14, spaceAfter=2)

    def hr(c=BARK, t=0.8):
        return HRFlowable(width="100%", thickness=t, color=c, spaceAfter=5, spaceBefore=5)
    def sp(h=0.25):
        return Spacer(1, h*cm)

    today = datetime.now().strftime("%d %B %Y")
    start = d["start_date"].strftime("%d %B %Y") if d.get("start_date") else "—"
    all_behaviors = d["behaviors"] + ([d["custom_behavior"]] if d.get("custom_behavior") else [])
    all_skills    = d.get("skills", []) + ([d["custom_skill"]] if d.get("custom_skill") else [])
    nursery   = d.get("nursery_name") or "Nursery"
    child     = d.get("child_name")   or "—"
    footer_t  = d.get("footer_text")  or "Confidential — For internal and family use only."

    story = []

    # -- Title block ------------------------------------------------------------
    story += [
        Paragraph("Individual Behavior Support Plan", sTitle),
        Paragraph(f"Developmental Support Programme — {nursery}", sSub2),
        sp(0.3), hr(INK, 1.5), sp(0.2),
    ]

    # -- Logo (if uploaded) ----------------------------------------------------
    if d.get("logo_bytes"):
        try:
            img_buf = BytesIO(d["logo_bytes"])
            img = RLImage(img_buf, width=3*cm, height=1.8*cm, kind="proportional")
            img.hAlign = "CENTER"
            story += [img, sp(0.2)]
        except Exception:
            pass

    # -- Demographic table ------------------------------------------------------
    dob_str = d["date_of_birth"].strftime("%d %B %Y") if d.get("date_of_birth") else "—"
    therapist = d.get("therapist_name") or "—"

    demo_rows = [
        [Paragraph("<b>Child Name</b>",     sSmall), Paragraph(child,                       sBody),
         Paragraph("<b>Date of Birth</b>",  sSmall), Paragraph(dob_str,                     sBody)],
        [Paragraph("<b>Age</b>",            sSmall), Paragraph(d.get("age","—"),            sBody),
         Paragraph("<b>Class / Group</b>",  sSmall), Paragraph(d.get("class_group","—"),    sBody)],
        [Paragraph("<b>Therapist</b>",      sSmall), Paragraph(therapist,                   sBody),
         Paragraph("<b>Start Date</b>",     sSmall), Paragraph(start,                       sBody)],
        [Paragraph("<b>Sessions/month</b>", sSmall), Paragraph(str(d.get("sessions_per_month","—")), sBody),
         Paragraph("<b>Frequency</b>",      sSmall), Paragraph(d.get("session_frequency","—"), sBody)],
    ]
    demo_tbl = Table(demo_rows, colWidths=[3.0*cm, 6.0*cm, 3.0*cm, 4.2*cm])
    demo_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0),(0,-1), LINEN),
        ("BACKGROUND",   (2,0),(2,-1), LINEN),
        ("ROWBACKGROUNDS",(0,0),(-1,-1),[WARM, WARM2, WARM, WARM2]),
        ("GRID",         (0,0),(-1,-1), 0.4, BORDER),
        ("VALIGN",       (0,0),(-1,-1), "MIDDLE"),
        ("TOPPADDING",   (0,0),(-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING",  (0,0),(-1,-1), 7),
    ]))
    story += [demo_tbl, sp(0.4)]

    # -- Confidential badge -----------------------------------------------------
    badge_row = [[Paragraph("<b>  CONFIDENTIAL DOCUMENT  </b>",
        S("CB", fontName="Helvetica-Bold", fontSize=9, textColor=BARK, alignment=TA_CENTER))]]
    badge_tbl = Table(badge_row, colWidths=[16.2*cm])
    badge_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0),(-1,-1), LINEN),
        ("TOPPADDING",   (0,0),(-1,-1), 7),
        ("BOTTOMPADDING",(0,0),(-1,-1), 7),
        ("BOX",          (0,0),(-1,-1), 0.5, BARK),
    ]))
    story += [badge_tbl, sp(0.5)]

    def sec_hdr(num, title):
        row = [[Paragraph(str(num),
                S("SN", fontName="Helvetica-Bold", fontSize=9, textColor=WHITE, alignment=TA_CENTER)),
                Paragraph(f"<b>{title}</b>",
                S("ST2", fontName="Helvetica-Bold", fontSize=11, textColor=INK))]]
        t = Table(row, colWidths=[0.7*cm, 15.5*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(0,0), INK),
            ("BACKGROUND",    (1,0),(1,0), LINEN),
            ("VALIGN",        (0,0),(-1,-1),"MIDDLE"),
            ("TOPPADDING",    (0,0),(-1,-1), 6),
            ("BOTTOMPADDING", (0,0),(-1,-1), 6),
            ("LEFTPADDING",   (0,0),(0,0), 0),
            ("LEFTPADDING",   (1,0),(1,0), 10),
            ("BOTTOMPADDING", (0,0),(-1,-1), 6),
            ("LINEBELOW",     (0,0),(-1,-1), 0.5, BARK),
        ]))
        return t

    def bullet_p(title, body):
        return Paragraph(f"<b>{title}:</b> {body}", sBullet)

    # -- Section 1 — Purpose ----------------------------------------------------
    story += [sec_hdr(1, "Programme Purpose & Scope"), sp(0.15)]
    story += [
        Paragraph(f"This Individual Behavior Support Plan has been developed to provide structured, "
                  f"evidence-informed support for <b>{child}'s</b> developmental needs. It is important to note "
                  f"that this plan is not a diagnostic document, nor does it imply the presence of any clinical "
                  f"condition or disorder.", sBody),
        Paragraph("The purpose of this programme is to build specific developmental skills — including "
                  "emotional regulation, impulse control, frustration tolerance, and functional communication — "
                  "within the safe, consistent, and nurturing environment of the nursery setting.", sBody),
        Paragraph("All interventions are grounded in a positive, strengths-based framework that emphasises "
                  "the child's capacity for growth. Progress is expected to occur gradually and non-linearly; "
                  "the programme is designed to be responsive and adaptive based on ongoing observation and "
                  "collaboration with the family.", sBody),
        sp(0.4),
    ]

    # -- Section 2 — Behavioral Concerns ----------------------------------------
    story += [sec_hdr(2, "Identified Behavioral Concerns"), sp(0.15)]
    story.append(Paragraph("The following behaviors have been identified as current areas of focus. "
                            "These behaviors are understood as the child's current best attempt to communicate "
                            "unmet needs or to manage experiences that exceed their present regulatory capacity:", sBody))

    if all_behaviors:
        beh_text = " | ".join(all_behaviors)
        story.append(Paragraph(beh_text, S("BT", fontName="Helvetica-Oblique", fontSize=9.5,
                                            textColor=STONE, leading=14, spaceBefore=4, spaceAfter=4,
                                            leftIndent=10, borderPad=4)))

    freq_bg, freq_fg = LEVEL_COLS.get(d.get("frequency","moderate"), LEVEL_COLS["moderate"])
    int_bg,  int_fg  = LEVEL_COLS.get(d.get("intensity","moderate"),  LEVEL_COLS["moderate"])

    level_row = [
        [Paragraph("<b>Frequency</b>", sSmall),
         Paragraph(d.get("frequency","—").capitalize(),
                   S("LV", fontName="Helvetica-Bold", fontSize=9, textColor=freq_fg)),
         Paragraph("<b>Intensity</b>", sSmall),
         Paragraph(d.get("intensity","—").capitalize(),
                   S("LV2", fontName="Helvetica-Bold", fontSize=9, textColor=int_fg))],
    ]
    lv_tbl = Table(level_row, colWidths=[2.5*cm, 3.5*cm, 2.5*cm, 3.5*cm])
    lv_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (1,0),(1,0), freq_bg),
        ("BACKGROUND",   (3,0),(3,0), int_bg),
        ("TOPPADDING",   (0,0),(-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING",  (0,0),(-1,-1), 6),
        ("VALIGN",       (0,0),(-1,-1), "MIDDLE"),
        ("BOX",          (1,0),(1,0), 0.4, BORDER),
        ("BOX",          (3,0),(3,0), 0.4, BORDER),
    ]))
    story += [lv_tbl, sp(0.1)]

    if d.get("pattern"):
        story.append(Paragraph(f"<b>Recent Pattern / Context:</b> {d['pattern']}", sBody))
    story.append(sp(0.4))

    # -- Section 3 — Skills -----------------------------------------------------
    story += [sec_hdr(3, "Core Skills Targeted"), sp(0.15)]
    story.append(Paragraph("The following skill areas have been selected as the primary targets for this "
                            "programme. Each skill area is interconnected; progress in one domain typically "
                            "supports development across others.", sBody))

    for s in all_skills:
        if s in SKILL_DETAILS:
            title, body = SKILL_DETAILS[s]
        else:
            title = s
            body = ("This is a custom skill area identified for this child's programme. "
                    "Specific goals and strategies will be developed collaboratively with the team based on the "
                    "child's individual profile and needs.")
        skill_row = [[
            Paragraph(f"<b>{title}</b>", sSkillT),
            Paragraph(body, sSkillB),
        ]]
        st2 = Table(skill_row, colWidths=[4.0*cm, 12.2*cm])
        st2.setStyle(TableStyle([
            ("BACKGROUND",   (0,0),(-1,-1), LINEN),
            ("LINEAFTER",    (0,0),(0,0), 3, BARK),
            ("TOPPADDING",   (0,0),(-1,-1), 8),
            ("BOTTOMPADDING",(0,0),(-1,-1), 8),
            ("LEFTPADDING",  (0,0),(-1,-1), 10),
            ("VALIGN",       (0,0),(-1,-1), "TOP"),
        ]))
        story += [st2, sp(0.15)]
    story.append(sp(0.3))

    # -- Section 4 — Activities -------------------------------------------------
    story += [sec_hdr(4, "Session Activities"), sp(0.15)]
    story.append(Paragraph("Sessions will incorporate a variety of structured, play-based, and reflective "
                            "activities designed to build the targeted skills and address the identified "
                            "behavioral concerns in an engaging and developmentally appropriate manner:", sBody))
    _acts = generate_dynamic_activities(
        tuple(all_skills), tuple(all_behaviors), child
    )
    for t_, b_ in _acts:
        story.append(bullet_p(t_, b_))
    story.append(sp(0.4))

    # -- Section 5 — Monitoring -------------------------------------------------
    story += [sec_hdr(5, "Session Monitoring Framework"), sp(0.15)]
    story.append(Paragraph("Each session will be monitored using the following framework to track behavioral "
                            "patterns, responses to intervention, and progress over time:", sBody))

    mon_data = [[Paragraph("<b>Monitoring Dimension</b>",
                            S("MH", fontName="Helvetica-Bold", fontSize=8.5, textColor=WHITE)),
                  Paragraph("<b>What Is Recorded</b>",
                            S("MH2", fontName="Helvetica-Bold", fontSize=8.5, textColor=WHITE))]]
    for d_, r in MONITORING_ROWS:
        mon_data.append([Paragraph(f"<b>{d_}</b>", sBody), Paragraph(r, sBody)])

    mon_tbl = Table(mon_data, colWidths=[5.5*cm, 10.7*cm])
    mon_tbl.setStyle(TableStyle([
        ("BACKGROUND",       (0,0),(-1,0), INK),
        ("ROWBACKGROUNDS",   (0,1),(-1,-1), [WARM, LINEN]),
        ("GRID",             (0,0),(-1,-1), 0.4, BORDER),
        ("TOPPADDING",       (0,0),(-1,-1), 6),
        ("BOTTOMPADDING",    (0,0),(-1,-1), 6),
        ("LEFTPADDING",      (0,0),(-1,-1), 8),
        ("VALIGN",           (0,0),(-1,-1), "MIDDLE"),
    ]))
    story += [mon_tbl, sp(0.4)]

    # -- Section 6 — Progress ---------------------------------------------------
    story += [sec_hdr(6, "Indicators of Progress"), sp(0.15)]
    story.append(Paragraph("Progress will be evaluated through direct observation, session notes, and regular "
                            "review meetings. The following indicators will be used as primary markers of "
                            "meaningful progress:", sBody))

    _prog = generate_dynamic_progress(
        tuple(all_skills), tuple(all_behaviors), child
    )
    prog_pairs = list(zip(_prog[::2], _prog[1::2]))
    for left, right in prog_pairs:
        row = [
            [Paragraph(f"<b>{left[0]} {left[1]}</b>",
              S("PC", fontName="Helvetica-Bold", fontSize=9, textColor=INK)),
             Paragraph(left[2], sSmall)],
            [Paragraph(f"<b>{right[0]} {right[1]}</b>",
              S("PC2", fontName="Helvetica-Bold", fontSize=9, textColor=INK)),
             Paragraph(right[2], sSmall)],
        ]
        pt = Table(
            [[row[0][0], row[1][0]], [row[0][1], row[1][1]]],
            colWidths=[8.1*cm, 8.1*cm]
        )
        pt.setStyle(TableStyle([
            ("BACKGROUND",   (0,0),(-1,-1), LINEN),
            ("BOX",          (0,0),(0,-1), 0.4, BORDER),
            ("BOX",          (1,0),(1,-1), 0.4, BORDER),
            ("TOPPADDING",   (0,0),(-1,-1), 6),
            ("BOTTOMPADDING",(0,0),(-1,-1), 4),
            ("LEFTPADDING",  (0,0),(-1,-1), 10),
            ("VALIGN",       (0,0),(-1,-1), "TOP"),
        ]))
        story += [pt, sp(0.1)]
    story.append(sp(0.3))

    # -- Section 7 — Collaboration ---------------------------------------------
    story += [sec_hdr(7, "Family & Team Collaboration"), sp(0.15)]
    story.append(Paragraph("This programme operates on the understanding that meaningful and lasting "
                            "developmental change occurs when the child's key environments are aligned in their "
                            "approach and consistent in their response.", sBody))

    collab = [
        ("Nursery Team", "Teaching staff are briefed on programme goals and monitoring frameworks. "
                         "Fortnightly check-ins ensure responsiveness and fidelity to the plan."),
        ("Behavioral Therapist", f"{'The therapist assigned to this programme is ' + d.get('therapist_name') + '.' if d.get('therapist_name') else 'The assigned behavioral therapist'} maintains session monitoring records and serves as "
                         "primary liaison between the programme and the classroom environment."),
        ("Family Partnership", "Parents receive regular written updates and recommended home reinforcement "
                               "strategies. Family input is actively valued and integral to all reviews."),
    ]
    col_data = [[
        [Paragraph(f"<b>{t}</b>", S("CT", fontName="Helvetica-Bold", fontSize=9, textColor=INK)),
         Paragraph(b, sSmall)]
        for t, b in collab
    ]]
    flat = [[col_data[0][i][0] for i in range(3)], [col_data[0][i][1] for i in range(3)]]
    ctbl = Table(flat, colWidths=[5.4*cm, 5.4*cm, 5.4*cm])
    ctbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0),(-1,-1), LINEN),
        ("GRID",         (0,0),(-1,-1), 0.4, BORDER),
        ("TOPPADDING",   (0,0),(-1,-1), 7),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING",  (0,0),(-1,-1), 8),
        ("VALIGN",       (0,0),(-1,-1), "TOP"),
    ]))
    story += [ctbl, sp(0.2)]

    if d.get("notes") and d["notes"].strip():
        story += [
            Paragraph("<b>Additional Notes</b>",
                      S("AN", fontName="Helvetica-Bold", fontSize=9.5, textColor=BARK,
                        spaceBefore=10, spaceAfter=4)),
            Paragraph(d["notes"], sBody),
        ]

    # -- Footer ----------------------------------------------------------------
    story += [
        sp(0.5), hr(STONE, 0.5),
        Paragraph(f"{footer_t}  |  Generated: {today}  |  {nursery}", sFooter),
    ]

    def add_page_number(canvas_obj, doc_obj):
        canvas_obj.saveState()
        canvas_obj.setFont("Helvetica", 8)
        canvas_obj.setFillColor(STONE)
        page_num = canvas_obj.getPageNumber()
        text = f"Page {page_num}"
        canvas_obj.drawRightString(A4[0] - 2.2*cm, 1.2*cm, text)
        canvas_obj.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return buf.getvalue()


# -----------------------------------------------------------------------------
# SUGGESTED BEHAVIORAL ACTIVITIES (for staff PDF)
# -----------------------------------------------------------------------------
SUGGESTED_ACTIVITIES = [
    (
        "Emotion Volcano 🌋",
        "Emotional Regulation / Self-Awareness",
        "Draw a volcano together and label the stages: calm (bottom), building up (middle), erupting (top). "
        "Ask the child to identify where they are on the volcano throughout the session. "
        "Practice 'cooling lava' strategies: slow belly breaths, squeezing a stress ball, or shaking hands and letting go. "
        "Over time, aim for the child to self-identify and self-regulate before reaching the 'eruption' stage.",
    ),
    (
        "Freeze & Think 🛑",
        "Impulse Control",
        "Play music and move around the room. When the music stops, freeze completely. "
        "While frozen, the child must answer a 'think' question (e.g. 'What would you do if a friend took your toy?'). "
        "Gradually extend freeze time. Use visual stop/go cards to generalise the skill to classroom triggers.",
    ),
    (
        "The Waiting Jar ⏳",
        "Frustration Tolerance / Delayed Gratification",
        "Place a small treat or sticker inside a clear jar. Tell the child they can have it at the end of the session "
        "if they practice waiting when something feels hard. Each time the child waits without escalating, add a "
        "marble or token to a second jar so they can see progress visually. Celebrate each success explicitly.",
    ),
    (
        "Feelings Detectives 🔍",
        "Emotional Literacy / Communication",
        "Use picture cards of faces showing different emotions. Ask the child to guess the feeling and think of a time "
        "they felt that way. Role-play what that character could say or do instead of acting out. "
        "Progress to using cartoon scenarios closer to real classroom situations.",
    ),
    (
        "The Replacement Behavior Rehearsal 🔄",
        "Functional Communication / Replacement Behaviors",
        "Identify the top two or three problem behaviors and their typical triggers. "
        "Explicitly teach and rehearse one replacement behavior for each (e.g., tapping the table instead of hitting; "
        "saying 'I need a break' instead of running away). Practice in low-stakes role-play first, "
        "then gradually introduce mild stressors to build generalisation.",
    ),
    (
        "Compliment Catch 🌟",
        "Social Skills / Peer Interaction",
        "Sit with the child and, where safe, a peer or nursery adult. Each person catches a soft ball and must say "
        "one genuine compliment to the next person before throwing. Begin therapist-to-child only, then introduce "
        "a trusted peer. Debrief: 'How did it feel to give/receive a compliment?'",
    ),
    (
        "Calm Corner Design 🎨",
        "Coping Strategy Development / Self-Regulation",
        "Collaboratively design the child's own 'calm corner' toolkit on paper or in a small box. "
        "Include: a chosen breathing exercise card, a sensory item, a picture of something calming, and a feelings chart. "
        "Practice using each item during the session so the child can access them independently in the classroom.",
    ),
    (
        "Puppet Problem-Solving 🎭",
        "Social Skills / Conflict Resolution",
        "Use puppets or soft toys to act out a common conflict scenario (e.g., two puppets want the same toy). "
        "Guide the child to direct the puppets to a solution step-by-step: stop → feel → think → act. "
        "Switch roles so the child also plays the 'problem-solving coach'. Link solutions to real nursery situations.",
    ),
    (
        "Body Check-In 🧘",
        "Emotional Regulation / Physiological Awareness",
        "Begin each session with a brief 'body scan': starting at the feet, notice what each part of the body feels. "
        "Introduce vocabulary (tight shoulders = nervous; heavy arms = tired; wobbly tummy = worried). "
        "Create a simple body map and mark where feelings live. "
        "Repeat at the end of the session to notice any changes after regulation practice.",
    ),
    (
        "The Success Story Book 📖",
        "Self-Awareness / Progress Tracking",
        "Each session, help the child dictate or draw one thing they did well this week — however small. "
        "Collect these in a personalised booklet. Review together to build a narrative of competence and growth. "
        "Share (with the child's permission) one page with parents each fortnight to reinforce the message at home.",
    ),
]


# -----------------------------------------------------------------------------
# STAFF PDF (with suggested activities appended)
# -----------------------------------------------------------------------------
def build_pdf_staff(d) -> bytes:
    """Wraps build_pdf and appends a 'Suggested Behavioral Activities' section."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, HRFlowable, PageBreak
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from io import BytesIO

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=2.2*cm, rightMargin=2.2*cm,
                             topMargin=2.5*cm, bottomMargin=2.5*cm)

    INK   = colors.HexColor("#2A2520")
    BARK  = colors.HexColor("#8B7355")
    STONE = colors.HexColor("#6B6456")
    LINEN = colors.HexColor("#F0EBE1")
    WARM  = colors.HexColor("#FAF8F4")
    WARM2 = colors.HexColor("#F5F0E8")
    BORDER= colors.HexColor("#DDD8CE")
    SAGE  = colors.HexColor("#7A8A72")
    WHITE = colors.white

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    sTitle  = S("AT", fontName="Helvetica-Bold",  fontSize=18, textColor=INK,
                      leading=24, alignment=TA_CENTER, spaceAfter=4)
    sSub    = S("AS", fontName="Helvetica",        fontSize=10, textColor=STONE,
                      leading=14, alignment=TA_CENTER, spaceAfter=10)
    sHead   = S("AH", fontName="Helvetica-Bold",   fontSize=11, textColor=INK,
                      leading=15, spaceBefore=6, spaceAfter=3)
    sDomain = S("AD", fontName="Helvetica-Oblique",fontSize=9,  textColor=BARK,
                      leading=13, spaceAfter=4)
    sBody   = S("AB", fontName="Helvetica",         fontSize=9.5, textColor=INK,
                      leading=15, spaceAfter=4)
    sFooter = S("AF", fontName="Helvetica-Oblique", fontSize=7.5, textColor=STONE,
                      alignment=TA_CENTER)
    sNum    = S("AN2",fontName="Helvetica-Bold",    fontSize=22, textColor=BARK,
                      leading=26, spaceAfter=0)

    def hr(c=BARK, t=0.8):
        return HRFlowable(width="100%", thickness=t, color=c, spaceAfter=5, spaceBefore=5)
    def sp(h=0.25):
        return Spacer(1, h*cm)

    # Reuse the main plan story from build_pdf, then append activities
    # We rebuild here using build_pdf as source but add the activities section
    # First get the main PDF bytes and then append — but since ReportLab doesn't
    # support merging easily, we rebuild fully: call build_pdf internals inline.

    today = datetime.now().strftime("%d %B %Y")
    nursery = d.get("nursery_name") or "Nursery"
    footer_t = d.get("footer_text") or "Confidential — For internal use only."
    child = d.get("child_name") or "—"

    # Get main plan story by calling the internal build
    main_bytes = build_pdf(d)

    # Now build the activities-only supplement and merge with PyPDF
    act_buf = BytesIO()
    act_doc = SimpleDocTemplate(act_buf, pagesize=A4,
                                 leftMargin=2.2*cm, rightMargin=2.2*cm,
                                 topMargin=2.5*cm, bottomMargin=2.5*cm)

    story = []
    story += [
        Paragraph("Suggested Behavioral Activities", sTitle),
        Paragraph(f"Session Resource Supplement — {child} — {nursery}", sSub),
        sp(0.2), hr(INK, 1.5), sp(0.3),
        Paragraph(
            "The following activities are recommended as structured, evidence-informed tools to support "
            "the skills targeted in this behaviour support plan. Each activity is designed to be engaging, "
            "developmentally appropriate, and adaptable to the child's current level. Activities should be "
            "selected based on the child's focus areas and may be used across multiple sessions.",
            sBody),
        sp(0.4),
    ]

    for i, (name, domain, description) in enumerate(SUGGESTED_ACTIVITIES, 1):
        act_row = [[
            Paragraph(str(i), sNum),
            [Paragraph(name, sHead),
             Paragraph(f"Target Area: {domain}", sDomain),
             Paragraph(description, sBody)]
        ]]
        act_tbl = Table(act_row, colWidths=[1.2*cm, 15.0*cm])
        act_tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0,0),(-1,-1), LINEN if i % 2 == 1 else WARM2),
            ("LINEAFTER",    (0,0),(0,0), 3, BARK),
            ("TOPPADDING",   (0,0),(-1,-1), 10),
            ("BOTTOMPADDING",(0,0),(-1,-1), 10),
            ("LEFTPADDING",  (0,0),(-1,-1), 10),
            ("VALIGN",       (0,0),(-1,-1), "TOP"),
            ("BOX",          (0,0),(-1,-1), 0.4, BORDER),
        ]))
        story += [act_tbl, sp(0.2)]

    story += [
        sp(0.5), hr(STONE, 0.5),
        Paragraph(f"{footer_t}  |  Behavioral Activities Supplement  |  Generated: {today}  |  {nursery}", sFooter),
    ]

    def add_page_num(canvas_obj, doc_obj):
        canvas_obj.saveState()
        canvas_obj.setFont("Helvetica", 8)
        canvas_obj.setFillColor(STONE)
        canvas_obj.drawRightString(A4[0] - 2.2*cm, 1.2*cm, f"Page {canvas_obj.getPageNumber()}")
        canvas_obj.restoreState()

    act_doc.build(story, onFirstPage=add_page_num, onLaterPages=add_page_num)
    act_bytes = act_buf.getvalue()

    # Merge: main plan + activities supplement
    try:
        from pypdf import PdfWriter, PdfReader
        writer = PdfWriter()
        for src in [main_bytes, act_bytes]:
            reader = PdfReader(BytesIO(src))
            for page in reader.pages:
                writer.add_page(page)
        merged = BytesIO()
        writer.write(merged)
        return merged.getvalue()
    except Exception:
        # Fallback: just return the main PDF if merge fails
        return main_bytes


# -----------------------------------------------------------------------------
# PARENT-FRIENDLY PDF
# -----------------------------------------------------------------------------
def build_pdf_parent(d) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, HRFlowable, Image as RLImage
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from io import BytesIO

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=2.5*cm, rightMargin=2.5*cm,
                             topMargin=2.5*cm, bottomMargin=2.5*cm)

    INK   = colors.HexColor("#2A2520")
    BARK  = colors.HexColor("#8B7355")
    STONE = colors.HexColor("#6B6456")
    LINEN = colors.HexColor("#F0EBE1")
    WARM  = colors.HexColor("#FAF8F4")
    WARM2 = colors.HexColor("#F5F0E8")
    BORDER= colors.HexColor("#DDD8CE")
    SAGE  = colors.HexColor("#7A8A72")
    WHITE = colors.white
    SAGE_BG = colors.HexColor("#EEF2EC")

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    sTitle  = S("PT", fontName="Helvetica-Bold",  fontSize=22, textColor=INK,
                      leading=28, alignment=TA_CENTER, spaceAfter=6)
    sSub    = S("PS", fontName="Helvetica",        fontSize=11, textColor=STONE,
                      leading=16, alignment=TA_CENTER, spaceAfter=4)
    sLead   = S("PL", fontName="Helvetica",        fontSize=11, textColor=INK,
                      leading=18, spaceAfter=8, alignment=TA_JUSTIFY)
    sSecHdr = S("PH", fontName="Helvetica-Bold",   fontSize=13, textColor=BARK,
                      leading=18, spaceBefore=14, spaceAfter=6)
    sBody   = S("PB", fontName="Helvetica",         fontSize=10.5, textColor=INK,
                      leading=17, spaceAfter=6, alignment=TA_JUSTIFY)
    sBullet = S("PBu",fontName="Helvetica",         fontSize=10.5, textColor=INK,
                      leading=17, leftIndent=14, spaceAfter=4)
    sEmph   = S("PE", fontName="Helvetica-Bold",    fontSize=10.5, textColor=BARK,
                      leading=17, spaceAfter=4)
    sSmall  = S("PSm",fontName="Helvetica",         fontSize=9,  textColor=STONE, leading=13)
    sFooter = S("PF", fontName="Helvetica-Oblique", fontSize=8,  textColor=STONE,
                      alignment=TA_CENTER)
    sBox    = S("PBx",fontName="Helvetica",         fontSize=10.5, textColor=INK,
                      leading=17, leftIndent=6, spaceAfter=4, alignment=TA_JUSTIFY)

    def hr(c=BARK, t=0.8):
        return HRFlowable(width="100%", thickness=t, color=c, spaceAfter=6, spaceBefore=6)
    def sp(h=0.3):
        return Spacer(1, h*cm)

    today    = datetime.now().strftime("%d %B %Y")
    child    = d.get("child_name") or "Your child"
    nursery  = d.get("nursery_name") or "Nursery"
    therapist= d.get("therapist_name") or "the assigned behavioral therapist"
    dob_str  = d["date_of_birth"].strftime("%d %B %Y") if d.get("date_of_birth") else "—"
    age      = d.get("age", "—")
    group    = d.get("class_group", "—")
    start    = d["start_date"].strftime("%d %B %Y") if d.get("start_date") else "—"
    freq     = d.get("session_frequency", "—")
    spm      = str(d.get("sessions_per_month", "—"))
    pattern  = d.get("pattern", "")
    notes    = d.get("notes", "")
    all_behaviors = d.get("behaviors", []) + ([d["custom_behavior"]] if d.get("custom_behavior") else [])
    skills   = d.get("skills", [])
    all_skills = skills + ([d["custom_skill"]] if d.get("custom_skill") else [])
    footer_t = d.get("footer_text") or "Confidential — For family use."

    story = []

    # -- Header ----------------------------------------------------------------
    story += [
        Paragraph(f"Your Child's Support Plan", sTitle),
        Paragraph(f"{nursery} — Individual Behaviour Support Programme", sSub),
        sp(0.15), hr(INK, 1.5), sp(0.3),
    ]

    # Logo
    if d.get("logo_bytes"):
        try:
            img = RLImage(BytesIO(d["logo_bytes"]), width=3*cm, height=1.8*cm, kind="proportional")
            img.hAlign = "CENTER"
            story += [img, sp(0.2)]
        except Exception:
            pass

    # Warm intro
    story += [
        Paragraph(
            f"Dear Parent / Guardian,",
            S("PG", fontName="Helvetica-Bold", fontSize=11, textColor=INK, spaceAfter=6)),
        Paragraph(
            f"We are so glad to share this with you. This document has been prepared to keep you fully informed "
            f"about the individual support programme we have put in place for <b>{child}</b>. "
            f"It is written in plain language — no jargon — so that you can understand exactly what we are doing, "
            f"why we are doing it, and how you can be part of {child}'s journey.",
            sLead),
        Paragraph(
            f"Please know that this plan is a sign of our care and commitment, not a cause for worry. "
            f"Many children benefit enormously from a little extra, focused support during their early years. "
            f"We are here as partners every step of the way.",
            sLead),
        sp(0.2),
    ]

    # Child info table
    info_data = [
        [Paragraph("<b>Child</b>", sSmall),       Paragraph(child, sBody),
         Paragraph("<b>Date of Birth</b>", sSmall), Paragraph(dob_str, sBody)],
        [Paragraph("<b>Age</b>", sSmall),          Paragraph(age, sBody),
         Paragraph("<b>Class / Group</b>", sSmall), Paragraph(group, sBody)],
        [Paragraph("<b>Therapist</b>", sSmall),    Paragraph(therapist, sBody),
         Paragraph("<b>Programme Start</b>", sSmall), Paragraph(start, sBody)],
        [Paragraph("<b>Sessions / Month</b>", sSmall), Paragraph(spm, sBody),
         Paragraph("<b>Session Frequency</b>", sSmall), Paragraph(freq, sBody)],
    ]
    info_tbl = Table(info_data, colWidths=[3.5*cm, 5.5*cm, 3.5*cm, 4.0*cm])
    info_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(0,-1), LINEN),
        ("BACKGROUND",    (2,0),(2,-1), LINEN),
        ("ROWBACKGROUNDS",(0,0),(-1,-1), [WARM, WARM2, WARM, WARM2]),
        ("GRID",          (0,0),(-1,-1), 0.4, BORDER),
        ("TOPPADDING",    (0,0),(-1,-1), 6),
        ("BOTTOMPADDING", (0,0),(-1,-1), 6),
        ("LEFTPADDING",   (0,0),(-1,-1), 8),
        ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
    ]))
    story += [info_tbl, sp(0.5)]

    # -- Section 1: What is this plan? ----------------------------------------
    story += [Paragraph("What Is This Support Plan?", sSecHdr), hr(BARK, 0.5), sp(0.1)]
    story += [
        Paragraph(
            f"This is an <b>Individual Behaviour Support Plan</b> — a personalised programme designed specifically "
            f"for <b>{child}</b>. It brings together structured sessions, targeted skill-building, and close teamwork "
            f"between our team and your family.",
            sBody),
        Paragraph(
            f"<b>It is important to understand:</b> this plan is <i>not</i> a diagnosis. It does not mean something "
            f"is 'wrong' with {child}. It simply means we have noticed some areas where a little extra support "
            f"could make a big difference — and we want to provide that support in the most effective way possible.",
            sBody),
        sp(0.2),
    ]

    # -- Section 2: What behaviors are we focusing on? -------------------------
    story += [Paragraph("What Are We Working On?", sSecHdr), hr(BARK, 0.5), sp(0.1)]
    story += [
        Paragraph(
            f"Our team has observed the following behaviours that we would like to support {child} with. "
            f"Please remember — these behaviours are <i>not</i> 'bad behaviour'. They are {child}'s way of telling "
            f"us that something feels hard right now. Our job is to help build the skills to handle those "
            f"hard moments better:",
            sBody),
    ]
    if all_behaviors:
        for b in all_behaviors:
            story.append(Paragraph(f"• {b}", sBullet))
    else:
        story.append(Paragraph("• To be specified with the team.", sBullet))

    freq_label = d.get("frequency", "moderate").capitalize()
    int_label  = d.get("intensity", "moderate").capitalize()
    story += [
        sp(0.1),
        Paragraph(f"<b>How often:</b> {freq_label} &nbsp;&nbsp; <b>Intensity:</b> {int_label}", sBody),
    ]
    if pattern:
        story += [
            Paragraph(f"<b>When does it tend to happen?</b> {pattern}", sBody),
        ]
    story.append(sp(0.3))

    # -- Section 3: What skills are we building? -------------------------------
    story += [Paragraph("What Skills Are We Building?", sSecHdr), hr(BARK, 0.5), sp(0.1)]
    story += [
        Paragraph(
            f"Our sessions are designed to build specific skills that will help {child} feel more confident, "
            f"calm, and capable. Here is what we are working on and what it means:",
            sBody),
    ]
    if all_skills:
        for sk in all_skills:
            if sk in SKILL_DETAILS:
                _, body = SKILL_DETAILS[sk]
                story += [
                    Paragraph(f"✦  {sk}", sEmph),
                    Paragraph(body, sBox),
                    sp(0.1),
                ]
            else:
                story += [
                    Paragraph(f"✦  {sk}", sEmph),
                    Paragraph("This is a custom skill area we are working on with your child. "
                              "Your child's therapist will share specific goals and home strategies with you directly.", sBox),
                    sp(0.1),
                ]
    else:
        story.append(Paragraph("Skills will be confirmed in your first review meeting.", sBody))
    story.append(sp(0.2))

    # -- Section 4: What do sessions look like? --------------------------------
    story += [Paragraph("What Happens in a Session?", sSecHdr), hr(BARK, 0.5), sp(0.1)]
    story += [
        Paragraph(
            f"Sessions are warm, playful, and child-led as much as possible. {child} will not know they are in "
            f"a 'behaviour session' — they will just be playing and connecting with {therapist}. "
            f"Sessions happen <b>{freq}</b> ({spm} times per month).",
            sBody),
    ]
    _parent_acts = generate_dynamic_activities(
        tuple(all_skills), tuple(all_behaviors), child
    )
    for t_, b_ in _parent_acts:
        story.append(Paragraph(f"• <b>{t_}:</b> {b_}", sBullet))
    story.append(sp(0.3))

    # -- Section 5: How will we know it's working? ----------------------------
    story += [Paragraph("How Will We Know It's Working?", sSecHdr), hr(BARK, 0.5), sp(0.1)]
    story += [
        Paragraph(
            f"Progress in early childhood is rarely a straight line — and that is completely normal. "
            f"We will look for small, meaningful signs that {child} is growing. Here is what to watch for:",
            sBody),
    ]
    _parent_prog = generate_dynamic_progress(
        tuple(all_skills), tuple(all_behaviors), child
    )
    for _, t_, b_ in _parent_prog:
        story.append(Paragraph(f"• <b>{t_}:</b> {b_}", sBullet))
    story.append(sp(0.3))

    # -- Section 6: How can YOU help at home? ----------------------------------
    story += [Paragraph("How Can You Help at Home?", sSecHdr), hr(BARK, 0.5), sp(0.1)]

    home_tips = [
        ("Stay consistent", f"Try to use the same calm, non-reactive approach when {child} struggles at home as we use here. "
                            "Consistency between home and nursery is one of the most powerful things you can do."),
        ("Name feelings out loud", f"When {child} seems upset, try saying 'I can see you feel frustrated' rather than 'stop that'. "
                                   "This builds emotional vocabulary and makes feelings feel safe to express."),
        ("Celebrate the small wins", "Every time you notice {child} waiting, communicating, or calming down — even a little — "
                                     "name it and celebrate it. 'I noticed you took a deep breath — that was brilliant.'"),
        ("Tell us what you notice", "You know your child best. If you see something at home — a new trigger, a new strategy that "
                                    "works, or a breakthrough — please share it with us. Your insights shape the programme."),
        ("Keep routines predictable", "Many children who struggle with regulation find surprise and change very hard. "
                                      "Where possible, give {child} advance notice of any changes to routine."),
        ("Ask for support too", "This can feel a lot for parents. Please reach out to us any time you need guidance, "
                                "reassurance, or just a conversation."),
    ]
    for title, tip in home_tips:
        tip_filled = tip.replace("{child}", child)
        story.append(Paragraph(f"✦  <b>{title}</b>", sEmph))
        story.append(Paragraph(tip_filled, sBox))
        story.append(sp(0.1))
    story.append(sp(0.2))

    # -- Section 7: Who to contact --------------------------------------------
    story += [Paragraph("Your Team", sSecHdr), hr(BARK, 0.5), sp(0.1)]
    contact_data = [
        [Paragraph("<b>🏫 Nursery Team</b>", sEmph),
         Paragraph("<b>👩‍🏫 Behavioral Therapist</b>", sEmph),
         Paragraph("<b>👨‍👩‍👧 You — the Family</b>", sEmph)],
        [Paragraph("The nursery team will keep you informed and support the plan throughout the day.", sSmall),
         Paragraph(f"{therapist} leads the sessions and will provide you with regular written updates.", sSmall),
         Paragraph(f"Your voice matters. Regular review meetings will include your feedback and observations.", sSmall)],
    ]
    ct = Table(contact_data, colWidths=[5.2*cm, 5.2*cm, 5.2*cm])
    ct.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,0), INK),
        ("TEXTCOLOR",     (0,0),(-1,0), WHITE),
        ("BACKGROUND",    (0,1),(-1,-1), LINEN),
        ("GRID",          (0,0),(-1,-1), 0.4, BORDER),
        ("TOPPADDING",    (0,0),(-1,-1), 8),
        ("BOTTOMPADDING", (0,0),(-1,-1), 8),
        ("LEFTPADDING",   (0,0),(-1,-1), 8),
        ("VALIGN",        (0,0),(-1,-1), "TOP"),
    ]))
    story += [ct, sp(0.3)]

    # Notes
    if notes and notes.strip():
        story += [
            Paragraph("<b>Additional Notes from the Team</b>",
                      S("PN", fontName="Helvetica-Bold", fontSize=10.5, textColor=BARK,
                        spaceBefore=8, spaceAfter=4)),
            Paragraph(notes, sBody),
            sp(0.2),
        ]

    # Closing warm message
    closing_row = [[Paragraph(
        f"We are proud to support {child} and grateful to have you as a partner in this journey. "
        f"Together, we can make a real difference. Please don't hesitate to reach out — our door is always open.",
        S("PC2", fontName="Helvetica-Oblique", fontSize=10.5, textColor=STONE, leading=17, alignment=TA_CENTER))]]
    ct2 = Table(closing_row, colWidths=[15.5*cm])
    ct2.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,-1), SAGE_BG),
        ("BOX",           (0,0),(-1,-1), 0.5, SAGE),
        ("TOPPADDING",    (0,0),(-1,-1), 14),
        ("BOTTOMPADDING", (0,0),(-1,-1), 14),
        ("LEFTPADDING",   (0,0),(-1,-1), 16),
        ("RIGHTPADDING",  (0,0),(-1,-1), 16),
    ]))
    story += [ct2, sp(0.4)]

    # Footer
    story += [
        hr(STONE, 0.5),
        Paragraph(f"{footer_t}  |  Prepared for family of {child}  |  {today}  |  {nursery}", sFooter),
    ]

    def add_page_num(canvas_obj, doc_obj):
        canvas_obj.saveState()
        canvas_obj.setFont("Helvetica", 8)
        canvas_obj.setFillColor(STONE)
        canvas_obj.drawRightString(A4[0] - 2.5*cm, 1.2*cm, f"Page {canvas_obj.getPageNumber()}")
        canvas_obj.restoreState()

    doc.build(story, onFirstPage=add_page_num, onLaterPages=add_page_num)
    return buf.getvalue()



def build_docx(d) -> bytes:
    from docx import Document as DocxDoc
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = DocxDoc()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # -- Helpers ----------------------------------------------------------------
    def set_para_spacing(para, before=0, after=6, line=None):
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(before))
        spacing.set(qn("w:after"),  str(after))
        if line:
            spacing.set(qn("w:line"), str(line))
            spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0x8B, 0x73, 0x55) if level == 2 else RGBColor(0x2A, 0x25, 0x20)
        run.font.size = Pt(14 if level == 1 else 11)
        set_para_spacing(p, before=120, after=60)
        return p

    def add_section_header(num, title):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        run_num = p.add_run(f"{num}. ")
        run_num.bold = True
        run_num.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
        run_num.font.size = Pt(11)
        run_title = p.add_run(title)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(0x2A, 0x25, 0x20)
        run_title.font.size = Pt(11)
        set_para_spacing(p, before=160, after=60)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "8B7355")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x3A, 0x35, 0x30)
        set_para_spacing(p, after=60, line=276)
        return p

    def add_bullet(title, body_text):
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(f"{title}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(body_text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x3A, 0x35, 0x30)
        set_para_spacing(p, after=40)
        return p

    def add_spacer():
        p = doc.add_paragraph()
        set_para_spacing(p, after=40)
        return p

    def shade_cell(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"),    str(val))
            el.set(qn("w:type"), "dxa")
            tcMar.append(el)
        tcPr.append(tcMar)

    today  = datetime.now().strftime("%d %B %Y")
    start  = d["start_date"].strftime("%d %B %Y") if d.get("start_date") else "—"
    all_b  = d["behaviors"] + ([d["custom_behavior"]] if d.get("custom_behavior") else [])
    all_sk = d.get("skills", []) + ([d["custom_skill"]] if d.get("custom_skill") else [])
    nursery = d.get("nursery_name") or "Nursery"
    child   = d.get("child_name")   or "—"
    footer_t = d.get("footer_text") or "Confidential — For internal and family use only."

    # -- Title -----------------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Individual Behavior Support Plan")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x2A, 0x25, 0x20)
    set_para_spacing(t, after=40)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(f"Developmental Support Programme — {nursery}")
    rs.font.size = Pt(10)
    rs.font.color.rgb = RGBColor(0x6B, 0x64, 0x56)
    rs.italic = True
    set_para_spacing(sub, after=100)

    # -- Info table -------------------------------------------------------------
    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"
    info_data = [
        ("Child Name", child,                      "Date",           today),
        ("Age",        d.get("age","—"),            "Class / Group",  d.get("class_group","—")),
        ("Start Date", start,                       "Sessions/month", str(d.get("sessions_per_month","—"))),
    ]
    col_w = [Cm(3.0), Cm(5.8), Cm(3.0), Cm(4.5)]
    for i, (l1, v1, l2, v2) in enumerate(info_data):
        row = tbl.rows[i]
        for j, (w, txt, is_label) in enumerate([(col_w[0],l1,True),(col_w[1],v1,False),
                                                  (col_w[2],l2,True),(col_w[3],v2,False)]):
            cell = row.cells[j]
            cell.width = w
            set_cell_margins(cell)
            shade_cell(cell, "F0EBE1" if is_label else "FAF8F4")
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = is_label
            run.font.size = Pt(9)
            if is_label:
                run.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)

    add_spacer()

    # -- Section 1 — Purpose ----------------------------------------------------
    add_section_header(1, "Programme Purpose & Scope")
    add_body(f"This Individual Behavior Support Plan has been developed to provide structured, "
             f"evidence-informed support for {child}'s developmental needs. It is important to note that "
             f"this plan is not a diagnostic document, nor does it imply the presence of any clinical "
             f"condition or disorder.")
    add_body("The purpose of this programme is to build specific developmental skills — including "
             "emotional regulation, impulse control, frustration tolerance, and functional communication — "
             "within the safe, consistent, and nurturing environment of the nursery setting.")
    add_body("All interventions are grounded in a positive, strengths-based framework that emphasises "
             "the child's capacity for growth. Progress is expected to occur gradually and non-linearly.")
    add_spacer()

    # -- Section 2 — Behaviors -------------------------------------------------
    add_section_header(2, "Identified Behavioral Concerns")
    add_body("The following behaviors have been identified as current areas of focus:")
    if all_b:
        p = doc.add_paragraph()
        for b in all_b:
            r = p.add_run(f"  {b}  ")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x6B, 0x64, 0x56)
        set_para_spacing(p, after=40)

    p_lv = doc.add_paragraph()
    r1 = p_lv.add_run("Frequency: ")
    r1.bold = True; r1.font.size = Pt(10)
    r2 = p_lv.add_run(d.get("frequency","—").capitalize() + "   ")
    r2.font.size = Pt(10)
    r3 = p_lv.add_run("Intensity: ")
    r3.bold = True; r3.font.size = Pt(10)
    r4 = p_lv.add_run(d.get("intensity","—").capitalize())
    r4.font.size = Pt(10)
    set_para_spacing(p_lv, after=40)

    if d.get("pattern"):
        p_pat = doc.add_paragraph()
        rp1 = p_pat.add_run("Recent Pattern / Context: ")
        rp1.bold = True; rp1.font.size = Pt(10)
        rp2 = p_pat.add_run(d["pattern"])
        rp2.font.size = Pt(10)
        set_para_spacing(p_pat, after=60)
    add_spacer()

    # -- Section 3 — Skills ----------------------------------------------------
    add_section_header(3, "Core Skills Targeted")
    add_body("The following skill areas have been selected as primary targets for this programme:")
    all_skills_docx = d.get("skills", []) + ([d["custom_skill"]] if d.get("custom_skill") else [])
    for s in all_skills_docx:
        if s in SKILL_DETAILS:
            title, body_text = SKILL_DETAILS[s]
        else:
            title = s
            body_text = ("This is a custom skill area identified for this child's programme. "
                         "Specific goals and strategies will be developed collaboratively with the team.")
        p_sk = doc.add_paragraph()
        r_t = p_sk.add_run(f"{title}:  ")
        r_t.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(0x2A, 0x25, 0x20)
        r_b = p_sk.add_run(body_text)
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = RGBColor(0x4A, 0x45, 0x40)
        # Light background shading via paragraph border
        pPr = p_sk._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),  "clear")
        shd.set(qn("w:fill"), "F0EBE1")
        pPr.append(shd)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "200")
        pPr.append(ind)
        set_para_spacing(p_sk, after=60, line=276)
    add_spacer()

    # -- Section 4 — Activities ------------------------------------------------
    add_section_header(4, "Session Activities")
    add_body("Sessions will incorporate structured, play-based, and reflective activities "
             "designed to address the targeted skills and behaviors:")
    _docx_acts = generate_dynamic_activities(tuple(all_sk), tuple(all_b), child)
    for t_, b_ in _docx_acts:
        add_bullet(t_, b_)
    add_spacer()

    # -- Section 5 — Monitoring ------------------------------------------------
    add_section_header(5, "Session Monitoring Framework")
    mon_tbl = doc.add_table(rows=len(MONITORING_ROWS)+1, cols=2)
    mon_tbl.style = "Table Grid"
    hdr_cells = mon_tbl.rows[0].cells
    for i, txt in enumerate(["Monitoring Dimension", "What Is Recorded"]):
        shade_cell(hdr_cells[i], "2A2520")
        set_cell_margins(hdr_cells[i])
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, (d_, r_) in enumerate(MONITORING_ROWS):
        row = mon_tbl.rows[ri+1]
        fill = "FAF8F4" if ri % 2 == 0 else "F0EBE1"
        for ci, txt in enumerate([d_, r_]):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = (ci == 0)
            run.font.size = Pt(9.5)
    add_spacer()

    # -- Section 6 — Progress -------------------------------------------------
    add_section_header(6, "Indicators of Progress")
    add_body("The following indicators will be used as primary markers of meaningful progress:")
    _docx_prog = generate_dynamic_progress(tuple(all_sk), tuple(all_b), child)
    for _, t_, b_ in _docx_prog:
        add_bullet(t_, b_)
    add_spacer()

    # -- Section 7 — Collaboration ---------------------------------------------
    add_section_header(7, "Family & Team Collaboration")
    add_body("This programme operates on the understanding that meaningful and lasting developmental "
             "change occurs when the child's key environments are aligned in their approach and consistent "
             "in their response.")
    collab_items = [
        ("Nursery Team", "Teaching staff are briefed on programme goals and monitoring frameworks. "
                         "Fortnightly check-ins ensure responsiveness and fidelity to the plan."),
        ("Key Worker",   "The assigned key worker maintains session records and serves as primary liaison "
                         "between the programme and the classroom environment."),
        ("Family Partnership", "Parents receive regular written updates and recommended home reinforcement "
                               "strategies. Family input is actively valued and integral to all reviews."),
    ]
    for t_, b_ in collab_items:
        add_bullet(t_, b_)

    if d.get("notes") and d["notes"].strip():
        add_spacer()
        p_nt = doc.add_paragraph()
        r_nt = p_nt.add_run("Additional Notes:  ")
        r_nt.bold = True
        r_nt.font.size = Pt(10)
        r_nt.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
        r_nt2 = p_nt.add_run(d["notes"])
        r_nt2.font.size = Pt(10)
        set_para_spacing(p_nt, after=60)

    # -- Footer ----------------------------------------------------------------
    add_spacer()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Horizontal rule via top border
    pPr = p_foot._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top_el = OxmlElement("w:top")
    top_el.set(qn("w:val"),   "single")
    top_el.set(qn("w:sz"),    "4")
    top_el.set(qn("w:space"), "1")
    top_el.set(qn("w:color"), "DDD8CE")
    pBdr.append(top_el)
    pPr.append(pBdr)

    rf = p_foot.add_run(f"{footer_t}  |  Generated: {today}  |  {nursery}")
    rf.font.size = Pt(8)
    rf.italic = True
    rf.font.color.rgb = RGBColor(0x6B, 0x64, 0x56)
    set_para_spacing(p_foot, before=60)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# -----------------------------------------------------------------------------
# SIDEBAR — FORM
# -----------------------------------------------------------------------------
def render_sidebar() -> dict:
    with st.sidebar:
        st.markdown("""
        <div style="display:flex;align-items:center;gap:0.6rem;padding:0.5rem 0 1.2rem;">
            <div style="width:34px;height:34px;background:#2A2520;border-radius:8px;
                        display:flex;align-items:center;justify-content:center;
                        font-size:1rem;color:#fff;">✦</div>
            <div>
                <div style="font-family:'Jost',sans-serif;font-size:1rem;font-weight:500;color:#2A2520;line-height:1.2;">
                    Plan Builder</div>
                <div style="font-size:0.68rem;color:#6B6456;">Individual Support Programme</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Templates
        with st.expander("📋 Load a Template", expanded=False):
            if st.button("Transition Difficulties", use_container_width=True, key="tpl1"):
                st.session_state.tpl = "transition"
            if st.button("Aggression Toward Peers", use_container_width=True, key="tpl2"):
                st.session_state.tpl = "aggression"

        st.markdown("### Child Information")
        child_name  = st.text_input("Child Name *", placeholder="e.g. Liam Hassan",
                                     key="child_name")
        dob         = st.date_input("Date of Birth", value=date(date.today().year - 3, date.today().month, 1),
                                     key="date_of_birth", min_value=date(date.today().year-12,1,1),
                                     max_value=date.today())
        # Auto-calculate age
        today_d = date.today()
        years = today_d.year - dob.year - ((today_d.month, today_d.day) < (dob.month, dob.day))
        months_total = (today_d.year * 12 + today_d.month) - (dob.year * 12 + dob.month)
        months = months_total % 12
        age = f"{years} year{'s' if years != 1 else ''} {months} month{'s' if months != 1 else ''}"
        st.caption(f"🎂 Age: **{age}**")

        class_group = st.text_input("Class / Group", placeholder="e.g. Sunflower Group",
                                     key="class_group")
        therapist_name = st.text_input("Behavioral Therapist Name", placeholder="e.g. Sara Ahmed",
                                        key="therapist_name")
        start_date  = st.date_input("Programme Start Date", value=date.today(), key="start_date")

        st.markdown("### Behavioral Concerns")
        behaviors = st.multiselect(
            "Main behaviors",
            BEHAVIOR_OPTIONS,
            key="behaviors",
            placeholder="Select all that apply…",
        )
        custom_behavior = st.text_input("Custom behavior (optional)",
                                         placeholder="Describe a specific behavior…",
                                         key="custom_behavior")

        col1, col2 = st.columns(2)
        with col1:
            frequency = st.selectbox("Frequency", ["low","moderate","high"],
                                      index=1, key="frequency")
        with col2:
            intensity = st.selectbox("Intensity", ["low","moderate","high"],
                                      index=1, key="intensity")

        pattern = st.text_area("Recent pattern / context",
                                placeholder="e.g. Behaviors occur primarily during transitions…",
                                key="pattern", height=80)

        st.markdown("### Skills to Target")
        skills = []
        for s in SKILL_DETAILS.keys():
            if st.checkbox(s, key=f"skill_{s}"):
                skills.append(s)
        custom_skill = st.text_input("Other skill (optional)",
                                      placeholder="Describe a custom skill to target…",
                                      key="custom_skill")

        st.markdown("### Programme Details")
        session_frequency = st.selectbox(
            "Sessions per week",
            ["1 session per week", "2 sessions per week", "3 sessions per week"],
            key="session_frequency",
        )
        _freq_map = {
            "1 session per week": 4,
            "2 sessions per week": 8,
            "3 sessions per week": 12,
        }
        sessions_per_month = _freq_map[session_frequency]
        st.caption(f"📅 Total sessions planned per month: **{sessions_per_month}**")
        notes = st.text_area("Additional notes (optional)",
                              placeholder="Context, preferences, observations…",
                              key="notes", height=80)

        st.markdown("### Nursery Information")
        nursery_name = st.text_input("Nursery Name", placeholder="e.g. Little Stars Nursery",
                                      key="nursery_name")
        logo_file    = st.file_uploader("Logo (PNG / JPG)", type=["png","jpg","jpeg"],
                                         key="logo_file")
        footer_text  = st.text_input("Footer text (optional)",
                                      placeholder="Confidential — For internal use only.",
                                      key="footer_text")

        logo_bytes = None
        if logo_file:
            logo_bytes = logo_file.read()

        # Apply template if requested
        if st.session_state.get("tpl"):
            tpl = st.session_state.pop("tpl")
            if tpl == "transition":
                st.session_state.behaviors      = ["Refusing transitions","Tantrums / meltdowns","Dropping to the floor"]
                st.session_state.frequency      = "high"
                st.session_state.intensity      = "moderate"
                st.session_state.pattern        = "Behaviors occur primarily during transitions between activities and at pickup time."
                st.session_state["skill_Emotional Regulation"] = True
                st.session_state["skill_Frustration Tolerance"] = True
                st.session_state["skill_Communication Skills"] = True
                st.rerun()
            elif tpl == "aggression":
                st.session_state.behaviors      = ["Hitting / kicking others","Throwing objects","Verbal aggression"]
                st.session_state.frequency      = "moderate"
                st.session_state.intensity      = "high"
                st.session_state.pattern        = "Most incidents occur during unstructured play, particularly when sharing is required."
                st.session_state["skill_Impulse Control"] = True
                st.session_state["skill_Social Skills"] = True
                st.session_state["skill_Frustration Tolerance"] = True
                st.rerun()

        return {
            "child_name":        child_name,
            "date_of_birth":     dob,
            "age":               age,
            "class_group":       class_group,
            "therapist_name":    therapist_name,
            "start_date":        start_date,
            "behaviors":         behaviors,
            "custom_behavior":   custom_behavior,
            "frequency":         frequency,
            "intensity":         intensity,
            "pattern":           pattern,
            "skills":            skills,
            "custom_skill":      custom_skill,
            "sessions_per_month": sessions_per_month,
            "session_frequency": session_frequency,
            "notes":             notes,
            "nursery_name":      nursery_name,
            "logo_bytes":        logo_bytes,
            "footer_text":       footer_text,
        }


# -----------------------------------------------------------------------------
# ACCESS CODE GATE
# -----------------------------------------------------------------------------

def get_valid_codes() -> list:
    """
    Load access codes from st.secrets.
    Expected secrets.toml entry:
        ACCESS_CODES = "CODE1,CODE2,CODE3"
    Returns a list of stripped uppercase codes.
    """
    try:
        raw = st.secrets["ACCESS_CODES"]
        return [c.strip().upper() for c in raw.split(",") if c.strip()]
    except Exception:
        return []


def render_login_screen():
    """
    Full-page login screen shown to unauthenticated users.
    Returns True if the user just entered a valid code (triggers rerun).
    """
    st.markdown("""
    <style>
    /* Hide sidebar entirely on login screen */
    section[data-testid="stSidebar"] { display: none !important; }
    .main .block-container {
        max-width: 480px !important;
        margin: 0 auto !important;
        padding-top: 8vh !important;
    }

    .login-wrap {
        background: #fff;
        border: 1px solid #DDD8CE;
        border-radius: 16px;
        padding: 3rem 2.5rem 2.5rem;
        box-shadow: 0 8px 40px rgba(42,37,32,0.10);
        text-align: center;
    }
    .login-mark {
        width: 52px; height: 52px;
        background: #2A2520; border-radius: 12px;
        display: flex; align-items: center; justify-content: center;
        font-size: 1.4rem; color: #fff;
        margin: 0 auto 1.2rem;
    }
    .login-title {
        font-family: 'Cormorant Garamond', serif !important;
        font-size: 1.75rem; font-weight: 400;
        color: #2A2520; margin-bottom: 0.4rem;
        line-height: 1.2;
    }
    .login-sub {
        font-size: 0.85rem; color: #6B6456;
        margin-bottom: 2rem; line-height: 1.6;
    }
    .login-error {
        background: #FAECE8; border-radius: 8px;
        padding: 0.6rem 1rem; font-size: 0.83rem;
        color: #8A3A2A; margin-top: 0.75rem;
    }
    .login-footer {
        font-size: 0.72rem; color: #B09A7A;
        margin-top: 1.5rem; line-height: 1.5;
    }
    </style>

    <div class="login-wrap">
        <div class="login-mark">✦</div>
        <div class="login-title">Behavior Plan Builder</div>
        <div class="login-sub">
            This tool is for authorised nursery staff only.<br>
            Please enter your access code to continue.
        </div>
    </div>
    """, unsafe_allow_html=True)

    # The input sits outside the HTML block so Streamlit renders it properly
    code_input = st.text_input(
        "Access Code",
        placeholder="Enter your code…",
        type="password",
        label_visibility="collapsed",
        key="access_code_input",
    )

    btn = st.button("Enter →", use_container_width=True, key="access_code_btn")

    if btn or (code_input and code_input.endswith("\n")):
        valid_codes = get_valid_codes()
        if not valid_codes:
            # No codes configured — fail open with a warning (dev mode)
            st.warning("No ACCESS_CODES found in secrets. Add them to enable access control.")
            st.session_state.authenticated = True
            st.rerun()
        elif code_input.strip().upper() in valid_codes:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.markdown(
                '<div class="login-error">⚠  Incorrect access code. Please try again.</div>',
                unsafe_allow_html=True,
            )

    st.markdown(
        '<div class="login-footer">Contact your administrator if you need an access code.</div>',
        unsafe_allow_html=True,
    )


# -----------------------------------------------------------------------------
# MAIN
# -----------------------------------------------------------------------------
def main():
    inject_styles()

    # -- Access gate -----------------------------------------------------------
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if not st.session_state.authenticated:
        render_login_screen()
        st.stop()   # Nothing below renders until authenticated

    # Session state
    if "tpl" not in st.session_state:
        st.session_state.tpl = None

    # App header
    st.markdown("""
    <div class="app-header">
        <div class="app-header-mark">✦</div>
        <div>
            <div class="app-header-title">Individual Behavior Support Plan Generator</div>
            <div class="app-header-sub">Fill in the sidebar → preview updates live → export Staff PDF or Parent PDF</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Collect form data from sidebar
    d = render_sidebar()

    # -- Live preview ----------------------------------------------------------
    if d["child_name"].strip():
        # Pre-warm AI generation with a spinner (cached after first call)
        _sk_t = tuple(d.get("skills", []) + ([d["custom_skill"]] if d.get("custom_skill") else []))
        _bh_t = tuple((d.get("behaviors") or []) + ([d["custom_behavior"]] if d.get("custom_behavior") else []))
        _cn   = d.get("child_name", "")
        if _sk_t or _bh_t:
            with st.spinner("✦ Tailoring activities and progress indicators to selected skills and behaviors…"):
                generate_dynamic_activities(_sk_t, _bh_t, _cn)
                generate_dynamic_progress(_sk_t, _bh_t, _cn)

        st.markdown(render_html_preview(d), unsafe_allow_html=True)

        # Export section
        st.markdown("""
        <div class="export-box">
            <div class="export-title">Export Documents</div>
        </div>
        """, unsafe_allow_html=True)

        col_pdf, col_parent = st.columns(2)

        with col_pdf:
            try:
                pdf_bytes = build_pdf_staff(d)
                filename_pdf = f"BehaviorPlan_Staff_{d['child_name'].replace(' ','_')}.pdf"
                st.download_button(
                    label="⬇  Staff PDF (+ Activities)",
                    data=pdf_bytes,
                    file_name=filename_pdf,
                    mime="application/pdf",
                    use_container_width=True,
                    help="Full professional report with suggested behavioral activities for sessions",
                )
            except Exception as e:
                st.error(f"Staff PDF error: {e}")

        with col_parent:
            try:
                parent_bytes = build_pdf_parent(d)
                filename_parent = f"BehaviorPlan_Parent_{d['child_name'].replace(' ','_')}.pdf"
                st.download_button(
                    label="⬇  Parent-Friendly PDF",
                    data=parent_bytes,
                    file_name=filename_parent,
                    mime="application/pdf",
                    use_container_width=True,
                    help="Warm, clear summary written for parents — no jargon",
                )
            except Exception as e:
                st.error(f"Parent PDF error: {e}")

        st.markdown("<div style='margin-top:0.6rem;'></div>", unsafe_allow_html=True)
        try:
            pack_bytes = build_activity_pack(d)
            filename_pack = f"ActivitySessionPack_{d['child_name'].replace(' ','_')}.pdf"
            st.download_button(
                label="⬇  Activity Session Pack (PDF)",
                data=pack_bytes,
                file_name=filename_pack,
                mime="application/pdf",
                use_container_width=True,
                help="Full activity session pack — conduct guides, tracking sheets, parent sheets, token board",
            )
        except Exception as e:
            st.error(f"Activity Pack error: {e}")

    else:
        st.markdown("""
        <div style="display:flex;flex-direction:column;align-items:center;justify-content:center;
                    height:55vh;gap:1rem;opacity:0.5;">
            <div style="font-size:3rem;">✦</div>
            <div style="font-family:'Cormorant Garamond',serif;font-size:1.6rem;font-weight:400;
                        color:#2A2520;">Your plan will appear here</div>
            <div style="font-size:0.85rem;color:#6B6456;">Enter a child name in the sidebar to begin</div>
        </div>
        """, unsafe_allow_html=True)

def build_activity_pack(d) -> bytes:
    import math, tempfile, os
    from io import BytesIO
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak
    )
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.colors import HexColor
    from pypdf import PdfWriter, PdfReader

    C_TEAL         = HexColor("#2E7D8C")
    C_TEAL_LIGHT   = HexColor("#E8F4F6")
    C_GOLD         = HexColor("#D4A017")
    C_GOLD_LIGHT   = HexColor("#FDF6E3")
    C_CORAL        = HexColor("#E05C5C")
    C_CORAL_LIGHT  = HexColor("#FDF0F0")
    C_GREEN        = HexColor("#3A7D44")
    C_GREEN_LIGHT  = HexColor("#EAF4EB")
    C_PURPLE       = HexColor("#6A4C93")
    C_PURPLE_LIGHT = HexColor("#F2EEF8")
    C_ORANGE       = HexColor("#C0622A")
    C_ORANGE_LIGHT = HexColor("#FDF0E8")
    C_BLUE         = HexColor("#2A5C8C")
    C_BLUE_LIGHT   = HexColor("#E8EFF8")
    C_DARK         = HexColor("#1C2B3A")
    C_MID_GREY     = HexColor("#6B7C93")
    C_LIGHT_GREY   = HexColor("#F5F7FA")
    C_WHITE        = colors.white
    C_BORDER       = HexColor("#D0D8E4")

    child       = d.get("child_name") or "Child"
    therapist   = d.get("therapist_name") or "Therapist"
    nursery     = d.get("nursery_name") or "Nursery"
    class_group = d.get("class_group") or "—"
    age         = d.get("age") or "—"
    start_str   = d["start_date"].strftime("%d %B %Y") if d.get("start_date") else "—"
    freq_label  = d.get("session_frequency") or "—"
    sessions    = int(d.get("sessions_per_month") or 4)
    skills      = d.get("skills") or []
    all_skills  = skills + ([d["custom_skill"]] if d.get("custom_skill") else [])
    behaviors   = (d.get("behaviors") or []) + ([d["custom_behavior"]] if d.get("custom_behavior") else [])
    footer_t    = d.get("footer_text") or "Confidential — For therapist and family use only."
    today_str   = datetime.now().strftime("%d %B %Y")

    ACTIVITY_LIBRARY = {
      "Emotional Regulation": [{
        "name": "Feelings Monster Bag",
        "skill_label": "Emotional Regulation",
        "color": C_TEAL, "light": C_TEAL_LIGHT,
        "objective": "{child} will learn to identify, name, and physically express core emotions (happy, sad, angry, scared) using visual cards and puppet prompts. By the end of the session series, {child} will spontaneously label an emotion when asked or when naturally occurring during the session.",
        "materials": [
          ("Primary Tools","Soft bag or box, emotion picture cards (x4-6), small hand puppets (optional), mirror"),
          ("Substitutes","Plastic bag + laminated drawings; hand-drawn cards; small toys as characters; camera app as mirror"),
          ("Reinforcers","Token board + preferred stickers or stamps identified with the family"),
        ],
        "steps": [
          ("Set up (before session)","Place the bag on the table with emotion cards inside. Have token board visible. Prepare mirror if available."),
          ("Introduce the bag (2 min)","Say: 'I have a special feelings bag today! There are some characters hiding inside.' Build anticipation before opening."),
          ("Therapist models first (3 min)","Pull out 'happy'. Make a big happy face. Say 'This one is HAPPY — look at my face!' Invite {child} to copy. Award token with labeled praise."),
          ("Child's turn (8 min)","{child} pulls cards one at a time. For each: ask 'What feeling is this?' model if needed, ask them to make the face, ask 'When do you feel ___?' Accept pointing, nodding, or one word."),
          ("What helps? (4 min)","For 'angry' and 'sad', show a visual with 2-3 coping options. Ask 'What can you do?' Coach toward the visual. This is replacement behaviour rehearsal."),
          ("Wrap-up (3 min)","Recap: 'Today we found feelings! You were amazing.' Count tokens. Let {child} pack cards back."),
        ],
        "adapt": [
          ("Refuses to engage","Therapist plays alone and narrates. Curiosity usually draws child in within 2-3 minutes."),
          ("Becomes dysregulated","Pause. Offer a brief movement break. Return to the bag calmly."),
          ("Cannot name emotions","Switch to matching: 'Which face looks like THIS face?' rather than open naming."),
          ("Short attention span","Use only 2-3 cards per session. Add novelty with a puppet voice."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Emotion cards identified correctly (# / total)","_______  out of  _______"),
          ("Emotion labelling modality","&#9744; Verbal  &#9744; Pointed  &#9744; Nodded  &#9744; No response"),
          ("Face mirroring quality","&#9744; Spontaneous  &#9744; Prompted  &#9744; Refused"),
          ("Response to 'What helps?' prompt","&#9744; Verbal  &#9744; Pointed to visual  &#9744; No response"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Intensity of escalation (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5  &#9744; N/A"),
          ("Recovery time","_______  min  &#9744; N/A"),
          ("Replacement behavior used?","&#9744; Spontaneous  &#9744; Prompted  &#9744; None"),
          ("Tokens earned","_______"),
          ("Prompt level","&#9744; Independent  &#9744; Gestural  &#9744; Verbal  &#9744; Full physical"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} played a feelings game using our special Feelings Bag! We explored emotion cards and practised making different faces. We also began talking about what helps when we feel a big feeling.",
        "parent_why": "When children can name their feelings, they are far less likely to express them through challenging behaviour. Naming is the very first step toward managing emotions and is foundational to everything else in this programme.",
        "parent_tips": [
          "When {child} shows a strong emotion, gently name it: 'It looks like you are feeling angry right now.' You do not need to fix it - just name it.",
          "Make it a game: pull funny faces together and ask 'What feeling is this?'",
          "If {child} uses any feeling word, even just 'mad', celebrate it immediately.",
          "Simple emotion cards on the fridge work well for continued practice at home.",
        ],
        "substitutes": [
          "Bag: any box, envelope, or pillowcase",
          "Emotion cards: hand-drawn faces on paper, printed A4 sheet, magazine cut-outs",
          "Mirror: phone camera app, or skip entirely",
          "Puppets: any small toys used as 'characters'",
        ],
      }],
      "Impulse Control": [{
        "name": "Red Light, Green Light",
        "skill_label": "Impulse Control (Inhibitory Control)",
        "color": C_GOLD, "light": C_GOLD_LIGHT,
        "objective": "{child} will develop the neurological pause between impulse and action by practising deliberate stopping of a preferred movement in response to a clear visual signal. Progress is measured by reduction in impulse errors and observable increase in latency between signal and response.",
        "materials": [
          ("Primary Tools","Large red card, large green card, yellow card (advanced rounds), token board"),
          ("Substitutes","Red and green paper, plates, blocks, or hats - any two strongly contrasting objects; yellow introduced only from session 3+"),
          ("Preferred action","Ball rolling, marching, clapping, jumping - identify with child at session start"),
          ("Reinforcers","Token board + preferred stickers or stamps"),
        ],
        "steps": [
          ("Setup and motivation check (2 min)","Show the two cards. Ask what {child} wants to do on 'green' - offer 2 choices. This creates buy-in and ensures movement is intrinsically motivating."),
          ("Therapist models (2 min)","Therapist goes first: show green - start jumping; show red - FREEZE. Make it fun and exaggerated. Let {child} be card-holder for one round."),
          ("Child plays, slow pace (5 min)","Switch signals slowly every 5-7 seconds. Give warm-up successes early. Award token every successful stop. Narrate: 'RED - and you STOPPED! Wow!'"),
          ("Increase pace (8 min)","Gradually reduce green intervals. Introduce unexpected red signals. Advanced: introduce yellow (slow down). Track impulse errors neutrally."),
          ("Child as card-holder (3 min)","Give {child} the cards; therapist is the mover. Builds turn-taking and consolidates rule understanding."),
          ("Wrap-up (2 min)","Count tokens. Review: 'You stopped your body many times today!'"),
        ],
        "adapt": [
          ("Will not stop at all","Reduce green interval drastically to create early success."),
          ("Becomes frustrated","Offer card-holder role. Redirect frustration into leadership within the activity."),
          ("Loses interest","Change movement: silly walking, animal movements, or drumming on the table."),
          ("Limited space","Seated version: hands on knees = stop; hands clapping = go."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Movement chosen this session","_____________________________"),
          ("Total red signals given","_______"),
          ("Successful stops (# / total)","_______  out of  _______"),
          ("Latency to stop","&#9744; Immediate  &#9744; 1-2 sec delay  &#9744; 3+ sec  &#9744; Did not stop"),
          ("Impulse errors (moved on red)","_______  errors"),
          ("Yellow card introduced?","&#9744; Yes  &#9744; No - not yet"),
          ("Response to correction","&#9744; Accepted calmly  &#9744; Resisted  &#9744; Escalated"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Intensity of escalation (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5  &#9744; N/A"),
          ("Recovery time","_______  min  &#9744; N/A"),
          ("Tokens earned","_______"),
          ("Prompt level","&#9744; Independent  &#9744; Gestural  &#9744; Verbal  &#9744; Full physical"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} played a movement game where they had to move on the green card and freeze completely on the red card. A classic game with a very serious developmental purpose!",
        "parent_why": "The ability to stop your body on purpose is one of the earliest and most important forms of self-control. This game trains the exact part of the brain that helps children pause before reacting impulsively.",
        "parent_tips": [
          "Play a version at home during any routine: Freeze! while walking to the table, getting dressed, or in the garden.",
          "Use 'red light' as a calm, playful signal when you need {child} to pause - not a command, but a shared game cue.",
          "When they stop, name what they did: 'You stopped your body! That was such good brain work.'",
          "Even one successful freeze deserves a big celebration. The pause is the goal.",
        ],
        "substitutes": [
          "Red/green cards: coloured plates, cups, hats, blocks, or scarves",
          "Open space: seated hand version (hands up/down) if space is limited",
          "Yellow card: simply say 'slow down' as a verbal cue instead",
        ],
      }],
      "Frustration Tolerance": [{
        "name": "The Waiting Tower",
        "skill_label": "Frustration Tolerance",
        "color": C_GREEN, "light": C_GREEN_LIGHT,
        "objective": "{child} will tolerate brief, structured waiting periods during a preferred activity without escalating to challenging behaviour. Waiting duration is systematically increased across sessions, and verbal expression of the wait is explicitly trained as the replacement behaviour.",
        "materials": [
          ("Primary Tools","Building blocks (DUPLO or wooden), sand timer or visual countdown, token board"),
          ("Substitutes","Any stackable items (cups, boxes, soft cubes); count aloud or clap rhythm as timer; a drawn number line to cross off"),
          ("No blocks?","Puzzle, threading beads, or any structured turn-taking craft - the waiting mechanic is the core"),
          ("Reinforcers","Token board + 2-minute free play reward at tower completion"),
        ],
        "steps": [
          ("Setup (before session)","Set up blocks and place sand timer visibly. Decide starting wait duration from previous session data (5 seconds if first session)."),
          ("Introduce the game (2 min)","'We are building a tower together! But there is a rule - we take turns. When the sand runs out, it is your turn.' Let {child} flip the timer once."),
          ("Therapist goes first (2 min)","Therapist places first block. Flips timer. Narrates the wait. First few turns should be very short to build confidence."),
          ("Structured turn-taking (12 min)","Continue building, alternating. Gradually extend timer. Every successful wait = token + specific praise. If {child} says 'my turn soon' spontaneously - bonus token."),
          ("Tower completion reward (3 min)","'You built it! You waited so many times!' Let {child} knock it over. Offer 2 min free block play as reward."),
          ("Wrap-up (3 min)","Record maximum wait achieved. Compare to last session. Share progress with {child} visually."),
        ],
        "adapt": [
          ("Cannot wait at all","Start at 3 seconds. Celebrate even partial waits. Build up one second at a time."),
          ("Escalates during wait","Shorten the wait immediately. Success is more important than duration."),
          ("Not interested in blocks","Switch to any preferred turn-taking activity."),
          ("No timer available","Count together aloud: one, two, three, your turn! Use clapping rhythm."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Target wait time this session","_______  seconds"),
          ("Total waiting turns given","_______"),
          ("Successful waits (# / total)","_______  out of  _______"),
          ("Maximum wait achieved","_______  seconds"),
          ("Behaviour during wait","&#9744; Calm  &#9744; Restless  &#9744; Verbal protest  &#9744; Physical escalation"),
          ("Verbal expression during wait","&#9744; Spontaneous ('my turn soon')  &#9744; Prompted  &#9744; None"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Intensity of escalation (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5  &#9744; N/A"),
          ("Recovery time","_______  min  &#9744; N/A"),
          ("Tokens earned","_______"),
          ("Previous session max wait / this session max","_______  sec  /  _______  sec"),
          ("Target for next session","_______  seconds"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} and I built a block tower together with a special rule: we had to take turns, and each person had to wait while the other placed their block. We used a small timer so {child} could see the wait.",
        "parent_why": "Many challenging moments happen when something {child} wants is briefly out of reach or delayed. Practising tiny, safe waits teaches the brain that waiting is survivable and always leads to something good.",
        "parent_tips": [
          "Use a small timer during moments that usually cause friction - waiting for food, a toy, or a turn.",
          "When {child} waits, name it: 'You waited! That was so grown-up.'",
          "Keep home waits very short at first - 5 to 10 seconds. Success matters far more than duration.",
          "If {child} starts to grab or protest, calmly say 'not yet - let us count together.'",
          "Never extend a wait as a punishment. The timer is a neutral, fair signal.",
        ],
        "substitutes": [
          "Blocks: stacking cups, cardboard boxes, soft toys, books",
          "Sand timer: count aloud, clap a rhythm, cross numbers off a drawn line",
          "No table: floor version works equally well",
        ],
      }],
      "Communication Skills": [{
        "name": "The 'I Need' Communication Board",
        "skill_label": "Functional Communication",
        "color": C_PURPLE, "light": C_PURPLE_LIGHT,
        "objective": "{child} will learn to use a simple picture communication board to express needs, preferences, and distress as a direct replacement for identified challenging behaviours. By the end of the session series, {child} will spontaneously point to or vocalise a card in the natural environment.",
        "materials": [
          ("Primary Tools","Printed and laminated picture board (6 cards: I need a break / Help me / No thank you / I am angry / I want / Stop please)"),
          ("Substitutes","Hand-drawn symbols on card; sticky notes with one word each; a single STOP card if simpler; any object {child} can point to"),
          ("Role-play props","Favourite toy to place out of reach; a transition bell; a box of items for practice"),
          ("Reinforcers","Token board + immediate response when card is used (CRITICAL: honours the request - this IS the reinforcer)"),
        ],
        "steps": [
          ("Introduce the board (3 min)","Lay board out. Point to each card and say its name with an exaggerated expression. Let {child} touch and point freely. Do not quiz yet."),
          ("Model the game (3 min)","Set up a scenario: favourite toy just out of reach. Therapist models pointing to 'I want'. Narrate replacement behaviour explicitly."),
          ("Role-play scenarios (10 min)","Run 4-6 brief scenarios mirroring real triggers. (1) Toy taken: I want / No thank you. (2) Tidy-up: No thank you / Stop. (3) Too hard: Help me / I need a break. (4) Frustrated: I am angry. Accept any pointing or vocalisation."),
          ("Generalisation test (4 min)","Without narrating, create a natural frustrating moment. See if {child} reaches for the board without prompting. Wait 5-10 seconds. Record result."),
          ("Board placement decision (2 min)","Decide with {child} where the board will live. Involvement increases ownership and use."),
          ("Wrap-up (2 min)","Review: 'Today you used your words instead of [behaviour]. That is such important growing up.'"),
        ],
        "adapt": [
          ("Ignores the board","Reduce to 2 cards only. Less choice = less cognitive load."),
          ("Cannot point reliably","Use hand-over-hand assistance. Gradually fade physical support."),
          ("Escalates in role-play","The scenario was too close to a real trigger. Scale back."),
          ("No laminator","Use a plastic sleeve, contact paper, or a simple drawn A4 sheet in a clear pocket."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Cards used this session","_____________________________"),
          ("Scenarios presented (# total)","_______"),
          ("Correct card used (# / total)","_______  out of  _______"),
          ("Communication modality","&#9744; Verbal  &#9744; Pointed  &#9744; Gestured  &#9744; No response"),
          ("Prompt level per scenario","&#9744; Independent  &#9744; Modelled  &#9744; Hand-over-hand"),
          ("Spontaneous generalisation observed?","&#9744; Yes - describe in notes  &#9744; No"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Intensity of escalation (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5  &#9744; N/A"),
          ("Recovery time","_______  min  &#9744; N/A"),
          ("Tokens earned","_______"),
          ("Strongest replacement behaviour","_____________________________"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} practised using a special picture board to tell us what they need. We acted out little pretend situations and {child} practised pointing to or saying a card instead of reacting with a tantrum or throwing.",
        "parent_why": "Most challenging behaviours are {child}'s way of communicating something they cannot yet say in words. This activity gives them a clearer, faster tool so they do not have to use their body to get their message across.",
        "parent_tips": [
          "Keep a copy of the board somewhere visible at home - fridge or play area works well.",
          "When {child} starts to escalate, calmly point to the board: 'Can you show me which one?' - before the peak, not after.",
          "If {child} points to or touches a card, respond immediately and specifically.",
          "Never ignore a board use, even a partial one. Even looking at the board deserves acknowledgement.",
          "If they use a word instead of the board - celebrate it equally. Any functional communication counts.",
        ],
        "substitutes": [
          "Printed board: hand-drawn symbols on A4; sticky notes with one word; a single STOP sign",
          "Laminator: plastic sleeve, contact paper, or zip-lock bag",
          "Role-play props: any favourite toy to place out of reach",
        ],
      }],
      "Social Skills": [{
        "name": "Compliment Catch and Social Stories",
        "skill_label": "Social Skills and Peer Interaction",
        "color": C_CORAL, "light": C_CORAL_LIGHT,
        "objective": "{child} will develop age-appropriate turn-taking, cooperative play, and positive social interaction skills. Spontaneous positive social initiations toward peers are the target outcome.",
        "materials": [
          ("Primary Tools","Soft ball, printed social story (1 page illustrated), token board"),
          ("Substitutes","Any soft object to pass (beanbag, stuffed toy, rolled socks); draw the social story together"),
          ("Peer involvement","A trusted classmate or nursery adult acts as peer"),
          ("Reinforcers","Token board + paired activity with the peer as completion reward"),
        ],
        "steps": [
          ("Social story (5 min)","Read the social story together: a character in a situation matching {child}'s real triggers. Ask: 'What did the character do? Was that a good idea?' Keep it light and playful."),
          ("Model the game (2 min)","Therapist models compliment catch: catch the ball, say one genuine kind thing, throw. Practice once with just therapist and {child}."),
          ("Play - therapist and child (5 min)","Play 3-4 rounds. Prompt gently if {child} is stuck: 'I notice something nice about you...' Accept any positive statement, however small."),
          ("Introduce peer (8 min)","Bring in a trusted peer or adult. Run 3-4 rounds with all three. Debrief: 'How did it feel to give and receive a compliment?'"),
          ("Wrap-up (3 min)","Identify one real moment this week where {child} could use the skill. Role-play briefly. Count tokens."),
        ],
        "adapt": [
          ("Refuses peer interaction","Stay therapist-only. Introduce a puppet as the 'peer' to bridge toward real interaction."),
          ("Cannot think of a compliment","Offer a menu: 'You could say: nice shirt / good throw / you are funny.'"),
          ("Social story is too long","Use just one illustration and one sentence. Simplicity is key."),
          ("Ball causes distraction","Use a beanbag or pass a card instead."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Social story comprehension","&#9744; Understood and discussed  &#9744; Needed support  &#9744; Disengaged"),
          ("Compliments given (# rounds)","_______"),
          ("Quality of compliments","&#9744; Spontaneous and genuine  &#9744; Prompted  &#9744; Imitated from menu"),
          ("Peer present this session?","&#9744; Yes  &#9744; No - therapist only"),
          ("Social initiation observed?","&#9744; Yes - describe in notes  &#9744; No"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Recovery time","_______  min  &#9744; N/A"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} practised social skills through a compliment catch game and a short social story. We worked on saying kind things, taking turns, and understanding what good interactions look like.",
        "parent_why": "Social skills are rarely innate - they need to be explicitly taught and practised. By rehearsing them in a safe, structured setting, {child} builds a repertoire of positive social behaviours for use with peers.",
        "parent_tips": [
          "Point out positive social moments at home: 'I noticed you shared that - that was really kind.'",
          "Play simple turn-taking games at home: board games, card games, taking turns in conversation.",
          "If {child} struggles with a peer situation, ask: 'What could you say instead?'",
          "Model compliments yourself - children learn social language by hearing it.",
        ],
        "substitutes": [
          "Ball: beanbag, stuffed toy, rolled socks",
          "Printed social story: draw it together with {child} in the session",
          "Peer: a trusted nursery adult can fulfil the peer role initially",
        ],
      }],
      "Self-Awareness": [{
        "name": "Body Check-In and Emotion Map",
        "skill_label": "Self-Awareness and Emotional Literacy",
        "color": C_ORANGE, "light": C_ORANGE_LIGHT,
        "objective": "{child} will build increasing awareness of their own emotional states and bodily sensations associated with emotion. By the end of the session series, {child} will spontaneously identify where they feel an emotion in their body and connect it to a named feeling state.",
        "materials": [
          ("Primary Tools","Large body outline drawing (A3), colouring markers (4+ colours), feelings chart"),
          ("Substitutes","Draw a body outline together on any large paper; use crayon or pencil; a simple smiley scale if no feelings chart"),
          ("Reinforcers","Token board + child keeps the body map they create"),
        ],
        "steps": [
          ("Body scan opening (3 min)","'Let us notice what your body feels today. What do your hands feel like? Your tummy?' Keep it light and curious, not clinical."),
          ("Introduce the body map (3 min)","Show a blank body outline. 'This is our feelings map. We are going to colour where feelings live in your body.' Model with one emotion first."),
          ("Child colours their map (10 min)","Guide {child} through 3-4 emotions. For each: name it, ask 'Where do you feel this?' Let {child} choose the colour and mark the spot. Accept any answer."),
          ("Connect to real moments (5 min)","Point to one marked area: 'Last week when [behaviour occurred], was your tummy feeling this?' Build the bridge gently."),
          ("Closing body scan (2 min)","Repeat the opening scan. Ask 'Does anything feel different now?'"),
        ],
        "adapt": [
          ("Will not engage with body map","Do a feelings puppet scan instead - the puppet has the feelings, not the child."),
          ("Cannot connect body to emotion","Use temperature language: 'Does it feel hot or cold inside when you are angry?'"),
          ("Becomes overwhelmed","Stay with only one emotion per session."),
          ("No large paper","Any paper works - even a simple outline drawn on A4."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Emotions mapped this session","_____________________________"),
          ("Body-emotion connection made?","&#9744; Yes - described bodily sensation  &#9744; Partial  &#9744; Not yet"),
          ("Verbal quality","&#9744; Spontaneous description  &#9744; Responded to prompts  &#9744; Pointed only"),
          ("Real-moment connection made?","&#9744; Yes  &#9744; Not yet"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Recovery time","_______  min  &#9744; N/A"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} worked on a Feelings Body Map - a large drawing where we coloured in where different emotions live in the body. We talked about what feelings feel like physically.",
        "parent_why": "Children who can notice body sensations have an early warning system for dysregulation. When {child} learns to notice 'my tummy feels tight' before a meltdown, they gain precious seconds to use a coping strategy.",
        "parent_tips": [
          "Ask 'What does your body feel like right now?' during calm moments - not just during hard ones.",
          "Use physical language yourself: 'My shoulders feel tight today - I think I am a bit stressed.'",
          "When {child} is escalating, gently ask: 'What does your tummy feel like right now?'",
          "Keep the body map somewhere visible - children often refer back to it spontaneously.",
        ],
        "substitutes": [
          "A3 body outline: trace around child's hand or arm on any paper",
          "Markers: crayons, pencils, or stickers to mark body zones",
          "Feelings chart: a simple hand-drawn smiley/sad/angry face scale",
        ],
      }],
      "Coping Strategies": [{
        "name": "My Calm Corner Toolkit",
        "skill_label": "Coping Strategy Development",
        "color": C_BLUE, "light": C_BLUE_LIGHT,
        "objective": "{child} will be introduced to and will practise a personalised set of coping strategies appropriate to their developmental level. The goal is for {child} to begin accessing these tools with decreasing levels of adult support over time.",
        "materials": [
          ("Primary Tools","Small box or bag, simple breathing exercise card, one sensory item, a picture of something calming, a feelings chart"),
          ("Substitutes","Any container; draw the breathing card together; sensory item can be textured fabric, smooth stone, or stress ball made from balloon and flour"),
          ("Reinforcers","Token board + child decorates and keeps their toolkit box"),
        ],
        "steps": [
          ("Introduce the idea (3 min)","'Everyone needs a toolkit for hard moments - even adults. Today we are building yours.' Show an example box. Let {child} explore the items."),
          ("Breathing practice (5 min)","Teach one simple breathing technique: slow belly breaths, blowing pretend bubbles, or smell the flowers / blow the candles. Practise together 3-4 times."),
          ("Sensory item selection (3 min)","Offer 2-3 sensory options. Let {child} choose one for their box. Ask 'What does this feel like? Does it feel calming?'"),
          ("Calming picture (3 min)","Draw or choose a picture of something calming to {child}. Attach to the box. 'When things feel hard, you can look at this.'"),
          ("Practice run (8 min)","Set up a mild frustrating scenario. Coach {child} to open the toolkit and choose a tool. Celebrate any tool use, however brief."),
          ("Wrap-up (3 min)","Decide where the toolkit will live in the classroom. 'It will always be here when you need it.'"),
        ],
        "adapt": [
          ("Refuses breathing","Do not force it. Offer the sensory item only. Introduce breathing casually next session."),
          ("Becomes dysregulated during practice","Use the toolkit for real right now. This is the best possible teaching moment."),
          ("Wants to play with items, not use them","That is fine initially. Familiarity with the tools is the first step."),
          ("No box available","Use a corner of the room, a special chair, or a drawer."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Breathing technique practised","_____________________________"),
          ("Breathing quality","&#9744; Slow and controlled  &#9744; Fast  &#9744; Variable  &#9744; Refused"),
          ("Sensory item selected","_____________________________"),
          ("Toolkit used during real distress?","&#9744; Yes - describe in notes  &#9744; No"),
          ("Tool initiated spontaneously?","&#9744; Yes  &#9744; Prompted  &#9744; Not yet"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Intensity of escalation (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5  &#9744; N/A"),
          ("Recovery time","_______  min  &#9744; N/A"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} built their very own Calm Corner Toolkit - a personal set of strategies to use when things feel hard. We practised slow breathing, chose a calming sensory item, and drew a picture of something that makes {child} feel safe.",
        "parent_why": "Children cannot use coping skills they have never practised. By building and rehearsing these tools in calm, playful moments, the brain starts to remember them when things get hard.",
        "parent_tips": [
          "Make a matching toolkit at home - {child} can help choose what goes in it.",
          "Practise the breathing technique during calm moments, not just hard ones.",
          "If {child} starts to escalate, offer the toolkit calmly: 'Do you want to use your toolkit?'",
          "Never force the tools during a meltdown. Introduce them just before or just after.",
          "Celebrate every time {child} uses a tool, even if it only helps a little.",
        ],
        "substitutes": [
          "Box: corner of the room, special chair, or any container",
          "Sensory item: smooth stone, textured fabric, squeezable sponge, homemade stress ball",
          "Breathing card: draw together; use hands as a guide (trace fingers = one breath)",
        ],
      }],
      "Attention and Concentration": [{
        "name": "Listening Detectives",
        "skill_label": "Attention and Concentration",
        "color": C_TEAL, "light": C_TEAL_LIGHT,
        "objective": "{child} will develop sustained attention by practicing focused listening tasks with gradually increasing duration and complexity. By the end of the session series, {child} will maintain on-task engagement for progressively longer periods with decreasing adult prompting.",
        "materials": [
          ("Primary Tools","Simple sound recordings or instruments, picture cards, token board"),
          ("Substitutes","Clap rhythms instead of recordings; any visual cards; tally marks for tracking"),
          ("Reinforcers","Token board + short preferred activity as end reward"),
        ],
        "steps": [
          ("Set up (before session)","Prepare 4-6 sounds or rhythm patterns. Place token board visibly. Decide on target on-task duration from previous data."),
          ("Introduce the Detective (2 min)","'Today you are a Listening Detective! Your job is to hear something very carefully and tell me what it is.' Show a detective badge if available."),
          ("Therapist models first (2 min)","Play a sound/rhythm, close eyes, listen, then describe it. Model the full listening behavior before {child}'s turn."),
          ("Detective rounds (10 min)","Take turns listening and identifying. Each correct response or genuine attempt earns a token. Gradually increase sound complexity or attention duration required."),
          ("Focus challenge (5 min)","Introduce a brief distraction (a knock, a movement) while {child} is listening. Coach them to ignore it and keep attending. This is the key generalization target."),
          ("Wrap-up (3 min)","Count tokens. 'You listened so carefully today — your brain worked really hard!'"),
        ],
        "adapt": [
          ("Cannot sustain 30 seconds","Start at 10 seconds. Build up 5 seconds per success."),
          ("Gets distracted easily","Reduce environmental noise first. Use a visual focus cue: point to ears."),
          ("Not interested in sounds","Switch to a visual attention task: 'Spot the difference' or sorting."),
          ("Impulsive responses","Use a counting delay: 'Listen first, then tell me after I count to 3.'"),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Target on-task duration this session","_______  seconds"),
          ("Maximum sustained attention achieved","_______  seconds"),
          ("Distractor introduced?","&#9744; Yes  &#9744; Not yet"),
          ("Response to distractor","&#9744; Maintained focus  &#9744; Briefly distracted, recovered  &#9744; Could not refocus"),
          ("Prompt level needed","&#9744; Independent  &#9744; Gestural  &#9744; Verbal  &#9744; Full physical"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} played a Listening Detectives game, focusing carefully on sounds and rhythms to identify them. We also practiced staying focused when a small distraction happened nearby.",
        "parent_why": "The ability to sustain focus is one of the most important skills for learning and social participation. Practicing it in a playful, low-stakes game is the safest way to build it.",
        "parent_tips": [
          "At home, try a 'listening moment': sit quietly for 30 seconds and name all the sounds you can hear together.",
          "When {child} is focused on something they enjoy, gently narrate it: 'Look how well you are concentrating!'",
          "Reduce background noise (TV, music) during tasks that require {child}'s attention.",
          "Short focused activities followed by movement breaks work better than long sit-down sessions.",
        ],
        "substitutes": [
          "Sound recordings: clap rhythms, tap patterns on a table, or hum melodies",
          "Picture cards: any images for sorting or identifying",
          "No badge: draw a detective hat on paper or use a sticker",
        ],
      }],
      "Decision Making": [{
        "name": "The Choice Station",
        "skill_label": "Decision Making and Problem Solving",
        "color": C_GOLD, "light": C_GOLD_LIGHT,
        "objective": "{child} will develop confidence in making choices, considering options, and anticipating simple consequences through structured choice-making games and guided reflection. By the end of the session series, {child} will demonstrate increasing independence in everyday decision-making situations.",
        "materials": [
          ("Primary Tools","Choice cards or visual choice board (2-3 options), simple scenario cards, token board"),
          ("Substitutes","Drawn pictures for choices; describe scenarios verbally; any objects representing options"),
          ("Reinforcers","Token board + child chooses the final reward activity"),
        ],
        "steps": [
          ("Set up (before session)","Prepare 3-4 simple scenario cards (snack choice, activity choice, toy choice). Have choice board ready."),
          ("Introduce the Choice Station (2 min)","'Today we are going to practice making choices — which is something really important grown-ups do every day!'"),
          ("Easy warm-up choices (5 min)","Start with preferred options only: 'Do you want to draw or play with blocks first?' Make the decision feel positive and powerful."),
          ("Scenario practice (10 min)","Present scenario cards. Ask {child} to pick an option. Then ask: 'What do you think will happen if you choose that?' Guide without correcting — explore consequences together."),
          ("Tough choice challenge (5 min)","Introduce a slightly harder scenario (e.g., sharing a toy vs keeping it). Guide using: 'Stop — Think — Choose — Check.' Reinforce any genuine reflection."),
          ("Wrap-up (3 min)","'Today you made so many great choices! Making good choices takes practice — and you practiced a lot today.'"),
        ],
        "adapt": [
          ("Freezes when asked to choose","Reduce to two very simple options. Use objects rather than cards."),
          ("Always chooses impulsively","Introduce a mandatory 3-second wait before any choice is accepted."),
          ("Becomes upset when wrong","Remove any framing of 'right/wrong'. All choices lead to learning."),
          ("Advanced learner","Introduce 3-option scenarios and add a consequence-reflection step."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Choices made this session","_______"),
          ("Decision style","&#9744; Impulsive  &#9744; Considered  &#9744; Hesitant/Frozen  &#9744; Variable"),
          ("Consequence awareness demonstrated?","&#9744; Yes - described outcome  &#9744; Partial  &#9744; Not yet"),
          ("Stop-Think-Choose-Check used?","&#9744; Spontaneous  &#9744; Prompted  &#9744; Not yet"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} practiced making choices using our Choice Station game. We worked on slowing down before choosing, and began thinking about what might happen after different choices.",
        "parent_why": "Children who feel confident making decisions are more independent, more resilient when things go wrong, and better equipped to handle peer pressure and social challenges.",
        "parent_tips": [
          "Give {child} real choices throughout the day: 'Do you want the red cup or the blue one?' Both options should be acceptable.",
          "After a choice, gently reflect: 'How did that work out? Would you choose the same next time?'",
          "Avoid choosing for {child} when they can manage it. Even a slow or 'wrong' choice is valuable practice.",
          "When {child} makes a good decision, name it specifically: 'I noticed you thought about that before you decided — brilliant.'",
        ],
        "substitutes": [
          "Choice cards: objects, pictures from a magazine, or hand-drawn options",
          "Scenario cards: describe verbally or use simple puppets to act out situations",
        ],
      }],
      "Self-Boundaries": [{
        "name": "My Body, My Space",
        "skill_label": "Self-Boundaries and Personal Space",
        "color": C_CORAL, "light": C_CORAL_LIGHT,
        "objective": "{child} will develop awareness of personal physical space, learn to identify comfortable and uncomfortable interactions, and practice using assertive, appropriate language to communicate boundaries. By the end of the series, {child} will demonstrate increasing confidence in recognising and expressing personal limits.",
        "materials": [
          ("Primary Tools","Hula hoop or chalk circle, 'comfortable / uncomfortable' picture cards, token board"),
          ("Substitutes","Draw a circle on paper to represent personal space; use thumbs up/down for comfortable/uncomfortable; any stuffed animal for role-play"),
          ("Reinforcers","Token board + child's choice of closing activity"),
        ],
        "steps": [
          ("Set up (before session)","Place hula hoop on floor. Prepare picture cards. Have token board ready."),
          ("My Space Circle (3 min)","Stand inside the hoop: 'This is your personal space — your bubble. Everyone has one.' Let {child} decorate or name their bubble."),
          ("Comfortable / Uncomfortable (5 min)","Show picture cards of different touch scenarios. Ask: 'Does this feel OK or not OK?' Validate all answers. There are no wrong feelings."),
          ("Role-play practice (8 min)","Use puppets or soft toys: Puppet A asks for a hug. Coach {child} to respond: 'Yes please / No thank you.' Practice both accepting and declining. Reinforce both as equally valid."),
          ("Saying 'Stop' practice (5 min)","Role-play a boundary being crossed. Coach {child} to say 'Stop, I don't like that' clearly. Practice volume and body posture. Award token for assertive (not aggressive) responses."),
          ("Wrap-up (3 min)","'Your body belongs to you. You are always allowed to say no. And we always listen.'"),
        ],
        "adapt": [
          ("Uncomfortable with topic","Keep all scenarios with puppets/toys only. Never use real touch demonstrations."),
          ("Already assertive but aggressive","Redirect toward assertive voice and words, distinguishing from shouting."),
          ("Shy or won't respond","Use a scale: thumbs up/sideways/down instead of verbal response."),
          ("Advanced learner","Introduce digital safety concepts: online spaces have boundaries too."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Personal space concept understood?","&#9744; Yes  &#9744; Partially  &#9744; Not yet"),
          ("Comfortable/uncomfortable discrimination","&#9744; Consistent  &#9744; Partial  &#9744; Needs more work"),
          ("Assertive response used?","&#9744; Spontaneous  &#9744; Prompted  &#9744; Not yet"),
          ("Voice quality during 'Stop' practice","&#9744; Clear and assertive  &#9744; Too quiet  &#9744; Aggressive  &#9744; No response"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} learned about personal space and practiced recognising comfortable and uncomfortable situations. We also practiced saying 'No thank you' and 'Stop, I don't like that' using puppets and role-play.",
        "parent_why": "Children who understand and can communicate their own boundaries are significantly safer and more socially confident. This skill protects {child} and helps them respect others too.",
        "parent_tips": [
          "Always respect {child}'s 'no' to physical contact — even hugs from family. This models that their boundaries matter.",
          "Practice 'Stop, I don't like that' as a game: take turns with a silly tickle game where {child} can say stop.",
          "Use books or stories where characters set boundaries and are respected — discuss them together.",
          "Never force {child} to hug or kiss anyone. Offer a wave or high five as an alternative.",
        ],
        "substitutes": [
          "Hula hoop: chalk circle on floor, a drawn circle on paper, or imaginary bubble",
          "Picture cards: describe scenarios verbally or draw simple stick figures",
          "Puppets: any stuffed animals or dolls",
        ],
      }],
      "Hyperactivity Reduction": [{
        "name": "Energy Dial",
        "skill_label": "Hyperactivity Reduction and Self-Regulation",
        "color": C_ORANGE, "light": C_ORANGE_LIGHT,
        "objective": "{child} will develop awareness of their own activity level and build strategies to modulate energy when the environment requires it. Sessions will use a visual Energy Dial (1 = very calm to 5 = very active) to build body awareness and self-regulation vocabulary.",
        "materials": [
          ("Primary Tools","Energy Dial visual (drawn 1-5 scale), movement activity cards, calm-down activity cards, token board"),
          ("Substitutes","Hand-drawn dial on paper; describe movement activities verbally; any sensory calming items"),
          ("Reinforcers","Token board + 2 minutes free choice movement at session end"),
        ],
        "steps": [
          ("Set up (before session)","Draw the Energy Dial. Prepare 3 movement cards (jumping, spinning, marching) and 3 calm-down cards (breathing, slow walking, squeezing a ball)."),
          ("Introduce the Dial (3 min)","'Your body has an energy level — like a TV volume! Sometimes it needs to be on 5, sometimes on 2.' Walk through each level with a body demonstration."),
          ("Where are you now? (2 min)","Ask {child} to point to their level on the dial. No right or wrong answer. Validate. Record starting level."),
          ("High energy activity (5 min)","Let {child} engage in an approved high-energy activity (jumping jacks, animal walks). 'Get that energy OUT!'"),
          ("Dial it down (8 min)","'Now let's bring the dial down to a 2.' Practice slow deep breaths, slow motion walking, or gentle squeezing. Check dial again. Award token for each dial reduction attempted."),
          ("Generalisation discussion (2 min)","'When do you need to bring your dial down at nursery?' Identify 2 real moments. This is the key generalisation step."),
          ("Wrap-up","Record dial start and end levels. 'You learned how to control your own dial today!'"),
        ],
        "adapt": [
          ("Cannot self-rate","Therapist rates together: 'I think your body is at a 4 right now — do you agree?'"),
          ("Refuses calm-down activities","Offer choice between two calm-down options. Reduce duration to 30 seconds initially."),
          ("High energy throughout","Ensure significant physical release before calm-down. Never skip the high energy phase."),
          ("Very low energy","Still teach the dial — help child recognize when they could bring energy UP (agency in both directions)."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Starting dial level (self-reported)","_______"),
          ("Ending dial level (self-reported)","_______"),
          ("Largest dial reduction achieved","_______  levels"),
          ("Calm-down strategy used","_____________________________"),
          ("Self-rating accuracy (vs therapist observation)","&#9744; Accurate  &#9744; Over-estimated  &#9744; Under-estimated"),
          ("Generalisation moment identified?","&#9744; Yes - recorded above  &#9744; Not yet"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} learned about their Energy Dial — a 1-5 scale for how active their body feels. We practiced getting energy OUT safely, and then bringing the dial back down using calm-down strategies.",
        "parent_why": "Children with high energy levels often struggle not because they are being naughty, but because they have not yet learned how to self-regulate their arousal. Teaching them to notice and name their energy level is the first step.",
        "parent_tips": [
          "Use the dial language at home: 'What level is your body on right now?' Accept any answer without judgment.",
          "Before a sit-down activity, offer 5 minutes of physical movement first — this helps the brain settle.",
          "Celebrate any moment you notice {child} slowing down or choosing a calmer activity independently.",
          "Keep calm-down strategies consistent between home and nursery — ask the therapist for the specific strategies we use.",
        ],
        "substitutes": [
          "Energy Dial: hold up fingers 1-5, or draw a simple thermometer scale",
          "Movement cards: describe activities verbally or let child choose from familiar options",
          "Calm-down items: any sensory object, a favourite quiet toy, or a cozy corner",
        ],
      }],
      "Basic Knowledge": [{
        "name": "Discovery Box Challenge",
        "skill_label": "Basic Knowledge and Foundational Concepts",
        "color": C_GREEN, "light": C_GREEN_LIGHT,
        "objective": "{child} will consolidate foundational concepts including colours, shapes, numbers, and simple categorisation within a structured, play-based exploration framework. Progress is measured by increasing accuracy and independence in identifying and using target concepts.",
        "materials": [
          ("Primary Tools","Small box filled with objects of different colours and shapes, simple sorting mats, number cards 1-10, token board"),
          ("Substitutes","Any household objects; hand-drawn sorting mats; finger counting in place of number cards"),
          ("Reinforcers","Token board + child decorates their own 'Discovery Box' at end of session series"),
        ],
        "steps": [
          ("Set up (before session)","Fill box with 8-10 objects covering target concepts. Prepare sorting mats. Decide which concepts to focus on this session."),
          ("Reveal the box (2 min)","'I have a mystery box with amazing things inside!' Build excitement before opening. Let {child} reach in and pull items out one by one."),
          ("Identify and sort (10 min)","For each object: name it, identify its colour, shape, or category. Use questioning: 'What colour is this? Where does it go?' Prompt as needed. Token for each correct attempt."),
          ("Number practice (5 min)","Count the objects in each sorted group. Point to and count together. Use number cards to match. 'How many red ones do we have?'"),
          ("Concept game (5 min)","Simon Says using concepts: 'Simon says touch something ROUND / touch something BLUE.' Quick-fire, fun, no wrong answers punished."),
          ("Wrap-up (3 min)","'You discovered so many things today!' Review two concepts learned or practiced. Count tokens."),
        ],
        "adapt": [
          ("Knows all presented concepts","Increase complexity: introduce letters, sequences, more/less, opposites."),
          ("Struggles with multiple concepts","Focus on one concept only per session until solid."),
          ("Not interested in objects","Switch to favourite materials: toy cars sorted by colour, food items for categories."),
          ("Language barrier","Use gesture-based sorting: pointing and placing, no verbal labels required initially."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Concepts targeted this session","_____________________________"),
          ("Accuracy on target concepts","_______  out of  _______  correct"),
          ("Prompt level needed","&#9744; Independent  &#9744; Gestural  &#9744; Verbal  &#9744; Full physical"),
          ("Number range accurately counted","1 to  _______"),
          ("Any concept fully mastered this session?","&#9744; Yes - note which  &#9744; Not yet"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} explored our Discovery Box — full of colourful objects! We practiced identifying colours, shapes, and categories, and did some counting together.",
        "parent_why": "Foundational concepts like colours, numbers, and shapes are the building blocks of all academic learning. Practising them through play in a relaxed, positive environment is the most effective way to make them stick.",
        "parent_tips": [
          "Name concepts naturally throughout the day: 'Can you pass me the ROUND orange? The big red cup?'",
          "Count everything together: stairs, grapes, socks, anything. Make it a routine.",
          "Sort laundry by colour or shape together — real-life practice is just as valuable as structured sessions.",
          "Celebrate every correct identification immediately and specifically: 'You knew that was a triangle!'",
        ],
        "substitutes": [
          "Discovery box: any bag, basket, or box; household items work perfectly",
          "Sorting mats: drawn circles on paper labelled by colour/shape/category",
          "Number cards: finger counting, drawn number lines, or number magnets",
        ],
      }],
      "Emotional Resilience": [{
        "name": "The Bounce-Back Story",
        "skill_label": "Emotional Resilience and Recovery",
        "color": C_PURPLE, "light": C_PURPLE_LIGHT,
        "objective": "{child} will develop a positive narrative around mistakes and setbacks, build perspective-taking skills after difficult moments, and strengthen their sense of themselves as capable and resilient. By the end of the session series, {child} will demonstrate increased emotional recovery speed and a more adaptive response to disappointment.",
        "materials": [
          ("Primary Tools","Simple picture books or story cards featuring characters overcoming challenges, a 'Bounce-Back Book' (blank booklet), token board"),
          ("Substitutes","Verbally narrate stories; draw simple stick figure scenarios; use puppets"),
          ("Reinforcers","Token board + child adds a page to their personal Bounce-Back Book each session"),
        ],
        "steps": [
          ("Set up (before session)","Choose a simple story featuring a character who faces a setback and recovers. Prepare blank booklet page."),
          ("Story time (5 min)","Read or narrate the story. Pause at the hard moment: 'Oh no — how do you think they feel?' Pause at recovery: 'What helped them bounce back?'"),
          ("Connect to real life (5 min)","'Has something like that ever happened to you? What happened?' Listen without fixing. Validate the hard feeling first, then explore recovery."),
          ("What helps me bounce back? (8 min)","Brainstorm together: 'What are YOUR bounce-back tools?' Draw or dictate 2-3 strategies and add them to the booklet page."),
          ("Resilience affirmations (3 min)","Teach 1-2 simple phrases: 'I can try again.' 'Making mistakes is how I learn.' 'I am getting stronger every time.' Practice saying them together."),
          ("Wrap-up","Add the booklet page. 'Look — you have your own Bounce-Back Book now!'"),
        ],
        "adapt": [
          ("Cannot recall a setback","Use only the story character. Direct personal connection is not required."),
          ("Becomes upset when discussing mistakes","Shift entirely to the fictional character. Keep safe distance."),
          ("Denies ever having hard moments","Accept it. Plant the seed. Try again in later sessions with a lighter touch."),
          ("Very resilient already","Focus on empathy: how might others feel when things go wrong? Build outward resilience perspective."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Story connection made?","&#9744; Engaged with story  &#9744; Partial  &#9744; Disengaged"),
          ("Personal setback shared?","&#9744; Yes - spontaneous  &#9744; Yes - prompted  &#9744; No"),
          ("Bounce-back strategies identified","_____________________________"),
          ("Affirmation practiced?","&#9744; Yes  &#9744; Not yet"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} heard a story about a character who faced a hard moment and bounced back. We talked about what helps us recover from difficult things and started building a personal Bounce-Back Book.",
        "parent_why": "Resilience is not about never struggling — it is about recovering. Children who believe they can bounce back from setbacks are significantly more confident and emotionally healthy.",
        "parent_tips": [
          "When things go wrong, name the feeling first: 'That was really disappointing.' Then, after a pause: 'What could we try next time?'",
          "Share your own small setbacks and how you recovered: 'I made a mistake today and felt frustrated — then I tried again.'",
          "Praise effort and recovery, not just success: 'I am so proud of how you kept trying even when it was hard.'",
          "Read picture books together about characters overcoming challenges — ask {child} to predict what the character will do.",
        ],
        "substitutes": [
          "Picture books: any story with a character facing and overcoming a challenge",
          "Bounce-Back Book: any blank paper stapled together; decorate the cover together",
          "Story cards: hand-drawn stick figure scenarios",
        ],
      }],
      "Independence and Self-Care": [{
        "name": "I Can Do It Steps",
        "skill_label": "Independence and Self-Care Skills",
        "color": C_BLUE, "light": C_BLUE_LIGHT,
        "objective": "{child} will develop age-appropriate independence in a targeted daily living skill using task analysis, visual step cards, and systematic reinforcement. By the end of the session series, {child} will complete the target skill with decreasing levels of adult prompting.",
        "materials": [
          ("Primary Tools","Visual step cards for the target skill (e.g., hand washing, putting on shoes), token board, the actual materials for the skill"),
          ("Substitutes","Hand-drawn step pictures; photograph sequence on a phone or tablet; any available real materials"),
          ("Reinforcers","Token board + child marks off each step on their own visual chart"),
        ],
        "steps": [
          ("Set up (before session)","Identify the specific self-care skill to target (selected with family). Prepare visual step cards. Gather materials."),
          ("Introduce the steps (3 min)","Show visual cards: 'These are the steps for [skill]. Let us count them together.' Walk through each step while narrating."),
          ("Full model (3 min)","Therapist models the complete skill from start to finish. Narrate each step clearly. Use the visual cards as reference."),
          ("Supported practice (10 min)","{child} attempts the skill with visual cards visible. Use a prompt hierarchy: wait → gesture → verbal → physical. Token for each step completed."),
          ("Independent attempt (5 min)","Remove prompts. Let {child} try with only the visual cards. Accept a partial attempt as success. Celebrate the furthest independent step reached."),
          ("Wrap-up (2 min)","Count steps completed independently. Record. 'Every session you do a little more on your own!'"),
        ],
        "adapt": [
          ("Task too complex","Break into smaller sub-steps. Target only step 1 until mastered."),
          ("Refuses to attempt","Reduce to observation only this session. Try backward chaining: complete all steps for them and let {child} do only the final step."),
          ("Completes task but needs prompting","Focus entirely on fading prompts rather than adding complexity."),
          ("Masters skill quickly","Add a new self-care target or increase independence expectation in unfamiliar environments."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Skill targeted this session","_____________________________"),
          ("Total steps in task","_______"),
          ("Steps completed independently","_______  out of  _______"),
          ("Prompt level required for remaining steps","&#9744; Gestural  &#9744; Verbal  &#9744; Physical  &#9744; N/A"),
          ("Furthest independent step reached","Step #  _______"),
          ("Any refusal?","&#9744; No  &#9744; Yes - managed how: _____"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} practiced [the target self-care skill] using a step-by-step visual guide. We worked on completing as many steps as possible independently, celebrating every step done without help.",
        "parent_why": "Independence in daily routines builds confidence, reduces conflict at transition times, and gives children a genuine sense of competence. Every step done alone is a real achievement.",
        "parent_tips": [
          "Use the same visual step cards at home if possible — ask the therapist for a copy.",
          "Give {child} time to attempt each step before stepping in to help.",
          "Use the same words and sequence we use in sessions to keep it consistent.",
          "Never take over a step {child} can manage — even if it takes longer. The time is worth it.",
        ],
        "substitutes": [
          "Visual step cards: photographs on a phone or tablet; drawn pictures; written words if {child} can read",
          "Real materials: whatever is available at home for the target skill",
        ],
      }],
      "Transitions and Flexibility": [{
        "name": "The Transition Train",
        "skill_label": "Transitions and Cognitive Flexibility",
        "color": C_TEAL, "light": C_TEAL_LIGHT,
        "objective": "{child} will develop improved tolerance for changes in activity and routine using visual transition tools, advance warning strategies, and structured preparation rituals. By the end of the session series, {child} will demonstrate smoother transitions with decreasing distress and adult support.",
        "materials": [
          ("Primary Tools","Visual schedule (4-5 activity pictures), countdown timer, transition signal (bell or clap pattern), token board"),
          ("Substitutes","Drawn pictures for schedule; count aloud instead of timer; consistent verbal warning phrase as signal"),
          ("Reinforcers","Token board + {child} gets to place each activity picture on the schedule themselves"),
        ],
        "steps": [
          ("Set up (before session)","Prepare visual schedule with 4-5 activities for today's session. Have timer ready."),
          ("Preview the session (2 min)","Show the schedule: 'Here is what we are doing today — and here is when each thing ends.' Let {child} place their own picture if possible. Predictability is the goal."),
          ("First transition practice (8 min)","Run first activity. Set timer. Give 2-minute warning: 'Two more minutes, then we go to [next activity].' When timer sounds, use the transition signal. Support calmly through the change."),
          ("Second transition practice (8 min)","Repeat with next activity. Reduce support if first went well. Token awarded for each calm transition."),
          ("Surprise transition (4 min)","Without warning: 'Actually — we are doing something different now!' Practice the unexpected change. Coach: 'Plans change sometimes — that is OK.' Support as needed but aim for less than first transition."),
          ("Wrap-up (3 min)","Count tokens. 'You switched activities so many times today and you did it!'"),
        ],
        "adapt": [
          ("Refuses all transitions","Use only 2 activities. Increase warning time to 5 minutes. Let {child} hold the timer."),
          ("Distressed by surprise change","Skip the surprise transition until planned transitions are solid. Introduce gently later."),
          ("Transitions fine, but rigid in preferences","Focus on flexibility within a single activity: change the rules slightly mid-game."),
          ("Already flexible","Add complexity: reduce warning time; increase number of transitions; introduce multi-step sequences."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Number of transitions attempted","_______"),
          ("Transitions completed calmly","_______  out of  _______"),
          ("Warning time used","_______  minutes"),
          ("Surprise transition introduced?","&#9744; Yes  &#9744; Not yet"),
          ("Response to surprise transition","&#9744; Calm  &#9744; Brief distress, recovered  &#9744; Significant distress"),
          ("Prompt level needed","&#9744; Independent  &#9744; Gestural  &#9744; Verbal  &#9744; Full support"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} practiced moving between activities using a visual schedule and a countdown timer. We worked on listening for the transition signal and switching activities calmly — including one surprise change!",
        "parent_why": "Transitions are one of the most common trigger points for challenging behaviour in young children. Giving advance warning and keeping routines predictable can dramatically reduce distress around changes.",
        "parent_tips": [
          "Give {child} a 2-5 minute warning before any activity change: 'In 3 minutes we are going to tidy up.'",
          "Use a consistent transition signal at home — the same phrase or sound every time.",
          "Keep the visual schedule visible at home for the day's routine — mornings and evenings especially.",
          "When {child} manages a hard transition, name it: 'That was a big change and you handled it so well.'",
        ],
        "substitutes": [
          "Visual schedule: drawn pictures; photographs; word list if child can read",
          "Timer: count aloud; use a sand timer; set a phone alarm with a gentle sound",
          "Transition signal: a consistent clap pattern, a specific phrase, or a small bell",
        ],
      }],
      "Play Skills": [{
        "name": "Imagination Station",
        "skill_label": "Play Skills and Imaginative Engagement",
        "color": C_GOLD, "light": C_GOLD_LIGHT,
        "objective": "{child} will develop structured and imaginative play skills, including purposeful engagement with play materials and the capacity for shared imaginative narratives with a play partner. By the end of the session series, {child} will demonstrate increasing complexity and duration of play engagement.",
        "materials": [
          ("Primary Tools","Assorted play materials (figures, vehicles, craft supplies, building blocks), role-play props, token board"),
          ("Substitutes","Any available toys; household objects as props; draw scenes if no figures available"),
          ("Reinforcers","Token board + child keeps a small drawing or creation from each session"),
        ],
        "steps": [
          ("Set up (before session)","Arrange 2-3 play material options. Have a simple play scenario idea ready but be prepared to follow {child}'s lead."),
          ("Child-led exploration (5 min)","Let {child} explore materials freely. Observe play style. Join as a parallel player without directing."),
          ("Extend the play (8 min)","Add an element to the play: a character, a problem, a destination. Follow {child}'s lead. Narrate collaboratively: 'And then what happened?' Scaffold without taking over."),
          ("Cooperative play challenge (5 min)","Introduce a simple shared goal: 'Let us build a house for all the animals together.' Practice sharing materials and negotiating roles."),
          ("Story capture (3 min)","Ask {child} to tell you or draw what happened in the play. 'If this was a book, what would it be called?' Build narrative language."),
          ("Wrap-up (2 min)","Count tokens. Celebrate imagination: 'The story we made today was amazing!'"),
        ],
        "adapt": [
          ("Functional play only, no imagination","Scaffold: 'What if this block was a bed? Who sleeps here?' Introduce fiction gently."),
          ("Solitary play only","Parallel play alongside without intrusion. Gradual joining over multiple sessions."),
          ("Directs therapist rigidly","Accept the role fully first. Gradually introduce small variations to test flexibility."),
          ("Very advanced imagination","Co-create more complex multi-session narratives. Introduce characters with emotions and dilemmas."),
        ],
        "tracking_rows": [
          ("Engagement level (1-5)","&#9744; 1  &#9744; 2  &#9744; 3  &#9744; 4  &#9744; 5"),
          ("Play type observed","&#9744; Solitary  &#9744; Parallel  &#9744; Associative  &#9744; Cooperative"),
          ("Imaginative elements present?","&#9744; Yes - describe in notes  &#9744; Functional only"),
          ("Duration of sustained play","_______  minutes"),
          ("Shared goal achieved?","&#9744; Yes  &#9744; Partially  &#9744; Not yet"),
          ("Narrative language quality","&#9744; Rich and varied  &#9744; Simple labels  &#9744; Non-verbal only"),
          ("Material sharing observed?","&#9744; Spontaneous  &#9744; Prompted  &#9744; Refused"),
          ("Any dysregulation?","&#9744; No  &#9744; Yes - describe in notes"),
          ("Tokens earned","_______"),
          ("Plan fidelity","&#9744; Fully implemented  &#9744; Modified - see notes"),
        ],
        "parent_what": "Today {child} and I created a story together using play materials. We practiced following each other's ideas, sharing, and building something together. {child}'s imagination is wonderful!",
        "parent_why": "Play is the primary way children learn about the world, relationships, and themselves. Rich, imaginative play builds language, problem-solving, social skills, and emotional understanding simultaneously.",
        "parent_tips": [
          "Set aside 10-15 minutes of uninterrupted, screen-free play time daily. Follow {child}'s lead entirely.",
          "Join the play as a character, not a director: 'Can I be the shopkeeper? What do you need today?'",
          "Narrate play aloud: 'Oh, the bear is feeling hungry — what should he do?' This builds storytelling language.",
          "Resist organising or correcting the play. There is no right way to imagine.",
        ],
        "substitutes": [
          "Play figures: any small objects, clothes pegs, or hand-drawn characters on paper",
          "Role-play props: household items work perfectly — a bowl becomes a cauldron, a box becomes a spaceship",
        ],
      }],
    }

    selected_skills = [s for s in all_skills if s in ACTIVITY_LIBRARY]
    if not selected_skills:
        selected_skills = list(ACTIVITY_LIBRARY.keys())[:3]

    activity_pool = []
    for sk in selected_skills:
        for act in ACTIVITY_LIBRARY[sk]:
            activity_pool.append(act)

    if not activity_pool:
        activity_pool = [ACTIVITY_LIBRARY["Emotional Regulation"][0]]

    num_acts = len(activity_pool)
    base, remainder = divmod(sessions, num_acts)
    session_counts = [base + (1 if i < remainder else 0) for i in range(num_acts)]
    session_counts = [max(c, 1) for c in session_counts]

    schedule = []
    session_num = 1
    for act, count in zip(activity_pool, session_counts):
        sess_nums = list(range(session_num, session_num + count))
        schedule.append((act, count, sess_nums))
        session_num += count

    def _S(name, **kw):
        return ParagraphStyle(name, **kw)

    ST = {
        "page_title": _S("pt", fontName="Helvetica-Bold", fontSize=18, textColor=C_DARK, leading=24),
        "h2":         _S("h2", fontName="Helvetica-Bold", fontSize=12, textColor=C_DARK, leading=16),
        "h3":         _S("h3", fontName="Helvetica",      fontSize=10, textColor=C_MID_GREY, leading=14),
        "body":       _S("bd", fontName="Helvetica",      fontSize=9.5, textColor=C_DARK, leading=14, alignment=TA_JUSTIFY),
        "body_sm":    _S("bs", fontName="Helvetica",      fontSize=8.5, textColor=C_DARK, leading=12),
        "bold_sm":    _S("bls",fontName="Helvetica-Bold", fontSize=8.5, textColor=C_DARK, leading=12),
        "label":      _S("lb", fontName="Helvetica-Bold", fontSize=8,   textColor=C_MID_GREY, leading=11),
        "bullet":     _S("bu", fontName="Helvetica",      fontSize=9.5, textColor=C_DARK, leading=14, leftIndent=14),
        "footer":     _S("ft", fontName="Helvetica",      fontSize=7.5, textColor=C_MID_GREY, alignment=TA_CENTER),
        "parent":     _S("pa", fontName="Helvetica",      fontSize=10,  textColor=C_DARK, leading=15, alignment=TA_JUSTIFY),
        "sec_hdr":    _S("sh", fontName="Helvetica-Bold", fontSize=11,  textColor=C_WHITE, alignment=TA_LEFT, leading=16, leftIndent=6),
    }

    def _sp(n=1): return Spacer(1, n * 0.3 * cm)
    def _hr(accent): return HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=6, spaceBefore=6)

    def _sec(label, accent):
        data = [[Paragraph(label, ST["sec_hdr"])]]
        t = Table(data, colWidths=["100%"])
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,-1), accent),
            ("TOPPADDING",    (0,0), (-1,-1), 6),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ("LEFTPADDING",   (0,0), (-1,-1), 10),
        ]))
        return t

    def _track_table(rows):
        data = [[Paragraph(a, ST["bold_sm"]), Paragraph(b, ST["body_sm"])] for a, b in rows]
        t = Table(data, colWidths=[7.5*cm, 10*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (0,-1), C_LIGHT_GREY),
            ("ROWBACKGROUNDS",(0,0), (-1,-1), [C_WHITE, HexColor("#F9FBFC")]),
            ("GRID",          (0,0), (-1,-1), 0.5, C_BORDER),
            ("VALIGN",        (0,0), (-1,-1), "TOP"),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 7),
            ("RIGHTPADDING",  (0,0), (-1,-1), 7),
        ]))
        return t

    def _tip_box(tips, accent, light):
        rows = [[Paragraph("HOME TIPS", ST["label"])]]
        for tip in tips:
            rows.append([Paragraph("*  " + tip.replace("{child}", child), ST["body_sm"])])
        t = Table(rows, colWidths=[17.5*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (0,0), accent),
            ("BACKGROUND",    (0,1), (0,-1), light),
            ("GRID",          (0,0), (-1,-1), 0.5, C_BORDER),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 8),
        ]))
        return t

    def _chrome(c_obj, doc_obj, bar_title, bar_sub, accent, tag=None, psize=A4):
        w, h = psize
        c_obj.setFillColor(accent)
        c_obj.rect(0, h - 2.2*cm, w, 2.2*cm, fill=1, stroke=0)
        c_obj.setFillColor(C_WHITE); c_obj.setFont("Helvetica-Bold", 12)
        c_obj.drawString(1.5*cm, h - 1.35*cm, bar_title)
        c_obj.setFont("Helvetica", 8.5)
        c_obj.drawString(1.5*cm, h - 1.85*cm, bar_sub)
        if tag:
            tx = w - 3.2*cm; ty = h - 1.65*cm
            c_obj.setFillColor(C_WHITE)
            c_obj.roundRect(tx, ty - 0.3*cm, 2.8*cm, 0.7*cm, 4, fill=1, stroke=0)
            c_obj.setFillColor(accent); c_obj.setFont("Helvetica-Bold", 7)
            c_obj.drawCentredString(tx + 1.4*cm, ty - 0.08*cm, tag)
        c_obj.setFillColor(C_LIGHT_GREY); c_obj.rect(0, 0, w, 1.0*cm, fill=1, stroke=0)
        c_obj.setFillColor(C_MID_GREY); c_obj.setFont("Helvetica", 7)
        c_obj.drawString(1.5*cm, 0.35*cm, f"{child}  |  {nursery}  |  Therapist: {therapist}")
        c_obj.drawRightString(w - 1.5*cm, 0.35*cm, "CONFIDENTIAL")
        c_obj.setFillColor(accent); c_obj.setFillAlpha(0.12)
        c_obj.rect(0, 1.0*cm, 0.35*cm, h - 3.2*cm, fill=1, stroke=0)
        c_obj.setFillAlpha(1.0)

    def _make_temp(content_fn, chrome_fn, psize=A4):
        buf2 = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        buf2.close()
        doc2 = SimpleDocTemplate(
            buf2.name, pagesize=psize,
            leftMargin=1.5*cm, rightMargin=1.5*cm,
            topMargin=3.0*cm, bottomMargin=1.8*cm,
        )
        doc2.build(content_fn(), onFirstPage=chrome_fn, onLaterPages=chrome_fn)
        return buf2.name

    temp_files = []

    # Cover
    def make_cover():
        tmp2 = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp2.close()
        c2 = rl_canvas.Canvas(tmp2.name, pagesize=A4)
        w, h = A4
        c2.setFillColor(C_TEAL); c2.rect(0, 0, w, h, fill=1, stroke=0)
        c2.setFillColor(HexColor("#1C5F6B")); c2.rect(0, 0, w, h*0.42, fill=1, stroke=0)
        c2.setFillColor(C_WHITE); c2.setFillAlpha(0.05)
        c2.circle(w+1*cm, h-2*cm, 9*cm, fill=1, stroke=0)
        c2.circle(-2*cm, 3*cm, 7*cm, fill=1, stroke=0)
        c2.setFillAlpha(1.0)
        c2.setFillColor(C_GOLD); c2.rect(0, h*0.42, w, 0.5*cm, fill=1, stroke=0)
        c2.setFillColor(C_GOLD)
        c2.roundRect(1.5*cm, h-3.2*cm, 5.5*cm, 0.9*cm, 5, fill=1, stroke=0)
        c2.setFillColor(C_DARK); c2.setFont("Helvetica-Bold", 9)
        c2.drawCentredString(1.5*cm+2.75*cm, h-2.85*cm, nursery.upper()[:28])
        c2.setFillColor(C_WHITE); c2.setFont("Helvetica-Bold", 36)
        c2.drawCentredString(w/2, h-5.5*cm, "Activity Session")
        c2.drawCentredString(w/2, h-6.5*cm, "Pack")
        c2.setFont("Helvetica", 14); c2.setFillColor(HexColor("#B2DDE5"))
        c2.drawCentredString(w/2, h-7.8*cm, "Individual Behaviour Support Programme")
        c2.setStrokeColor(C_GOLD); c2.setLineWidth(1.5)
        c2.line(w/2-4*cm, h-8.4*cm, w/2+4*cm, h-8.4*cm)
        c2.setFillColor(HexColor("#1C5F6B"))
        c2.roundRect(w/2-5.5*cm, h*0.42+1*cm, 11*cm, 5.5*cm, 8, fill=1, stroke=0)
        c2.setFillColor(C_WHITE); c2.setFont("Helvetica-Bold", 22)
        c2.drawCentredString(w/2, h*0.42+5.3*cm, child)
        c2.setFont("Helvetica", 11); c2.setFillColor(HexColor("#B2DDE5"))
        c2.drawCentredString(w/2, h*0.42+4.5*cm, f"Age: {age}   |   Class: {class_group}")
        c2.setFont("Helvetica", 10); c2.setFillColor(HexColor("#7ABEC9"))
        c2.drawCentredString(w/2, h*0.42+3.6*cm, f"Therapist: {therapist}")
        c2.drawCentredString(w/2, h*0.42+2.9*cm, f"Sessions: {freq_label}  |  {sessions} per month")
        c2.drawCentredString(w/2, h*0.42+2.2*cm, f"Start Date: {start_str}")
        c2.drawCentredString(w/2, h*0.42+1.5*cm, f"Activities Planned: {len(schedule)}")
        act_colors_cv = [C_TEAL, C_GOLD, C_GREEN, C_PURPLE, C_CORAL, C_ORANGE, C_BLUE]
        tag_w2 = min(3.5*cm, (w-3*cm)/max(len(schedule),1))
        tag_gap2 = 0.3*cm
        total_w2 = len(schedule)*tag_w2 + (len(schedule)-1)*tag_gap2
        start_x2 = (w-total_w2)/2
        for i2, (act2, cnt2, _) in enumerate(schedule):
            col2 = act_colors_cv[i2 % len(act_colors_cv)]
            tx2 = start_x2 + i2*(tag_w2+tag_gap2)
            c2.setFillColor(col2)
            c2.roundRect(tx2, 2.8*cm, tag_w2, 1.2*cm, 5, fill=1, stroke=0)
            c2.setFillColor(C_WHITE); c2.setFont("Helvetica-Bold", 6.5)
            words2 = act2["name"].split()
            l1 = " ".join(words2[:2]); l2 = " ".join(words2[2:]) if len(words2)>2 else ""
            c2.drawCentredString(tx2+tag_w2/2, 3.7*cm, l1[:16])
            if l2: c2.drawCentredString(tx2+tag_w2/2, 3.25*cm, l2[:16])
        c2.setFillColor(HexColor("#154952")); c2.rect(0, 0, w, 1.8*cm, fill=1, stroke=0)
        c2.setFillColor(HexColor("#7ABEC9")); c2.setFont("Helvetica", 7.5)
        c2.drawCentredString(w/2, 0.9*cm, "CONFIDENTIAL - For therapist and family use only.")
        c2.drawCentredString(w/2, 0.45*cm, f"Generated: {today_str}  |  {nursery}")
        c2.save()
        return tmp2.name

    temp_files.append(make_cover())

    # Schedule overview page
    def sched_content():
        s = [_sp(4)]
        s.append(Paragraph("Session Schedule Overview", ST["page_title"]))
        s.append(Paragraph(f"Child: {child}  |  Therapist: {therapist}  |  {freq_label}  |  {sessions} sessions planned", ST["h3"]))
        s.append(_hr(C_TEAL)); s.append(_sp())
        s.append(Paragraph(
            f"This pack contains {len(schedule)} activit{'y' if len(schedule)==1 else 'ies'} "
            f"planned across {sessions} sessions this month. Activities have been selected based on "
            f"the targeted skill areas and distributed to ensure balanced coverage. "
            f"Where a skill requires reinforcement, an activity is assigned across multiple sessions - "
            f"this is intentional and reflects best practice for early childhood skill acquisition.",
            ST["body"]))
        s.append(_sp(1.5))
        if behaviors:
            s.append(_sec("Behaviours Being Addressed", C_TEAL)); s.append(_sp(0.5))
            s.append(Paragraph("  *  ".join(behaviors), ST["body_sm"]))
            s.append(_sp(1.5))
        s.append(_sec("Activity Distribution Across Sessions", C_TEAL)); s.append(_sp(0.5))
        hdr = [Paragraph(x, ST["bold_sm"]) for x in ["Act #","Activity Name","Skill Targeted","Sessions Assigned","Session Numbers"]]
        tdata = [hdr]
        for i3, (act3, cnt3, snums3) in enumerate(schedule, 1):
            rnote = f"{cnt3} session{'s' if cnt3>1 else ''}"
            if cnt3 > 1:
                rnote += " (repeated - foundational skill)"
            tdata.append([
                Paragraph(str(i3), ST["body_sm"]),
                Paragraph(act3["name"], ST["bold_sm"]),
                Paragraph(act3["skill_label"], ST["body_sm"]),
                Paragraph(rnote, ST["body_sm"]),
                Paragraph(", ".join([f"#{n}" for n in snums3]), ST["body_sm"]),
            ])
        t3 = Table(tdata, colWidths=[1.2*cm, 5.5*cm, 4.5*cm, 4*cm, 2.3*cm])
        t3.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,0), C_TEAL),
            ("TEXTCOLOR",     (0,0), (-1,0), C_WHITE),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [C_TEAL_LIGHT, C_WHITE]),
            ("GRID",          (0,0), (-1,-1), 0.5, C_BORDER),
            ("VALIGN",        (0,0), (-1,-1), "TOP"),
            ("TOPPADDING",    (0,0), (-1,-1), 6),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ("LEFTPADDING",   (0,0), (-1,-1), 7),
        ]))
        s.append(t3); s.append(_sp(1.5))
        s.append(_sec("Recommended Session Structure (Every Session)", C_TEAL)); s.append(_sp(0.5))
        phases = [
            ("0-3 min",   "Arrival and settling - greet child warmly, allow brief free-choice moment."),
            ("3-7 min",   "Warm-up check-in - use emotion cards as visual prompt regardless of the day's activity."),
            ("7-25 min",  "Core activity - follow the Conduct Guide. Use token board throughout."),
            ("25-28 min", "Cool-down - brief positive reflection. Avoid lengthy discussion."),
            ("28-30 min", "Transition out - clear verbal and visual signal. Praise any calm transition."),
        ]
        pdata2 = [[Paragraph(a2, ST["bold_sm"]), Paragraph(b2, ST["body_sm"])] for a2,b2 in phases]
        pt2 = Table(pdata2, colWidths=[2.5*cm, 15*cm])
        pt2.setStyle(TableStyle([
            ("ROWBACKGROUNDS",(0,0), (-1,-1), [C_TEAL_LIGHT, C_WHITE]),
            ("GRID",          (0,0), (-1,-1), 0.5, C_BORDER),
            ("VALIGN",        (0,0), (-1,-1), "TOP"),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 7),
        ]))
        s.append(pt2)
        return s

    def sched_chrome(c2, doc2):
        _chrome(c2, doc2, "Session Schedule Overview", f"{child}  |  {sessions} sessions  |  {len(schedule)} activities", C_TEAL, "SCHEDULE")

    temp_files.append(_make_temp(sched_content, sched_chrome))

    # Per-activity pages
    for act_idx, (act, cnt, sess_nums) in enumerate(schedule):
        accent = act["color"]; light = act["light"]
        aname = act["name"]; anum = act_idx + 1
        sess_label = f"Sessions: {', '.join([f'#{n}' for n in sess_nums])}"
        repeat_note = (
            f"This activity is conducted across {cnt} sessions (Sessions {', '.join([str(n) for n in sess_nums])}) - "
            f"repeated because '{act['skill_label']}' is a foundational skill requiring consistent, spaced practice."
            if cnt > 1 else ""
        )

        def make_conduct_content(act=act, accent=accent, light=light, anum=anum, aname=aname, sess_label=sess_label, repeat_note=repeat_note):
            def content():
                s = [_sp(4)]
                s.append(Paragraph(f"Activity {anum}: {aname}", ST["page_title"]))
                s.append(Paragraph(f"Skill: {act['skill_label']}  |  Duration: 20-30 min  |  {sess_label}", ST["h3"]))
                s.append(_hr(accent))
                if repeat_note:
                    rt = Table([[Paragraph(repeat_note, ST["body_sm"])]], colWidths=[17.5*cm])
                    rt.setStyle(TableStyle([
                        ("BACKGROUND",(0,0),(-1,-1),light),
                        ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
                        ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
                        ("LEFTPADDING",(0,0),(-1,-1),8),
                    ]))
                    s.append(rt); s.append(_sp())
                s.append(_sec("Objective", accent)); s.append(_sp(0.5))
                s.append(Paragraph(act["objective"].format(**{"child":child}), ST["body"]))
                s.append(_sp(1.5))
                s.append(_sec("Materials Needed", accent)); s.append(_sp(0.5))
                mdata = [[Paragraph(a4,ST["bold_sm"]),Paragraph(b4,ST["body_sm"])] for a4,b4 in act["materials"]]
                mt = Table(mdata, colWidths=[4*cm,13.5*cm])
                mt.setStyle(TableStyle([
                    ("ROWBACKGROUNDS",(0,0),(-1,-1),[light,C_WHITE]),
                    ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
                    ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
                    ("LEFTPADDING",(0,0),(-1,-1),7),("VALIGN",(0,0),(-1,-1),"TOP"),
                ]))
                s.append(mt); s.append(_sp(1.5))
                s.append(_sec("Step-by-Step Session Guide", accent)); s.append(_sp(0.5))
                for stitle, sbody in act["steps"]:
                    s.append(Paragraph(f"<b>{stitle}</b>", ST["h2"]))
                    s.append(Paragraph(sbody.format(**{"child":child}), ST["body"]))
                    s.append(_sp(0.5))
                s.append(_sp())
                s.append(_sec("Adaptations and Troubleshooting", accent)); s.append(_sp(0.5))
                adata = [[Paragraph(a5,ST["bold_sm"]),Paragraph(b5.format(**{"child":child}),ST["body_sm"])] for a5,b5 in act["adapt"]]
                at2 = Table(adata, colWidths=[5*cm,12.5*cm])
                at2.setStyle(TableStyle([
                    ("ROWBACKGROUNDS",(0,0),(-1,-1),[light,C_WHITE]),
                    ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
                    ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
                    ("LEFTPADDING",(0,0),(-1,-1),7),("VALIGN",(0,0),(-1,-1),"TOP"),
                ]))
                s.append(at2)
                return s
            return content

        def make_conduct_chrome(accent=accent, anum=anum, aname=aname):
            def chrome(c2, doc2):
                _chrome(c2, doc2, f"Activity {anum}: {aname}", "Session Conduct Guide", accent, f"ACT {anum} - CONDUCT")
            return chrome

        temp_files.append(_make_temp(make_conduct_content(), make_conduct_chrome()))

        def make_tracking_content(act=act, accent=accent, light=light, anum=anum, aname=aname):
            def content():
                s = [_sp(4)]
                s.append(Paragraph(f"In-Session Tracking Sheet - Activity {anum}", ST["page_title"]))
                s.append(Paragraph(f"{aname}  |  CONFIDENTIAL - Therapist Use Only", ST["h3"]))
                s.append(_hr(accent))
                admin = [
                    [Paragraph("Child:",ST["bold_sm"]),Paragraph(child,ST["body_sm"]),Paragraph("Date:",ST["bold_sm"]),Paragraph("_____________________",ST["body_sm"])],
                    [Paragraph("Session #:",ST["bold_sm"]),Paragraph("___________",ST["body_sm"]),Paragraph("Duration:",ST["bold_sm"]),Paragraph("________ min",ST["body_sm"])],
                    [Paragraph("Therapist:",ST["bold_sm"]),Paragraph(therapist,ST["body_sm"]),Paragraph("Setting:",ST["bold_sm"]),Paragraph("_____________________",ST["body_sm"])],
                ]
                at3 = Table(admin, colWidths=[3*cm,5.5*cm,3*cm,6*cm])
                at3.setStyle(TableStyle([
                    ("BACKGROUND",(0,0),(-1,-1),light),("GRID",(0,0),(-1,-1),0.5,C_BORDER),
                    ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),("LEFTPADDING",(0,0),(-1,-1),7),
                ]))
                s.append(at3); s.append(_sp(1.5))
                s.append(_sec("Behavioral Observations", accent)); s.append(_sp(0.5))
                s.append(_track_table(act["tracking_rows"]))
                s.append(_sp(1.5))
                s.append(_sec("Session Notes and Observations", accent)); s.append(_sp(0.5))
                nt = Table([["_"*110]]*6, colWidths=[17.5*cm])
                nt.setStyle(TableStyle([
                    ("TEXTCOLOR",(0,0),(-1,-1),C_LIGHT_GREY),
                    ("TOPPADDING",(0,0),(-1,-1),8),("BOTTOMPADDING",(0,0),(-1,-1),8),
                    ("LEFTPADDING",(0,0),(-1,-1),7),("LINEBELOW",(0,0),(-1,-1),0.5,C_BORDER),
                ]))
                s.append(nt); s.append(_sp())
                prog = [
                    ["Progress indicator flagged today:",
                     "[ ] Increased latency  [ ] Faster recovery  [ ] Improved communication  [ ] Spontaneous replacement"],
                    ["Next session focus:", "______________________________________________"],
                ]
                pt3 = Table([[Paragraph(a6,ST["bold_sm"]),Paragraph(b6,ST["body_sm"])] for a6,b6 in prog], colWidths=[5*cm,12.5*cm])
                pt3.setStyle(TableStyle([
                    ("ROWBACKGROUNDS",(0,0),(-1,-1),[light,C_WHITE]),
                    ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
                    ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),("LEFTPADDING",(0,0),(-1,-1),7),
                ]))
                s.append(pt3)
                return s
            return content

        def make_tracking_chrome(accent=accent, anum=anum, aname=aname):
            def chrome(c2, doc2):
                _chrome(c2, doc2, f"Activity {anum}: {aname}", "In-Session Tracking Sheet  |  CONFIDENTIAL", accent, f"ACT {anum} - TRACKING")
            return chrome

        temp_files.append(_make_temp(make_tracking_content(), make_tracking_chrome()))

        def make_parent_content(act=act, accent=accent, light=light, anum=anum, aname=aname):
            def content():
                s = [_sp(4)]
                s.append(Paragraph("Parent Update Sheet", ST["page_title"]))
                s.append(Paragraph(f"Activity {anum}: {aname}  |  Date: ___________________", ST["h3"]))
                s.append(_hr(accent))
                s.append(Paragraph("Dear Family,", ST["h2"])); s.append(_sp())
                s.append(Paragraph(act["parent_what"].format(**{"child":child}), ST["parent"]))
                s.append(_sp(1.5))
                s.append(_sec("Why This Matters", accent)); s.append(_sp(0.5))
                s.append(Paragraph(act["parent_why"].format(**{"child":child}), ST["parent"]))
                s.append(_sp(1.5))
                s.append(_sec("How They Did Today", accent)); s.append(_sp(0.5))
                checks = ["Engaged enthusiastically","Engaged with some encouragement","Needed a lot of support today","Had a difficult session - see therapist note"]
                cells = [Paragraph(f"[ ]  {c}", ST["body_sm"]) for c in checks]
                cht = Table([cells[:2],cells[2:]], colWidths=[8.75*cm,8.75*cm])
                cht.setStyle(TableStyle([("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
                s.append(cht); s.append(_sp(1.5))
                s.append(_tip_box(act["parent_tips"], accent, light))
                s.append(_sp(1.5))
                sub_rows = [[Paragraph("NURSERY-FRIENDLY ALTERNATIVES FOR HOME", ST["label"])]]
                for sub in act["substitutes"]:
                    sub_rows.append([Paragraph("*  " + sub.format(**{"child":child}), ST["body_sm"])])
                subt = Table(sub_rows, colWidths=[17.5*cm])
                subt.setStyle(TableStyle([
                    ("BACKGROUND",(0,0),(0,0),C_MID_GREY),("BACKGROUND",(0,1),(0,-1),C_LIGHT_GREY),
                    ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
                    ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),("LEFTPADDING",(0,0),(-1,-1),8),
                ]))
                s.append(subt); s.append(_sp(1.5))
                s.append(_sec("Therapist Note", accent)); s.append(_sp(0.5))
                nt2 = Table([["_"*110]]*4, colWidths=[17.5*cm])
                nt2.setStyle(TableStyle([
                    ("TEXTCOLOR",(0,0),(-1,-1),C_LIGHT_GREY),
                    ("TOPPADDING",(0,0),(-1,-1),8),("BOTTOMPADDING",(0,0),(-1,-1),8),
                    ("LEFTPADDING",(0,0),(-1,-1),7),("LINEBELOW",(0,0),(-1,-1),0.5,C_BORDER),
                ]))
                s.append(nt2); s.append(_sp())
                s.append(Paragraph(
                    f"Questions? Please speak with {therapist} at pick-up or contact the nursery team. "
                    f"Your observations at home are a vital part of {child}'s progress.", ST["body_sm"]))
                return s
            return content

        def make_parent_chrome(accent=accent, anum=anum, aname=aname):
            def chrome(c2, doc2):
                _chrome(c2, doc2, f"Activity {anum}: {aname}", "Parent Update Sheet", accent, f"ACT {anum} - PARENT")
            return chrome

        temp_files.append(_make_temp(make_parent_content(), make_parent_chrome()))

    # Token board
    def make_token_board():
        tmp3 = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp3.close()
        lsz = landscape(A4)
        c3 = rl_canvas.Canvas(tmp3.name, pagesize=lsz)
        w3, h3 = lsz; lm3 = 1.5*cm; rm3 = w3-1.5*cm
        c3.setFillColor(C_TEAL); c3.rect(0, h3-2.0*cm, w3, 2.0*cm, fill=1, stroke=0)
        c3.setFillColor(C_WHITE); c3.setFont("Helvetica-Bold", 13)
        c3.drawString(1.5*cm, h3-1.3*cm, "Token Board")
        c3.setFont("Helvetica", 9)
        c3.drawString(1.5*cm, h3-1.75*cm, f"{child}  |  {nursery}  |  Print and laminate for repeated use")
        c3.setFillColor(C_LIGHT_GREY); c3.rect(0, 0, w3, 0.9*cm, fill=1, stroke=0)
        c3.setFillColor(C_MID_GREY); c3.setFont("Helvetica", 7)
        c3.drawCentredString(w3/2, 0.3*cm, f"CONFIDENTIAL - {nursery} - Therapist: {therapist}")
        panel_w3 = 6.5*cm
        c3.setFillColor(C_TEAL_LIGHT)
        c3.roundRect(lm3, 1.2*cm, panel_w3, h3-3.4*cm, 6, fill=1, stroke=0)
        c3.setStrokeColor(C_TEAL); c3.setLineWidth(1.5)
        c3.roundRect(lm3, 1.2*cm, panel_w3, h3-3.4*cm, 6, fill=0, stroke=1)
        px3=lm3+0.8*cm; py3=h3-5.8*cm; ps3=4.5*cm
        c3.setFillColor(C_WHITE); c3.roundRect(px3, py3, ps3, ps3, 4, fill=1, stroke=0)
        c3.setStrokeColor(C_TEAL); c3.setLineWidth(1); c3.roundRect(px3, py3, ps3, ps3, 4, fill=0, stroke=1)
        c3.setFillColor(C_MID_GREY); c3.setFont("Helvetica", 8)
        c3.drawCentredString(px3+ps3/2, py3+ps3/2+0.3*cm, "Child's")
        c3.drawCentredString(px3+ps3/2, py3+ps3/2-0.3*cm, "Photo")
        for lbl3, yo3 in [("Name:",1.2),("Date:",2.0),("Session #:",2.8)]:
            c3.setFillColor(C_DARK); c3.setFont("Helvetica-Bold", 10)
            c3.drawString(lm3+0.5*cm, py3-yo3*cm, lbl3)
            c3.setStrokeColor(C_TEAL); c3.setLineWidth(1)
            c3.line(lm3+2.8*cm, py3-yo3*cm, lm3+panel_w3-0.5*cm, py3-yo3*cm)
        c3.setFillColor(C_TEAL); c3.roundRect(lm3+0.3*cm, 2.0*cm, panel_w3-0.6*cm, 1.0*cm, 4, fill=1, stroke=0)
        c3.setFillColor(C_WHITE); c3.setFont("Helvetica-Bold", 9)
        c3.drawCentredString(lm3+panel_w3/2, 2.45*cm, "My Reward:")
        tax3=lm3+panel_w3+0.8*cm; taw3=rm3-tax3; tay3=1.2*cm; tah3=h3-3.4*cm
        c3.setFillColor(C_GOLD); c3.roundRect(tax3, tay3+tah3-1.4*cm, taw3, 1.4*cm, 4, fill=1, stroke=0)
        c3.setFillColor(C_DARK); c3.setFont("Helvetica-Bold", 14)
        c3.drawCentredString(tax3+taw3/2, tay3+tah3-0.85*cm, "My Tokens")
        cols_n3=5; rows_n3=2; cp3=0.4*cm
        gh3=tah3-1.4*cm-1.0*cm-0.5*cm
        cw3=(taw3-cp3*(cols_n3+1))/cols_n3; ch3=(gh3-cp3*(rows_n3+1))/rows_n3
        tok_cols3=[C_TEAL,C_GOLD,C_GREEN,C_PURPLE,C_CORAL]
        cell_n3=1
        for row3 in range(rows_n3):
            for col3 in range(cols_n3):
                cx3=tax3+cp3+col3*(cw3+cp3)
                cy3=tay3+gh3+0.8*cm-cp3-(row3+1)*ch3-row3*cp3
                col_c3=tok_cols3[col3%len(tok_cols3)]
                c3.setFillColor(col_c3); c3.setFillAlpha(0.1)
                c3.roundRect(cx3, cy3, cw3, ch3, 6, fill=1, stroke=0); c3.setFillAlpha(1.0)
                c3.setStrokeColor(col_c3); c3.setLineWidth(1.5)
                c3.roundRect(cx3, cy3, cw3, ch3, 6, fill=0, stroke=1)
                star_sz3=min(cw3,ch3)*0.38; scx3=cx3+cw3/2; scy3=cy3+ch3/2+0.3*cm
                pts3=[]
                for i4 in range(5):
                    oa3=math.radians(90+i4*72); ia3=math.radians(90+i4*72+36)
                    pts3.append((scx3+star_sz3*math.cos(oa3), scy3+star_sz3*math.sin(oa3)))
                    pts3.append((scx3+star_sz3*0.4*math.cos(ia3), scy3+star_sz3*0.4*math.sin(ia3)))
                path3=c3.beginPath(); path3.moveTo(pts3[0][0],pts3[0][1])
                for px4,py4 in pts3[1:]: path3.lineTo(px4,py4)
                path3.close()
                c3.setFillColor(col_c3); c3.setFillAlpha(0.15); c3.drawPath(path3,fill=1,stroke=1); c3.setFillAlpha(1.0)
                c3.setFillColor(col_c3); c3.setFont("Helvetica-Bold",11)
                c3.drawCentredString(scx3, cy3+0.35*cm, str(cell_n3)); cell_n3+=1
        c3.setFillColor(C_LIGHT_GREY); c3.rect(tax3, tay3, taw3, 0.8*cm, fill=1, stroke=0)
        c3.setFillColor(C_MID_GREY); c3.setFont("Helvetica", 7.5)
        c3.drawCentredString(tax3+taw3/2, tay3+0.25*cm,
            "Mark each token as it is earned. When all 10 tokens are collected, the reward is earned!")
        c3.save()
        return tmp3.name

    temp_files.append(make_token_board())

    # Materials summary
    def mat_content():
        s = [_sp(4)]
        s.append(Paragraph("Materials and Substitutes - All Activities", ST["page_title"]))
        s.append(Paragraph("Quick reference for session preparation", ST["h3"]))
        s.append(_sp(1.5))
        for i5, (act5, cnt5, _) in enumerate(schedule):
            acc5=act5["color"]; lgt5=act5["light"]
            s.append(_sec(f"Activity {i5+1}: {act5['name']}", acc5)); s.append(_sp(0.5))
            hdr5=[[Paragraph("<b>Item</b>",ST["bold_sm"]),Paragraph("<b>If unavailable - use instead</b>",ST["bold_sm"])]]
            for itm5,sub5 in act5["materials"]:
                hdr5.append([Paragraph(itm5,ST["body_sm"]),Paragraph(sub5,ST["body_sm"])])
            mt5=Table(hdr5, colWidths=[5*cm,12.5*cm])
            mt5.setStyle(TableStyle([
                ("BACKGROUND",(0,0),(-1,0),acc5),("TEXTCOLOR",(0,0),(-1,0),C_WHITE),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[lgt5,C_WHITE]),
                ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
                ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
                ("LEFTPADDING",(0,0),(-1,-1),7),("VALIGN",(0,0),(-1,-1),"TOP"),
            ]))
            s.append(mt5); s.append(_sp(1.5))
        s.append(HRFlowable(width="100%", thickness=1, color=C_BORDER))
        s.append(_sp(0.5))
        s.append(Paragraph(
            "<b>Note:</b> A token board and a set of preferred reinforcers (identified with the family "
            "before the first session) are required across ALL activities.", ST["body_sm"]))
        return s

    def mat_chrome(c2, doc2):
        _chrome(c2, doc2, "Materials and Substitutes", "All activities - quick reference", C_DARK, "MATERIALS")

    temp_files.append(_make_temp(mat_content, mat_chrome))

    # Merge all
    writer = PdfWriter()
    for tf in temp_files:
        reader = PdfReader(tf)
        for page in reader.pages:
            writer.add_page(page)
    out_buf = BytesIO()
    writer.write(out_buf)
    for tf in temp_files:
        try: os.unlink(tf)
        except: pass
    return out_buf.getvalue()


if __name__ == "__main__":
    main()
